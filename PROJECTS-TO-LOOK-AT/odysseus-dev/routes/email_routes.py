"""
email_routes.py

FastAPI route handlers for the email feature. All non-route logic
(IMAP connection helpers, message parsing, account config, the
auto-summarize + scheduled-email pollers, Pydantic models) lives in:

    routes/email_helpers.py   — synchronous helpers + models + constants
    routes/email_pollers.py   — background loops, started by `_start_poller`

Importing from the helpers module brings in everything those route
handlers need. The split is mechanical — no behavior change.
"""

import asyncio
import os
import sqlite3 as _sql3
import time
import email as email_mod
import email.header
import email.utils
import smtplib
import json
import re
import html
from html.parser import HTMLParser as _HTMLParser
import logging
import uuid
from datetime import datetime
from pathlib import Path

from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

from fastapi import APIRouter, Query, UploadFile, File, BackgroundTasks, HTTPException, Depends, Request
from fastapi.responses import FileResponse
from src.constants import DATA_DIR

from src.llm_core import llm_call_async
from src.upload_limits import read_upload_limited, EMAIL_COMPOSE_UPLOAD_MAX_BYTES

from routes.email_helpers import (
    _strip_think, _extract_reply, _apply_email_style_mechanics, require_owner, require_user, _assert_owns_account,
    _q, _attach_compose_uploads, _cleanup_compose_uploads,
    _load_settings, _save_settings, _get_email_config,
    _send_smtp_message, _smtp_security_mode,
    _IMAP_TIMEOUT_SECONDS, _open_imap_connection,
    make_oauth_state, verify_oauth_state,
    _imap_connect, _imap, _decode_header, _detect_sent_folder, _detect_drafts_folder,
    _extract_attachment_text, _list_attachments_from_msg,
    _extract_attachment_to_disk, _extract_html, _extract_text,
    _fetch_sender_thread_context, _pre_retrieve_context,
    _EMAIL_REPLY_SYS_PROMPT_BASE, _POOL_HOOKS,
    _friendly_email_auth_error,
    SendEmailRequest, ExtractStyleRequest,
    ATTACHMENTS_DIR, COMPOSE_UPLOADS_DIR, SCHEDULED_DB,
    attachment_extract_dir, _email_cache_owner_clause,
)
from routes.email_pollers import _start_poller

logger = logging.getLogger(__name__)

ODYSSEUS_MAIL_ORIGIN = "odysseus-ui"


def _email_tag_owner_aliases(account_id: str | None, owner: str = "") -> list[str]:
    aliases = [owner or ""]
    try:
        from core.database import SessionLocal as _SL, EmailAccount as _EA
        db = _SL()
        try:
            resolved_account_id = account_id
            if not resolved_account_id:
                try:
                    cfg = _get_email_config(None, owner=owner)
                    resolved_account_id = cfg.get("account_id") or None
                    aliases.extend([
                        cfg.get("imap_user") or "",
                        cfg.get("smtp_user") or "",
                        cfg.get("from_address") or "",
                    ])
                except Exception as _e:
                    logger.warning("Failed to resolve email account alias", exc_info=_e)
                    resolved_account_id = None
            row = db.get(_EA, resolved_account_id) if resolved_account_id else None
            if row:
                aliases.extend([row.owner or "", row.imap_user or "", row.from_address or ""])
        finally:
            db.close()
    except Exception as _e:
        logger.warning("Failed to load email aliases", exc_info=_e)
    out = []
    for a in aliases:
        a = (a or "").strip()
        if a not in out:
            out.append(a)
    return out or [""]


def _email_tag_owner_clause(account_id: str | None, owner: str = "") -> tuple[str, list[str]]:
    aliases = _email_tag_owner_aliases(account_id, owner)
    placeholders = ",".join("?" * len(aliases))
    # In configured multi-user mode, do not treat legacy owner='' rows as
    # visible to everyone. Single-user/unconfigured mode keeps legacy rows.
    if owner:
        return f"owner IN ({placeholders})", aliases
    return f"(owner IN ({placeholders}) OR owner IS NULL)", aliases


def _record_email_received_events(owner: str, account_id: str | None, folder: str, emails: list[dict]):
    """Baseline inbox messages, then fire `email_received` for new arrivals."""
    if not owner or (folder or "INBOX").upper() != "INBOX" or not emails:
        return
    try:
        from src.event_bus import fire_event
        account_key = (account_id or "default").strip() or "default"
        now = datetime.utcnow().isoformat() + "Z"
        keys = []
        for e in emails:
            key = (e.get("message_id") or e.get("uid") or "").strip()
            if key and key not in keys:
                keys.append(key)
        if not keys:
            return

        conn = _sql3.connect(SCHEDULED_DB)
        try:
            conn.execute(
                "CREATE TABLE IF NOT EXISTS email_event_seen ("
                "owner TEXT NOT NULL, account_key TEXT NOT NULL, folder TEXT NOT NULL, "
                "message_key TEXT NOT NULL, first_seen_at TEXT NOT NULL, "
                "PRIMARY KEY (owner, account_key, folder, message_key))"
            )
            count = conn.execute(
                "SELECT COUNT(*) FROM email_event_seen WHERE owner=? AND account_key=? AND folder=?",
                (owner, account_key, folder),
            ).fetchone()[0]
            existing = set()
            if count:
                placeholders = ",".join("?" * len(keys))
                rows = conn.execute(
                    f"SELECT message_key FROM email_event_seen "
                    f"WHERE owner=? AND account_key=? AND folder=? AND message_key IN ({placeholders})",
                    (owner, account_key, folder, *keys),
                ).fetchall()
                existing = {r[0] for r in rows}
            new_keys = [k for k in keys if k not in existing]
            conn.executemany(
                "INSERT OR IGNORE INTO email_event_seen "
                "(owner, account_key, folder, message_key, first_seen_at) VALUES (?, ?, ?, ?, ?)",
                [(owner, account_key, folder, k, now) for k in keys],
            )
            conn.commit()
        finally:
            conn.close()

        if count and new_keys:
            for _ in new_keys[:50]:
                fire_event("email_received", owner)
            logger.info("Fired email_received for %d new message(s)", min(len(new_keys), 50))
    except Exception:
        logger.debug("email_received event detection skipped", exc_info=True)


def _folder_name_from_list_line(line) -> str | None:
    decoded = line.decode() if isinstance(line, bytes) else str(line)
    match = re.search(r'"([^"]*)"\s*$|(\S+)\s*$', decoded)
    if not match:
        return None
    return match.group(1) or match.group(2)


def _list_imap_folders(conn) -> tuple[list, list[str]]:
    try:
        status, folders = conn.list()
        if status != "OK" or not folders:
            return [], []
        names = [name for name in (_folder_name_from_list_line(f) for f in folders) if name]
        return folders, names
    except Exception:
        return [], []


def _resolve_mail_folder(conn, preferred: str, role: str = "") -> str:
    """Resolve provider-specific names such as Gmail's [Gmail]/Bin/Spam."""
    folders, names = _list_imap_folders(conn)
    if preferred and preferred in names:
        return preferred
    role_flags = {
        "trash": ("\\Trash",),
        "archive": ("\\Archive", "\\All"),
        "junk": ("\\Junk",),
    }.get(role, ())
    for f in folders:
        decoded = f.decode() if isinstance(f, bytes) else str(f)
        if any(flag in decoded for flag in role_flags):
            name = _folder_name_from_list_line(f)
            if name:
                return name
    candidates = {
        "trash": ("Trash", "[Gmail]/Trash", "[Google Mail]/Trash", "Bin", "[Gmail]/Bin", "Deleted Messages", "Deleted Items"),
        "archive": ("Archive", "Archives", "[Gmail]/All Mail", "[Google Mail]/All Mail", "All Mail"),
        "junk": ("Junk", "Spam", "[Gmail]/Spam", "[Google Mail]/Spam"),
    }.get(role, ())
    lower_map = {n.lower(): n for n in names}
    for candidate in candidates:
        found = lower_map.get(candidate.lower())
        if found:
            return found
    return preferred


def _folder_role_from_name(name: str) -> str:
    lower = (name or "").lower()
    if "trash" in lower or "bin" in lower or "deleted" in lower:
        return "trash"
    if "spam" in lower or "junk" in lower:
        return "junk"
    if "archive" in lower or "all mail" in lower:
        return "archive"
    return ""


def _uid_bytes(uid: str | bytes) -> bytes:
    return uid if isinstance(uid, bytes) else str(uid).encode()


def _uid_exists(conn, uid: str) -> bool:
    try:
        status, data = conn.uid("FETCH", _uid_bytes(uid), "(UID)")
        if status != "OK":
            return False
        for part in data or []:
            meta = part[0] if isinstance(part, tuple) else part
            meta_b = meta if isinstance(meta, bytes) else str(meta).encode()
            if re.search(rb"\bUID\s+\d+\b", meta_b):
                return True
        return False
    except Exception:
        return False


def _imap_uid_search(conn, criteria: str):
    return conn.uid("SEARCH", None, criteria)


def _imap_uid_fetch(conn, uid_set: str | bytes, query: str):
    return conn.uid("FETCH", _uid_bytes(uid_set), query)


def _uid_from_fetch_meta(meta_b: bytes) -> str:
    m = re.search(rb"\bUID\s+(\d+)\b", meta_b)
    return m.group(1).decode() if m else ""


_FETCH_SEQ_RE = re.compile(rb"^(\d+)\s+\(")


def _group_uid_fetch_records(msg_data) -> list:
    """Group an imaplib UID FETCH response into per-message (meta, payload).

    imaplib yields an interleaved list: ``(meta, literal)`` tuples for
    attributes that carry a literal (``RFC822.HEADER {n}`` etc.) plus bare
    ``bytes`` elements for everything the server sends outside a literal.
    Where each attribute lands is server-specific: Dovecot sends FLAGS
    *before* the header literal (so it ends up inside the tuple meta), while
    Gmail sends FLAGS *after* it, arriving as a bare ``b' FLAGS (\\Seen))'``
    element. Dropping bare elements therefore silently loses FLAGS on Gmail
    and every message renders as unread/unflagged.

    A tuple whose meta starts with a sequence number opens a new record;
    every other part — continuation tuple or bare bytes — is folded into the
    current record's meta so attribute regexes see the full meta text.
    Plain ``b')'`` terminators get folded in too, which is harmless.
    """
    grouped: list = []  # list of (meta_bytes, payload_bytes_or_None)
    for part in (msg_data or []):
        if isinstance(part, tuple):
            meta_b = part[0] if isinstance(part[0], (bytes, bytearray)) else str(part[0]).encode()
            if _FETCH_SEQ_RE.match(meta_b):
                grouped.append((meta_b, part[1]))
            elif grouped:
                cur_meta, cur_payload = grouped[-1]
                grouped[-1] = (cur_meta + b" " + meta_b, cur_payload or part[1])
        elif isinstance(part, (bytes, bytearray)) and grouped:
            cur_meta, cur_payload = grouped[-1]
            grouped[-1] = (cur_meta + b" " + bytes(part), cur_payload)
    return grouped


def _smtp_ready(cfg: dict) -> bool:
    if not cfg.get("smtp_host") or not cfg.get("smtp_user"):
        return False
    return bool(cfg.get("smtp_password") or cfg.get("oauth_provider"))


def _resolve_send_config(account_id: str | None = None, owner: str = "") -> dict:
    """Resolve an account for outbound SMTP.

    If the caller explicitly picked an account, use only that account and
    return a clear error when it cannot send. If no account was picked and
    the default is receive-only, fall back to the first SMTP-capable account
    owned by the same user.
    """
    cfg = _get_email_config(account_id, owner=owner)
    if _smtp_ready(cfg):
        return cfg
    if account_id:
        raise ValueError(f"Email account {cfg.get('account_name') or account_id} has no SMTP configured")
    try:
        from core.database import SessionLocal as _SL, EmailAccount as _EA
        from sqlalchemy import and_, or_
        db = _SL()
        try:
            q = db.query(_EA).filter(_EA.enabled == True)  # noqa: E712
            if owner:
                unowned = or_(_EA.owner == None, _EA.owner == "")  # noqa: E711
                same_mailbox = or_(_EA.imap_user == owner, _EA.from_address == owner)
                q = q.filter(or_(_EA.owner == owner, and_(unowned, same_mailbox)))
            for row in q.order_by(_EA.is_default.desc(), _EA.created_at.asc()).all():
                trial = _get_email_config(account_id=row.id, owner=owner)
                if _smtp_ready(trial):
                    return trial
        finally:
            db.close()
    except Exception as e:
        logger.debug(f"SMTP-capable account fallback failed: {e}")
    raise ValueError("No SMTP-capable email account configured")


def _store_email_flag(conn, uid: str, flag: str, add: bool = True) -> bool:
    op = "+FLAGS" if add else "-FLAGS"
    if _uid_exists(conn, uid):
        status, _ = conn.uid("STORE", _uid_bytes(uid), op, flag)
    else:
        status, _ = conn.store(_uid_bytes(uid), op, flag)
    return status == "OK"


def _move_email_message(conn, uid: str, dest: str, role: str = "") -> bool:
    dest = _resolve_mail_folder(conn, dest, role or _folder_role_from_name(dest))
    if _uid_exists(conn, uid):
        status, _ = conn.uid("MOVE", _uid_bytes(uid), _q(dest))
        if status == "OK":
            return True
        status, _ = conn.uid("COPY", _uid_bytes(uid), _q(dest))
        if status != "OK":
            return False
        status, _ = conn.uid("STORE", _uid_bytes(uid), "+FLAGS", "\\Deleted")
    else:
        status, _ = conn.copy(_uid_bytes(uid), _q(dest))
        if status != "OK":
            return False
        status, _ = conn.store(_uid_bytes(uid), "+FLAGS", "\\Deleted")
    if status == "OK":
        conn.expunge()
        return True
    return False


def _apply_odysseus_headers(msg, kind: str | None = None, ref_id: str | None = None):
    msg["X-Odysseus-Origin"] = ODYSSEUS_MAIL_ORIGIN
    if kind:
        msg["X-Odysseus-Kind"] = re.sub(r"[^A-Za-z0-9_.-]", "-", kind)[:64]
    if ref_id:
        msg["X-Odysseus-Ref"] = re.sub(r"[^A-Za-z0-9_.:-]", "-", ref_id)[:128]


def _envelope_recipients(*fields: str) -> list:
    """Extract bare SMTP envelope addresses from one or more To/Cc/Bcc header
    strings. A naive `field.split(",")` corrupts display names that contain a
    comma (e.g. `"Smith, John" <john@corp.com>`, the canonical Outlook form):
    it splits into `"Smith` and `John" <john@corp.com>`, breaking delivery.
    email.utils.getaddresses parses the address grammar correctly."""
    out = []
    for _name, addr in email.utils.getaddresses([f for f in fields if f]):
        addr = (addr or "").strip()
        if addr:
            out.append(addr)
    return out


def _md_to_email_html(text: str) -> str:
    """Render the compose markdown body to a SAFE HTML fragment for the email's
    text/html part. Everything is HTML-escaped FIRST (so a pasted <script> /
    <img onerror=...> can never become live HTML in the recipient's client),
    then the toolbar's formatting is layered on with controlled regex: bold,
    italic, strike, inline code, http(s) links, headings, and bullet/numbered
    lists. Plain-text readers still get the raw markdown via the text/plain part.
    """
    def _inline(s: str) -> str:
        s = html.escape(s)                                  # escape BEFORE formatting
        s = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", s)
        s = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", s)
        s = re.sub(r"~~([^~]+)~~", r"<del>\1</del>", s)
        s = re.sub(r"`([^`]+)`", r"<code>\1</code>", s)
        # links: text + http(s) url only (escape() already neutralised quotes)
        s = re.sub(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", r'<a href="\2">\1</a>', s)
        return s

    parts: list[str] = []
    in_ul = in_ol = False
    for ln in (text or "").split("\n"):
        m_h = re.match(r"^(#{1,3})\s+(.*)$", ln)
        m_ul = re.match(r"^\s*[-*]\s+(.*)$", ln)
        m_ol = re.match(r"^\s*\d+\.\s+(.*)$", ln)
        if m_h:
            if in_ul: parts.append("</ul>"); in_ul = False
            if in_ol: parts.append("</ol>"); in_ol = False
            lvl = len(m_h.group(1))
            parts.append(f"<h{lvl}>{_inline(m_h.group(2))}</h{lvl}>")
        elif m_ul:
            if in_ol: parts.append("</ol>"); in_ol = False
            if not in_ul: parts.append("<ul>"); in_ul = True
            parts.append(f"<li>{_inline(m_ul.group(1))}</li>")
        elif m_ol:
            if in_ul: parts.append("</ul>"); in_ul = False
            if not in_ol: parts.append("<ol>"); in_ol = True
            parts.append(f"<li>{_inline(m_ol.group(1))}</li>")
        else:
            if in_ul: parts.append("</ul>"); in_ul = False
            if in_ol: parts.append("</ol>"); in_ol = False
            parts.append(_inline(ln) + "<br>")
    if in_ul: parts.append("</ul>")
    if in_ol: parts.append("</ol>")
    return "<html><body>" + "\n".join(parts) + "</body></html>"


# Tags the WYSIWYG email composer may legitimately produce.
_EMAIL_ALLOWED_TAGS = {
    "b", "strong", "i", "em", "u", "s", "strike", "del", "a", "br", "p", "div",
    "ul", "ol", "li", "blockquote", "span", "h1", "h2", "h3", "code", "pre",
}


class _EmailHtmlSanitizer(_HTMLParser):
    """Allowlist sanitizer for WYSIWYG-composed email HTML. Emits only known
    formatting tags (all attributes dropped except a safe href on <a>), escapes
    all text, and discards <script>/<style> content entirely — so client-sent
    HTML can never carry live script/handlers into the recipient's client."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.out = []
        self._skip = 0  # depth inside <script>/<style>

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._skip += 1
            return
        if tag == "br":
            self.out.append("<br>")
            return
        if tag not in _EMAIL_ALLOWED_TAGS:
            return
        if tag == "a":
            href = ""
            for k, v in attrs:
                if k.lower() == "href" and v and re.match(r"^(https?:|mailto:)", v.strip(), re.I):
                    href = v.strip()
            self.out.append(
                f'<a href="{html.escape(href, quote=True)}" target="_blank" rel="noopener noreferrer">'
                if href else "<a>")
        else:
            self.out.append(f"<{tag}>")

    def handle_startendtag(self, tag, attrs):
        if tag == "br":
            self.out.append("<br>")

    def handle_endtag(self, tag):
        if tag in ("script", "style"):
            if self._skip:
                self._skip -= 1
            return
        if tag == "br" or tag not in _EMAIL_ALLOWED_TAGS:
            return
        self.out.append(f"</{tag}>")

    def handle_data(self, data):
        if self._skip:
            return
        self.out.append(html.escape(data))


def _sanitize_email_html(raw: str) -> str:
    """Return a safe <html><body>…</body></html> from client-supplied compose
    HTML, or None if it can't be parsed."""
    p = _EmailHtmlSanitizer()
    try:
        p.feed(raw or "")
        p.close()
    except Exception:
        return None
    inner = "".join(p.out).strip()
    if not inner:
        return None
    return f"<html><body>{inner}</body></html>"


def setup_email_routes():
    _start_poller()
    router = APIRouter(prefix="/api/email", tags=["email"])

    # ── In-memory cache + prefetch + IMAP connection pool ──
    # Three layers stacked because every cold click was hitting Dovecot
    # over a fresh TCP+TLS+LOGIN handshake plus a full RFC822 fetch.
    #   1. _LIST_CACHE: list-emails responses keyed by (account, folder, filter,
    #      limit, offset). 8s TTL — short enough that flag changes show up
    #      quickly but long enough to absorb burst polls and tab switches.
    #   2. _READ_CACHE: per-(account, folder, uid) parsed email bodies.
    #      60s TTL — bodies don't change.
    #   3. _IMAP_POOL: per-account live IMAP connection reused across
    #      requests. Recycled if NOOP fails or it's been idle >60s.
    #   4. Prefetch task: after a list load, kick off background reads of
    #      the top-N visible UIDs so clicks land in the read cache.
    import asyncio as _asyncio
    import time as _time
    import threading as _threading

    _LIST_CACHE = {}  # key → (expires_at, response_dict)
    _LIST_TTL = 8.0
    _READ_CACHE = {}  # key → (expires_at, response_dict)
    _READ_TTL = 30 * 60.0
    _IMAP_POOL = {}   # account_id → (conn, last_used_at)
    _IMAP_IDLE_MAX = 60.0
    _WARMING_READS = set()
    _WARM_READ_LIMIT = 1
    _WARM_MAX_BYTES = 128 * 1024
    _WARM_RECENT_SECONDS = 7 * 24 * 60 * 60
    _pool_lock = _threading.Lock()

    def _pooled_connect(account_id, owner=""):
        """Reuse a live IMAP connection if one is in the pool and still
        responsive. Otherwise open fresh and store it. Caller must release
        via _pooled_release after use (not strictly required — the pool
        holds the same conn handle, and we lock to serialize access).

        SECURITY: `owner` is forwarded to `_imap_connect` so the fallback
        config lookup (when `account_id` is None) is scoped to this user's
        accounts only. The pool key is (account_id, owner) so two users
        with `account_id=None` don't share a pooled connection.
        """
        pool_key = (account_id, owner)
        now = _time.monotonic()
        with _pool_lock:
            entry = _IMAP_POOL.get(pool_key)
            if entry:
                conn, last_used = entry
                if (now - last_used) < _IMAP_IDLE_MAX:
                    try:
                        conn.noop()
                        # Pop it out of the pool while we use it (serialize)
                        del _IMAP_POOL[pool_key]
                        return conn, True  # reused
                    except Exception:
                        try: conn.logout()
                        except Exception: pass
                        del _IMAP_POOL[pool_key]
                else:
                    try: conn.logout()
                    except Exception: pass
                    del _IMAP_POOL[pool_key]
        # Fresh connection
        return _imap_connect(account_id, owner=owner), False

    def _pooled_release(account_id, conn, ok=True, owner=""):
        # SECURITY: match the (account_id, owner) key used by _pooled_connect
        # so a pooled handle is returned to the same per-user slot.
        if not ok:
            try: conn.logout()
            except Exception: pass
            return
        with _pool_lock:
            _IMAP_POOL[(account_id, owner)] = (conn, _time.monotonic())

    def _list_cache_key(account_id, folder, filter_, limit, offset, from_addr=""):
        return (account_id or "", folder, filter_, int(limit), int(offset), from_addr or "")

    def _read_cache_key(account_id, folder, uid, owner=""):
        # SECURITY: include owner so two users with `account_id == ""` /
        # None (i.e. resolved through the per-user default) don't share
        # a cached message body.
        return (account_id or "", folder, str(uid), owner)

    def _list_cache_get(key):
        v = _LIST_CACHE.get(key)
        if not v: return None
        if v[0] < _time.monotonic():
            _LIST_CACHE.pop(key, None)
            return None
        return v[1]

    def _list_cache_put(key, value):
        _LIST_CACHE[key] = (_time.monotonic() + _LIST_TTL, value)
        # Cap size
        if len(_LIST_CACHE) > 64:
            for k in list(_LIST_CACHE.keys())[:-32]:
                _LIST_CACHE.pop(k, None)

    def _invalidate_list_cache(account_id=None, folder=None):
        """Drop list cache entries that the caller's mutation may have stale-ed.

        Called from flag-mutating endpoints (mark-read/unread/answered, archive,
        delete, move) so the UI doesn't show stale read/unread counts for up to
        the 8s TTL after a manual flag change. With no args, clears everything.
        """
        if account_id is None and folder is None:
            _LIST_CACHE.clear()
            return
        for k in list(_LIST_CACHE.keys()):
            k_acct = k[0] if len(k) > 0 else ""
            k_folder = k[1] if len(k) > 1 else ""
            if (account_id is None or k_acct == (account_id or "")) and \
               (folder is None or k_folder == folder):
                _LIST_CACHE.pop(k, None)

    def _read_cache_get(key):
        v = _READ_CACHE.get(key)
        if not v: return None
        if v[0] < _time.monotonic():
            _READ_CACHE.pop(key, None)
            return None
        return v[1]

    def _read_cache_put(key, value):
        _READ_CACHE[key] = (_time.monotonic() + _READ_TTL, value)
        if len(_READ_CACHE) > 256:
            for k in list(_READ_CACHE.keys())[:-128]:
                _READ_CACHE.pop(k, None)

    # Expose helpers in the closure to be used by handlers below
    router._email_pool = {
        "connect": _pooled_connect,
        "release": _pooled_release,
        "list_cache_get": _list_cache_get,
        "list_cache_put": _list_cache_put,
        "list_cache_key": _list_cache_key,
        "read_cache_get": _read_cache_get,
        "read_cache_put": _read_cache_put,
        "read_cache_key": _read_cache_key,
    }
    # Wire the module-level _imap() context manager into the pool so every
    # `with _imap(account_id, owner=owner) as conn:` reuses an existing connection
    # instead of paying TCP+TLS+LOGIN per request.
    _POOL_HOOKS["connect"] = _pooled_connect
    _POOL_HOOKS["release"] = _pooled_release

    def _list_emails_sync(folder, limit, offset, filter_, account_id, from_addr=None, has_attachments_only=False, owner=""):
        """Sync IMAP work — call from async handler via asyncio.to_thread so
        it doesn't block the event loop.

        When `has_attachments_only` is True, IMAP doesn't have a portable
        HASATTACH keyword, so we widen the fetch (up to ~400 most-recent
        UIDs in the folder slice) and post-filter by Content-Type. Total
        count then reflects matches in that scanned window, not the whole
        folder.

        SECURITY: `owner` is propagated so when `account_id` is missing,
        the fallback config lookup is scoped to this user's accounts only.
        """
        conn = None
        try:
            conn = _imap_connect(account_id, owner=owner)
            select_status, _ = conn.select(_q(folder), readonly=True)
            if select_status != "OK":
                return {"emails": [], "total": 0, "folder": folder, "error": f"Folder not found: {folder}"}

            from_clause = ""
            if from_addr:
                # Escape quotes/backslashes for IMAP SEARCH FROM
                _safe = from_addr.replace("\\", "\\\\").replace('"', '\\"')
                from_clause = f' FROM "{_safe}"'

            if filter_ == "unread":
                status, data = _imap_uid_search(conn, f"(UNSEEN{from_clause})")
            elif filter_ == "favorites":
                # Flagged/favorited emails (the star toggle sets the \Flagged flag).
                status, data = _imap_uid_search(conn, f"(FLAGGED{from_clause})")
            elif filter_ == "unanswered":
                status, data = _imap_uid_search(conn, f"(UNSEEN UNANSWERED{from_clause})")
            elif filter_ == "undone":
                # All emails NOT marked as answered/done (read or unread).
                status, data = _imap_uid_search(conn, f"(UNANSWERED{from_clause})")
            elif filter_ == "reminders":
                # Prefer the Odysseus marker header, but include the subject
                # fallback too. The fallback uses a distinct Odysseus prefix
                # so ordinary emails containing "Reminder" don't get mixed in.
                status, data = _imap_uid_search(
                    conn,
                    f'(OR HEADER X-Odysseus-Kind "reminder" SUBJECT "Reminder (Odysseus):"{from_clause})',
                )
            elif filter_ == "pending_30d":
                # "What's pending in the last month" — UNANSWERED + delivered
                # within the last 30 days. SINCE takes a DD-Mon-YYYY date.
                from datetime import datetime as _dt, timedelta as _td
                _since = (_dt.utcnow() - _td(days=30)).strftime("%d-%b-%Y")
                status, data = _imap_uid_search(conn, f'(UNANSWERED SINCE "{_since}"{from_clause})')
            elif filter_ == "stale_30d":
                # "What's been sitting too long" — UNANSWERED + delivered
                # MORE than 30 days ago. BEFORE excludes the cutoff date itself.
                from datetime import datetime as _dt, timedelta as _td
                _before = (_dt.utcnow() - _td(days=30)).strftime("%d-%b-%Y")
                status, data = _imap_uid_search(conn, f'(UNANSWERED BEFORE "{_before}"{from_clause})')
            elif filter_ and filter_.startswith("tag:"):
                # Tag-based filter — resolve UIDs from email_tags first, then
                # ask IMAP for those messages by Message-ID. `tag:spam` reads
                # spam_verdict=1; any other tag matches JSON-array membership
                # in `tags`.
                _tag_name = filter_[len("tag:"):].strip().lower()
                _tag_message_ids = []
                _tag_seq_fallback = []
                try:
                    import sqlite3 as _sql3t
                    _ct = _sql3t.connect(SCHEDULED_DB)
                    _owner_clause, _owner_params = _email_tag_owner_clause(account_id, owner)
                    # SECURITY: owner-scope the lookup (review C2/H8). Without
                    # this, user A's `tag:urgent` filter would surface UIDs
                    # written by user B and IMAP would return whatever
                    # happens to live at those UIDs in A's mailbox. Account
                    # mailbox aliases are included because the background
                    # urgency task may be owned by the mailbox address while
                    # the UI is owned by the app user.
                    if _tag_name == "spam":
                        rows_t = _ct.execute(
                            "SELECT message_id, uid FROM email_tags "
                            "WHERE folder=? AND spam_verdict=1 "
                            f"AND {_owner_clause}",
                            (folder, *_owner_params),
                        ).fetchall()
                        for mid, uid in rows_t:
                            if mid:
                                _tag_message_ids.append(str(mid).strip())
                            elif uid:
                                _tag_seq_fallback.append(str(uid).strip())
                    else:
                        rows_t = _ct.execute(
                            "SELECT message_id, uid, tags FROM email_tags "
                            "WHERE folder=? AND tags IS NOT NULL AND tags != '' "
                            f"AND {_owner_clause}",
                            (folder, *_owner_params),
                        ).fetchall()
                        for r in rows_t:
                            try:
                                tg = json.loads(r[2] or "[]")
                                wanted = {_tag_name}
                                if _tag_name == "marketing":
                                    wanted.add("promo")
                                row_tags = {str(t).strip().lower().replace("_", "-") for t in tg} if isinstance(tg, list) else set()
                                if wanted.intersection(row_tags):
                                    if r[0]:
                                        _tag_message_ids.append(str(r[0]).strip())
                                    elif r[1]:
                                        _tag_seq_fallback.append(str(r[1]).strip())
                            except Exception:
                                continue
                    _ct.close()
                except Exception as _te:
                    logger.warning(f"tag filter lookup failed: {_te}")
                if not _tag_message_ids and not _tag_seq_fallback:
                    conn.logout()
                    return {"emails": [], "total": 0, "folder": folder}
                # Prefer stable Message-ID rows. Older tag rows may have only
                # numeric ids; those were sequence numbers historically, but
                # may be real UIDs for newer rows. Treat them as UIDs only.
                def _imap_search_quote(value: str) -> str:
                    return '"' + str(value or "").replace("\\", "\\\\").replace('"', '\\"') + '"'
                _uids = set()
                for _mid in dict.fromkeys(_tag_message_ids):
                    if not _mid:
                        continue
                    st_m, data_m = _imap_uid_search(conn, f'(HEADER Message-ID {_imap_search_quote(_mid)}{from_clause})')
                    if st_m == "OK" and data_m and data_m[0]:
                        _uids.update(data_m[0].split())
                for _uid in _tag_seq_fallback:
                    if _uid:
                        _uids.add(str(_uid).encode())
                if not _uids:
                    conn.logout()
                    return {"emails": [], "total": 0, "folder": folder}
                data = [b" ".join(sorted(_uids, key=lambda x: int(x) if str(x, "ascii", "ignore").isdigit() else 0))]
                status = "OK"
            elif from_clause:
                status, data = _imap_uid_search(conn, f"({from_clause.strip()})")
            else:
                status, data = _imap_uid_search(conn, "ALL")

            if status != "OK" or not data[0]:
                conn.logout()
                return {"emails": [], "total": 0, "folder": folder}

            uid_list = data[0].split()
            total = len(uid_list)
            # Reverse for newest first, apply pagination
            uid_list = list(reversed(uid_list))
            if has_attachments_only:
                # Can't filter via IMAP — widen the window so post-filter
                # still yields enough rows to fill `limit` after dropping
                # rows without attachments.
                scan_window = max(400, offset + limit * 8)
                uid_list = uid_list[:scan_window]
            else:
                uid_list = uid_list[offset:offset + limit]

            # Preload tag rows once — keyed by uid (as str) for the emails we'll render
            _tag_by_uid = {}
            try:
                import sqlite3 as _sql3
                _c = _sql3.connect(SCHEDULED_DB)
                _uid_strs = [u.decode() for u in uid_list]
                if _uid_strs:
                    placeholders = ",".join("?" * len(_uid_strs))
                    _owner_clause, _owner_params = _email_tag_owner_clause(account_id, owner)
                    rows = _c.execute(
                        f"SELECT uid, tags, spam_verdict FROM email_tags "
                        f"WHERE folder=? AND {_owner_clause} AND uid IN ({placeholders})",
                        [folder, *_owner_params, *_uid_strs],
                    ).fetchall()
                    for r in rows:
                        try:
                            tg = json.loads(r[1] or "[]")
                        except Exception:
                            tg = []
                        if isinstance(tg, list):
                            tg = ["marketing" if str(t).strip().lower().replace("_", "-") == "promo" else t for t in tg]
                        _tag_by_uid[r[0]] = {"tags": tg, "spam": bool(r[2])}
                _c.close()
            except Exception as e:
                logger.warning(f"Tag preload failed: {e}")

            # Batch fetch ALL requested UIDs in a single IMAP round-trip.
            # Per-UID fetch was the dominant cost — N round-trips × (~5-20ms
            # each on localhost) made 50-message lists take 250ms-1s+. The
            # batched form trades a slightly bigger response for one round-trip.
            emails = []
            if uid_list:
                fetch_set = b",".join(uid_list)
                try:
                    status, msg_data = _imap_uid_fetch(conn, fetch_set, "(UID FLAGS RFC822.HEADER RFC822.SIZE)")
                except Exception as e:
                    logger.warning(f"Batch fetch failed, falling back to per-UID: {e}")
                    status, msg_data = "NO", []
                # Group the batched response into per-message (meta, payload)
                # records. Bare bytes parts must be kept: Gmail returns FLAGS
                # after the header literal as a bare element, and dropping it
                # rendered every Gmail message as unread/unflagged.
                grouped = _group_uid_fetch_records(msg_data)

                if status != "OK" and not grouped:
                    conn.logout()
                    return {"emails": [], "total": total, "folder": folder, "offset": offset}

                _tag_by_message_id = {}
                try:
                    header_ids = []
                    for _, raw_header in grouped:
                        if not raw_header:
                            continue
                        mid = (email_mod.message_from_bytes(raw_header).get("Message-ID", "") or "").strip()
                        if mid:
                            header_ids.append(mid)
                    if header_ids:
                        import sqlite3 as _sql3m
                        _cm = _sql3m.connect(SCHEDULED_DB)
                        _owner_clause_m, _owner_params_m = _email_tag_owner_clause(account_id, owner)
                        _mid_ph = ",".join("?" * len(header_ids))
                        rows_m = _cm.execute(
                            f"SELECT message_id, tags, spam_verdict FROM email_tags "
                            f"WHERE folder=? AND {_owner_clause_m} "
                            f"AND message_id IN ({_mid_ph})",
                            [folder, *_owner_params_m, *header_ids],
                        ).fetchall()
                        _cm.close()
                        for mid, tags_raw, spam_raw in rows_m:
                            try:
                                tags = json.loads(tags_raw or "[]")
                            except Exception:
                                tags = []
                            if isinstance(tags, list):
                                tags = ["marketing" if str(t).strip().lower().replace("_", "-") == "promo" else t for t in tags]
                            _tag_by_message_id[(mid or "").strip()] = {
                                "tags": tags if isinstance(tags, list) else [],
                                "spam": bool(spam_raw),
                            }
                except Exception as e:
                    logger.warning(f"Message-ID tag preload failed: {e}")

                for meta_b, raw_header in grouped:
                    try:
                        meta = meta_b.decode(errors="replace")
                        uid_num = _uid_from_fetch_meta(meta_b)
                        if not uid_num:
                            continue
                        flag_m = re.search(r'FLAGS \(([^)]*)\)', meta)
                        flags = flag_m.group(1) if flag_m else ""
                        size_m = re.search(r'RFC822\.SIZE (\d+)', meta)
                        size = int(size_m.group(1)) if size_m else 0
                        if not raw_header:
                            continue

                        msg = email_mod.message_from_bytes(raw_header)
                        subject = _decode_header(msg.get("Subject", "(no subject)"))
                        sender = _decode_header(msg.get("From", "unknown"))
                        date_str = msg.get("Date", "")
                        message_id = msg.get("Message-ID", "")
                        sender_name, sender_addr = email.utils.parseaddr(sender)
                        # To/Cc — needed for the from-sender sidebar's
                        # multi-tag filter ("emails involving ALL these
                        # people"). Decoded raw strings; client splits.
                        to_str = _decode_header(msg.get("To", ""))
                        cc_str = _decode_header(msg.get("Cc", ""))
                        parsed_date = email.utils.parsedate_to_datetime(date_str) if date_str else None
                        # Normalise tz-naive parses to UTC so timestamp() is
                        # deterministic across hosts.
                        if parsed_date and parsed_date.tzinfo is None:
                            from datetime import timezone as _tz
                            parsed_date = parsed_date.replace(tzinfo=_tz.utc)
                        iso_date = parsed_date.isoformat() if parsed_date else ""
                        date_epoch = parsed_date.timestamp() if parsed_date else 0.0
                        is_read = "\\Seen" in flags
                        is_answered = "\\Answered" in flags
                        is_flagged = "\\Flagged" in flags
                        ct = msg.get("Content-Type", "")
                        has_attachments = "multipart/mixed" in ct.lower() or "multipart/related" in ct.lower()
                        tag_entry = _tag_by_message_id.get(message_id.strip()) or _tag_by_uid.get(uid_num, {})
                        emails.append({
                            "uid": uid_num,
                            "message_id": message_id.strip(),
                            "subject": subject,
                            "from_name": sender_name or sender_addr,
                            "from_address": sender_addr,
                            "to": to_str,
                            "cc": cc_str,
                            "date": iso_date,
                            "date_display": date_str,
                            "date_epoch": date_epoch,
                            "size": size,
                            "is_read": is_read,
                            "is_answered": is_answered,
                            "is_flagged": is_flagged,
                            "flags": flags,
                            "has_attachments": has_attachments,
                            "tags": tag_entry.get("tags", []),
                            "is_spam_verdict": tag_entry.get("spam", False),
                        })
                    except Exception as e:
                        logger.warning(f"Error parsing batched email entry: {e}")
                        continue
                # IMAP returns batched results in seq-set order, not the
                # newest-first order we want. Sort by the parsed UTC epoch
                # so cross-timezone dates compare chronologically (ISO-string
                # sort had `+02:00` beating `+00:00` at the same local time).
                emails.sort(key=lambda x: x.get("date_epoch") or 0.0, reverse=True)

            if has_attachments_only:
                emails = [e for e in emails if e.get("has_attachments")]
                # Total now reflects matches inside the scanned window, not
                # the whole folder — see scan_window above.
                total = len(emails)
                emails = emails[offset:offset + limit]

            # Bulk-attach cached AI summaries by Message-ID so the frontend
            # can show them on hover (avoids a per-card round-trip).
            try:
                ids = [e.get("message_id", "") for e in emails if e.get("message_id")]
                if ids:
                    import sqlite3 as _sql3
                    _c = _sql3.connect(SCHEDULED_DB)
                    placeholders = ",".join("?" * len(ids))
                    owner_clause, owner_params = _email_cache_owner_clause(owner)
                    rows = _c.execute(
                        f"SELECT message_id, summary FROM email_summaries "
                        f"WHERE message_id IN ({placeholders}) AND {owner_clause}",
                        (*ids, *owner_params),
                    ).fetchall()
                    _c.close()
                    by_id = {r[0]: r[1] for r in rows}
                    for e in emails:
                        s = by_id.get(e.get("message_id", ""))
                        if s:
                            e["cached_summary"] = s
            except Exception as _summary_err:
                logger.debug(f"Bulk summary attach skipped: {_summary_err}")

            return {"emails": emails, "total": total, "folder": folder, "offset": offset}
        except Exception as e:
            logger.error(f"Failed to list emails: {e}")
            detail = str(e).strip()
            return {"emails": [], "total": 0, "error": f"Mail operation failed: {detail[:180]}" if detail else "Mail operation failed"}
        finally:
            if conn:
                try:
                    conn.logout()
                except Exception:
                    pass

    @router.get("/list")
    async def list_emails(
        folder: str = Query("INBOX"),
        limit: int = Query(50),
        offset: int = Query(0),
        filter: str = Query("all"),  # all, unread, unanswered
        from_addr: str | None = Query(None, alias="from"),
        account_id: str | None = Query(None),
        has_attachments: int = Query(0),
        cache_bust: str | None = Query(None, alias="_"),
        owner: str = Depends(require_owner),
    ):
        """List emails. Uses an 8s in-memory cache + offloads blocking IMAP
        calls to a worker thread so the event loop never stalls."""
        _deferred = getattr(_start_poller, '_deferred', None)
        if _deferred:
            await _deferred()
        # SECURITY: include `owner` in the cache key so two users with
        # different account scopes don't share a cached list.
        ck = _list_cache_key(account_id, folder, filter, limit, offset, from_addr or "") + (int(bool(has_attachments)), owner)
        if not cache_bust:
            cached = _list_cache_get(ck)
            if cached is not None:
                _schedule_recent_email_warm(cached.get("emails") or [], folder, account_id, owner)
                return cached
        result = await _asyncio.to_thread(
            _list_emails_sync, folder, limit, offset, filter, account_id, from_addr,
            bool(has_attachments), owner,
        )
        if result and not result.get("error"):
            if offset == 0 and not from_addr and not has_attachments and filter in ("all", "unread", "unanswered", "undone"):
                _record_email_received_events(owner, account_id, folder, result.get("emails") or [])
                _schedule_recent_email_warm(result.get("emails") or [], folder, account_id, owner)
            _list_cache_put(ck, result)
        return result

    @router.post("/{uid}/unflag-spam")
    async def unflag_spam(uid: str, owner: str = Depends(require_owner)):
        """User override — mark email as not spam."""
        try:
            owner_clause, owner_params = _email_tag_owner_clause(None, owner)
            _c = _sql3.connect(SCHEDULED_DB)
            _c.execute(
                f"UPDATE email_tags SET spam_verdict=0, spam_reason='' WHERE uid=? AND {owner_clause}",
                [uid, *owner_params],
            )
            _c.commit()
            _c.close()
            return {"ok": True}
        except Exception as e:
            logger.error(f"unflag-spam failed: {e}")
            return {"ok": False, "error": "Mail operation failed"}

    @router.get("/contacts")
    async def list_contacts(
        q: str = Query(""),
        limit: int = Query(20),
        owner: str = Depends(require_owner),
    ):
        """Distinct name/address pairs aggregated from the email_tags table
        — used by the from-sender sidebar's autocomplete to convert typed
        names into chips. Backed by the AI-classification cache so it's a
        cheap SQL read; people you've never received a tagged email from
        won't appear yet."""
        ql = (q or "").strip().lower()
        try:
            conn = _sql3.connect(SCHEDULED_DB)
            owner_clause, owner_params = _email_tag_owner_clause(None, owner)
            rows = conn.execute(
                f"SELECT sender FROM email_tags WHERE sender IS NOT NULL AND sender != '' AND {owner_clause}",
                owner_params,
            ).fetchall()
            conn.close()
            seen = {}
            for (s,) in rows:
                try:
                    name, addr = email.utils.parseaddr(s or "")
                except Exception:
                    continue
                if not addr:
                    continue
                addr_l = addr.lower()
                if ql and ql not in (name or "").lower() and ql not in addr_l:
                    continue
                if addr_l in seen:
                    continue
                seen[addr_l] = {"name": (name or addr).strip(), "address": addr}
            items = list(seen.values())
            # Prefer entries whose name starts with the query, then alphabetical.
            items.sort(key=lambda c: (
                0 if ql and (c["name"] or "").lower().startswith(ql) else 1,
                (c["name"] or c["address"]).lower(),
            ))
            return {"contacts": items[: max(1, int(limit))]}
        except Exception as e:
            logger.error(f"contacts list failed: {e}")
            return {"contacts": [], "error": "Mail operation failed"}

    @router.get("/search")
    # Sync def: the body is blocking IMAP I/O with no awaits. As `async def` it ran
    # directly on the event loop and stalled the whole app during a search; as a sync
    # def FastAPI runs it in a threadpool, keeping the loop responsive.
    def search_emails(
        q: str = Query(""),
        folder: str = Query("INBOX"),
        limit: int = Query(50),
        account_id: str | None = Query(None),
        owner: str = Depends(require_owner),
    ):
        """Search emails server-side via IMAP SEARCH. Matches subject, from, or body text.

        When the caller asks for INBOX and the account has an "All Mail"
        folder (Gmail does), we transparently swap to All Mail so the
        search surfaces archived / labelled emails too. Plain IMAP
        accounts fall back to whatever folder the caller specified."""
        if not q or len(q) < 2:
            return {"emails": [], "total": 0, "query": q}
        # CRLF in q would terminate the IMAP command early — reject defensively.
        if "\r" in q or "\n" in q:
            raise HTTPException(400, "Invalid query")
        try:
            with _imap(account_id, owner=owner) as conn:
                # If the user asked for INBOX, try to upgrade to All Mail —
                # one folder == every email on Gmail-class servers.
                effective_folder = folder
                if (folder or "").upper() == "INBOX":
                    try:
                        status, folder_lines = conn.list()
                        if status == "OK" and folder_lines:
                            for raw in folder_lines:
                                if isinstance(raw, bytes):
                                    raw = raw.decode("utf-8", errors="replace")
                                m = re.match(r"\((?P<flags>[^)]*)\)\s+\"[^\"]*\"\s+(?P<name>.+)", raw)
                                if not m:
                                    continue
                                flags = (m.group("flags") or "").lower()
                                name = m.group("name").strip().strip('"')
                                if "\\all" in flags or "all mail" in name.lower():
                                    effective_folder = name
                                    break
                    except Exception:
                        pass
                conn.select(_q(effective_folder), readonly=True)

                # Escape backslash and quote for the IMAP-SEARCH quoted-string.
                q_escaped = q.replace('\\', '\\\\').replace('"', '\\"')
                search_cmd = f'(OR OR FROM "{q_escaped}" SUBJECT "{q_escaped}" TEXT "{q_escaped}")'

                status, data = _imap_uid_search(conn, search_cmd)
                if status != "OK" or not data[0]:
                    return {"emails": [], "total": 0, "query": q, "folder": effective_folder}

                uid_list = data[0].split()
                total = len(uid_list)
                uid_list = list(reversed(uid_list))[:limit]

                emails = []
                for uid in uid_list:
                    try:
                        status, msg_data = _imap_uid_fetch(conn, uid, "(UID FLAGS RFC822.HEADER)")
                        if status != "OK":
                            continue
                        raw_header = None
                        flags = ""
                        # Same Gmail caveat as the list route: FLAGS may
                        # arrive after the header literal, so group bare
                        # parts back into the message meta before scanning.
                        for meta_b, payload in _group_uid_fetch_records(msg_data):
                            if payload and b"RFC822.HEADER" in meta_b:
                                raw_header = payload
                            flag_match = re.search(rb'FLAGS \(([^)]*)\)', meta_b)
                            if flag_match:
                                flags = flag_match.group(1).decode(errors="replace")
                        if not raw_header:
                            continue
                        msg = email_mod.message_from_bytes(raw_header)
                        subject = _decode_header(msg.get("Subject", "(no subject)"))
                        sender = _decode_header(msg.get("From", "unknown"))
                        date_str = msg.get("Date", "")
                        message_id = msg.get("Message-ID", "")
                        sender_name, sender_addr = email.utils.parseaddr(sender)
                        to_str = _decode_header(msg.get("To", ""))
                        cc_str = _decode_header(msg.get("Cc", ""))
                        parsed_date = email.utils.parsedate_to_datetime(date_str) if date_str else None
                        if parsed_date and parsed_date.tzinfo is None:
                            from datetime import timezone as _tz
                            parsed_date = parsed_date.replace(tzinfo=_tz.utc)
                        iso_date = parsed_date.isoformat() if parsed_date else ""
                        date_epoch = parsed_date.timestamp() if parsed_date else 0.0
                        ct = msg.get("Content-Type", "")
                        has_attachments = "multipart/mixed" in ct.lower() or "multipart/related" in ct.lower()

                        stable_uid = ""
                        for part in msg_data:
                            if isinstance(part, tuple):
                                meta_b = part[0] if isinstance(part[0], bytes) else str(part[0]).encode()
                                stable_uid = _uid_from_fetch_meta(meta_b) or stable_uid
                        if not stable_uid:
                            continue
                        emails.append({
                            "uid": stable_uid,
                            "message_id": message_id.strip(),
                            "subject": subject,
                            "from_name": sender_name or sender_addr,
                            "from_address": sender_addr,
                            "to": to_str,
                            "cc": cc_str,
                            "date": iso_date,
                            "date_display": date_str,
                            "date_epoch": date_epoch,
                            "is_read": "\\Seen" in flags,
                            "is_answered": "\\Answered" in flags,
                            "is_flagged": "\\Flagged" in flags,
                            "flags": flags,
                            "has_attachments": has_attachments,
                            # Stamp the folder so the frontend opens each
                            # email from the folder it actually lives in
                            # (the search may have run against All Mail
                            # even though the caller asked for INBOX),
                            # otherwise clicks open whatever happens to
                            # have the same UID in INBOX → wrong email.
                            "folder": effective_folder,
                        })
                    except Exception as e:
                        logger.warning(f"Error parsing search result {uid}: {e}")
                        continue

                return {"emails": emails, "total": total, "query": q}
        except Exception as e:
            logger.error(f"Search failed: {e}")
            return {"emails": [], "total": 0, "error": "Mail operation failed"}

    def _read_email_sync(uid, folder, account_id, owner, mark_seen=True):
        """Sync IMAP read — wrapped in to_thread by the async handler.

        Two-phase: read body in readonly to avoid races with concurrent reads
        of the same UID, then flip \\Seen in a separate readwrite session.
        BODY.PEEK[] keeps the fetch itself from tripping \\Seen.
        """
        import time as _t
        _t0 = _t.monotonic()
        raw = None
        _t_select = 0.0
        _t_fetch = 0.0
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                _t_select = _t.monotonic() - _t0
                status, msg_data = _imap_uid_fetch(conn, uid, "(BODY.PEEK[])")
                _t_fetch = _t.monotonic() - _t0
                if status != "OK":
                    return {"error": f"Email UID {uid} not found"}
                raw = msg_data[0][1]

            msg = email_mod.message_from_bytes(raw)

            subject = _decode_header(msg.get("Subject", "(no subject)"))
            sender = _decode_header(msg.get("From", "unknown"))
            to = _decode_header(msg.get("To", ""))
            cc = _decode_header(msg.get("Cc", ""))
            date_str = msg.get("Date", "")
            message_id = msg.get("Message-ID", "")
            in_reply_to = msg.get("In-Reply-To", "")
            references = msg.get("References", "")
            body = _extract_text(msg)
            body_html = _extract_html(msg)

            sender_name, sender_addr = email.utils.parseaddr(sender)
            parsed_date = email.utils.parsedate_to_datetime(date_str) if date_str else None
            attachments = _list_attachments_from_msg(msg)

            if mark_seen:
                # Set \Seen in a separate readwrite session so concurrent reads
                # of the same UID don't fight over a shared SELECT state.
                try:
                    with _imap(account_id, owner=owner) as conn2:
                        conn2.select(_q(folder))
                        conn2.uid("STORE", _uid_bytes(uid), "+FLAGS", "\\Seen")
                except Exception:
                    pass
            _t_total = _t.monotonic() - _t0
            if _t_total > 2.0:
                logger.warning(
                    f"Slow email read uid={uid} folder={folder} "
                    f"select={_t_select*1000:.0f}ms fetch={_t_fetch*1000:.0f}ms "
                    f"size={len(raw)} total={_t_total*1000:.0f}ms"
                )

            # Look up cached summary, AI reply, and LLM-detected boundaries
            # by Message-ID
            cached_summary = None
            cached_ai_reply = None
            cached_boundaries = None
            try:
                import sqlite3 as _sql3
                _c = _sql3.connect(SCHEDULED_DB)
                owner_clause, owner_params = _email_cache_owner_clause(owner)
                _row = _c.execute(
                    f"SELECT summary FROM email_summaries WHERE message_id = ? AND {owner_clause}",
                    (message_id.strip(), *owner_params),
                ).fetchone()
                if _row:
                    cached_summary = _row[0]
                _row2 = _c.execute(
                    f"SELECT reply FROM email_ai_replies WHERE message_id = ? AND {owner_clause}",
                    (message_id.strip(), *owner_params),
                ).fetchone()
                if _row2:
                    cached_ai_reply = _apply_email_style_mechanics(_extract_reply(_row2[0] or ""))
                _row3 = _c.execute(
                    "SELECT sig_start, quote_start, turns_json FROM email_boundaries WHERE message_id = ?",
                    (message_id.strip(),),
                ).fetchone()
                cached_turns = None
                cached_sender_sig = None
                # Look up a per-sender cached signature (built by the
                # `learn_sender_signatures` action). Used by the renderer
                # to fold sigs consistently from the same address.
                try:
                    if sender_addr:
                        _rs = _c.execute(
                            f"SELECT signature_text FROM sender_signatures "
                            f"WHERE from_address = ? AND {owner_clause}",
                            (sender_addr.lower().strip(), *owner_params),
                        ).fetchone()
                        if _rs and _rs[0]:
                            cached_sender_sig = _rs[0]
                except Exception:
                    pass
                if _row3:
                    cached_boundaries = {"sig_start": _row3[0], "quote_start": _row3[1]}
                    if _row3[2]:
                        try:
                            from src.email_thread_parser import THREAD_PARSER_VERSION
                            _parsed = json.loads(_row3[2])
                            # Versioned envelope: {"v": N, "turns": [...]}.
                            # Anything else (bare list from older code, wrong
                            # version) is treated as a cache miss so the
                            # on-the-fly parser re-runs and the next write
                            # warms the cache with the current shape.
                            if (
                                isinstance(_parsed, dict)
                                and _parsed.get("v") == THREAD_PARSER_VERSION
                                and isinstance(_parsed.get("turns"), list)
                            ):
                                cached_turns = _parsed["turns"]
                        except Exception:
                            cached_turns = None
                _c.close()
            except Exception:
                pass

            # If no cached turns, parse on-the-fly so the client never has
            # to do the heavy lifting. Cheap on a 50KB body, free for short
            # ones. The background task warms the cache for next reads.
            if cached_turns is None:
                try:
                    from src.email_thread_parser import parse_thread
                    cached_turns = parse_thread(body_html, body)
                except Exception as _pe:
                    logger.debug(f"thread parse on read failed: {_pe}")
                    cached_turns = None

            return {
                "uid": uid,
                "folder": folder,
                "message_id": message_id.strip(),
                "subject": subject,
                "from_name": sender_name or sender_addr,
                "from_address": sender_addr,
                "to": to,
                "cc": cc,
                "date": parsed_date.isoformat() if parsed_date else "",
                "in_reply_to": in_reply_to.strip(),
                "references": references.strip(),
                "body": body,
                "body_html": body_html,
                "attachments": attachments,
                "cached_summary": cached_summary,
                "cached_ai_reply": cached_ai_reply,
                "boundaries": cached_boundaries,
                "thread_turns": cached_turns,
                "sender_signature": cached_sender_sig,
            }
        except Exception as e:
            logger.error(f"Failed to read email {uid}: {e}")
            return {"error": "Mail operation failed"}

    def _mark_email_seen_sync(uid, folder, account_id, owner):
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                conn.uid("STORE", _uid_bytes(uid), "+FLAGS", "\\Seen")
            _invalidate_list_cache(account_id, folder)
        except Exception as e:
            logger.debug(f"mark-seen after cached read failed uid={uid}: {e}")

    @router.get("/read/{uid}")
    async def read_email_by_uid(
        uid: str,
        folder: str = Query("INBOX"),
        account_id: str | None = Query(None),
        mark_seen: bool = Query(True),
        owner: str = Depends(require_owner),
    ):
        """Read email body. Cached for 30m, sync IMAP work runs in a thread."""
        ck = _read_cache_key(account_id, folder, uid, owner=owner)
        cached = _read_cache_get(ck)
        if cached is not None:
            if mark_seen:
                try:
                    _asyncio.create_task(_asyncio.to_thread(_mark_email_seen_sync, uid, folder, account_id, owner))
                except RuntimeError:
                    pass
            return cached
        result = await _asyncio.to_thread(_read_email_sync, uid, folder, account_id, owner, mark_seen)
        if result and not result.get("error"):
            _read_cache_put(ck, result)
        return result

    def _schedule_recent_email_warm(emails: list, folder: str, account_id: str | None, owner: str):
        if not emails or folder == "__scheduled__":
            return
        now = _time.time()
        selected = []
        for em in emails:
            uid = str((em or {}).get("uid") or "").strip()
            if not uid:
                continue
            try:
                epoch = float((em or {}).get("date_epoch") or 0)
            except Exception:
                epoch = 0
            if epoch and now - epoch > _WARM_RECENT_SECONDS:
                continue
            try:
                size = int((em or {}).get("size") or 0)
            except Exception:
                size = 0
            if size > _WARM_MAX_BYTES:
                continue
            ck = _read_cache_key(account_id, folder, uid, owner=owner)
            if _read_cache_get(ck) is not None or ck in _WARMING_READS:
                continue
            _WARMING_READS.add(ck)
            selected.append((uid, ck))
            if len(selected) >= _WARM_READ_LIMIT:
                break
        if not selected:
            return

        async def _warm():
            for uid, ck in selected:
                if _read_cache_get(ck) is not None:
                    _WARMING_READS.discard(ck)
                    continue
                try:
                    result = await _asyncio.to_thread(_read_email_sync, uid, folder, account_id, owner, False)
                    if result and not result.get("error"):
                        _read_cache_put(ck, result)
                except Exception as e:
                    logger.debug(f"email read warm skipped uid={uid}: {e}")
                finally:
                    _WARMING_READS.discard(ck)
                    await _asyncio.sleep(0.05)

        try:
            _asyncio.create_task(_warm())
        except RuntimeError:
            pass

    @router.get("/attachments/{uid}")
    async def list_attachments(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """List attachments for an email."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"attachments": [], "error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)
            attachments = _list_attachments_from_msg(msg)
            return {"attachments": attachments, "uid": uid}
        except Exception as e:
            logger.error(f"Failed to list attachments for {uid}: {e}")
            return {"attachments": [], "error": "Mail operation failed"}

    @router.get("/attachment/{uid}/{index}")
    async def download_attachment(uid: str, index: int, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Download a specific attachment by email UID and attachment index. Saves to local disk and returns the file."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)

            # Extract to a per-email folder
            target_dir = attachment_extract_dir(folder, uid)
            filepath = _extract_attachment_to_disk(msg, index, target_dir)
            if not filepath:
                return {"error": f"Attachment index {index} not found"}

            return FileResponse(
                path=str(filepath),
                filename=filepath.name,
                media_type="application/octet-stream",
            )
        except Exception as e:
            logger.error(f"Failed to download attachment {uid}/{index}: {e}")
            return {"error": "Mail operation failed"}

    @router.post("/attachment-as-doc/{uid}/{index}")
    async def attachment_as_doc(uid: str, index: int, request: Request, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Extract an email attachment and open it in the document editor.

        Supported extensions:
          - .pdf   → rendered as PDF Document (existing flow)
          - .docx  → text extracted to markdown Document
          - .txt / .md → loaded directly as a markdown Document

        Returns {doc_id} so the frontend can open it as a tab in the doc panel.
        Other types are rejected — caller should fall back to download.
        """
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)

            target_dir = attachment_extract_dir(folder, uid)
            filepath = _extract_attachment_to_disk(msg, index, target_dir)
            if not filepath:
                return {"error": f"Attachment index {index} not found"}

            from pathlib import Path as _Path
            base = _Path(filepath).name
            if base.startswith("."):
                return {"error": "Invalid filename", "filename": base}
            ext = _Path(base).suffix.lower()

            import os as _os
            title = _os.path.splitext(filepath.name)[0]

            # Capture the source email's identity so the doc can later be used
            # to thread a signed-reply back to the original sender.
            src_message_id = (msg.get("Message-ID") or "").strip()
            def _tag_doc_with_source(doc_id_to_tag: str):
                if not doc_id_to_tag:
                    return
                try:
                    from src.database import SessionLocal as _SL, Document as _Doc
                    _db = _SL()
                    try:
                        d = _db.query(_Doc).filter(_Doc.id == doc_id_to_tag).first()
                        if d:
                            d.source_email_uid = str(uid)
                            d.source_email_folder = folder
                            d.source_email_account_id = account_id or ""
                            d.source_email_message_id = src_message_id
                            _db.commit()
                    finally:
                        _db.close()
                except Exception as _e:
                    logger.warning(f"tag doc source-email failed: {_e}")

            # Extracted docs MUST belong to a session the caller owns — a
            # session-less ("orphan") doc is rejected by get_document's owner
            # check (404), so the frontend's loadDocument() throws and nothing
            # opens (the "open in document didn't open" bug). Attach it to the
            # user's most-recent session so it's fetchable + ownable.
            from src.auth_helpers import get_current_user as _gcu
            _doc_user = _gcu(request)
            def _resolve_doc_session():
                try:
                    from src.database import SessionLocal as _SL, Session as _Sess
                    _db = _SL()
                    try:
                        _q2 = _db.query(_Sess)
                        if _doc_user:
                            _q2 = _q2.filter(_Sess.owner == _doc_user)
                        s = _q2.order_by(_Sess.updated_at.desc()).first()
                        return s.id if s else None
                    finally:
                        _db.close()
                except Exception as _e:
                    logger.warning(f"resolve doc session failed: {_e}")
                    return None
            doc_session_id = _resolve_doc_session()

            # ── PDF path (existing) ────────────────────────────────────
            if ext == ".pdf":
                import shutil as _shutil
                from src.constants import UPLOAD_DIR
                from src.pdf_forms import has_form_fields, extract_fields
                from src.pdf_form_doc import (
                    save_field_sidecar,
                    create_form_markdown_document,
                    create_plain_pdf_document,
                )

                upload_id = f"{uuid.uuid4().hex}.pdf"
                today = datetime.utcnow().strftime("%Y/%m/%d")
                dated_dir = _os.path.join(UPLOAD_DIR, today)
                _os.makedirs(dated_dir, exist_ok=True)
                dest_path = _os.path.join(dated_dir, upload_id)
                _shutil.copyfile(str(filepath), dest_path)

                is_form = False
                try:
                    is_form = has_form_fields(dest_path)
                except Exception as e:
                    logger.warning(f"has_form_fields failed for attachment PDF: {e}")

                if is_form:
                    fields = extract_fields(dest_path)
                    save_field_sidecar(dest_path, fields)
                    doc_id = create_form_markdown_document(
                        session_id=doc_session_id,
                        fields=fields,
                        upload_id=upload_id,
                        title=title,
                        intro_text=None,
                    )
                else:
                    doc_id = create_plain_pdf_document(
                        session_id=doc_session_id,
                        upload_id=upload_id,
                        title=title,
                    )

                if not doc_id:
                    return {"error": "Failed to create document"}
                _tag_doc_with_source(doc_id)
                return {"doc_id": doc_id, "filename": filepath.name}

            # ── DOCX path: extract text → markdown document ───────────
            if ext == ".docx":
                try:
                    from docx import Document as _Docx
                except ImportError:
                    return {"error": "python-docx not installed", "filename": base}
                try:
                    d = _Docx(str(filepath))
                except Exception as e:
                    return {"error": f"Failed to read docx: {e}", "filename": base}
                # Convert paragraphs to markdown — preserve heading styles as #/##/###,
                # bullet lists as `- `, numbered lists as `1.`, and keep tables as
                # simple pipe-delimited rows.
                lines: list[str] = []
                for p in d.paragraphs:
                    text = p.text or ""
                    style = (p.style.name if p.style else "") or ""
                    if not text.strip():
                        lines.append("")
                        continue
                    if style.startswith("Heading 1"): lines.append(f"# {text}")
                    elif style.startswith("Heading 2"): lines.append(f"## {text}")
                    elif style.startswith("Heading 3"): lines.append(f"### {text}")
                    elif style.startswith("Heading "): lines.append(f"#### {text}")
                    elif style.startswith("List Bullet"): lines.append(f"- {text}")
                    elif style.startswith("List Number"): lines.append(f"1. {text}")
                    else: lines.append(text)
                for tbl in d.tables:
                    lines.append("")
                    for ri, row in enumerate(tbl.rows):
                        cells = [(c.text or "").replace("|", "\\|").replace("\n", " ").strip() for c in row.cells]
                        lines.append("| " + " | ".join(cells) + " |")
                        if ri == 0:
                            lines.append("|" + "|".join(["---"] * len(cells)) + "|")
                    lines.append("")
                content = "\n".join(lines).strip() or f"_(empty {base})_"

                from src.database import SessionLocal as _SL, Document as _Doc, DocumentVersion as _DV
                doc_id = str(uuid.uuid4())
                ver_id = str(uuid.uuid4())
                _db = _SL()
                try:
                    _db.query(_Doc).filter(_Doc.is_active == True).update({"is_active": False})
                    _db.add(_Doc(
                        id=doc_id, session_id=doc_session_id, title=title,
                        language="markdown", current_content=content,
                        version_count=1, is_active=True,
                    ))
                    _db.add(_DV(
                        id=ver_id, document_id=doc_id, version_number=1,
                        content=content, summary="Imported from DOCX", source="upload",
                    ))
                    _db.commit()
                finally:
                    _db.close()
                _tag_doc_with_source(doc_id)
                return {"doc_id": doc_id, "filename": filepath.name}

            # ── Plain text / markdown ────────────────────────────────
            if ext in (".txt", ".md", ".markdown"):
                try:
                    content = filepath.read_text(encoding="utf-8", errors="replace")
                except Exception as e:
                    return {"error": f"Failed to read text file: {e}", "filename": base}
                from src.database import SessionLocal as _SL, Document as _Doc, DocumentVersion as _DV
                doc_id = str(uuid.uuid4())
                ver_id = str(uuid.uuid4())
                _db = _SL()
                try:
                    _db.query(_Doc).filter(_Doc.is_active == True).update({"is_active": False})
                    _db.add(_Doc(
                        id=doc_id, session_id=doc_session_id, title=title,
                        language="markdown", current_content=content,
                        version_count=1, is_active=True,
                    ))
                    _db.add(_DV(
                        id=ver_id, document_id=doc_id, version_number=1,
                        content=content, summary="Imported from email attachment", source="upload",
                    ))
                    _db.commit()
                finally:
                    _db.close()
                _tag_doc_with_source(doc_id)
                return {"doc_id": doc_id, "filename": filepath.name}

            return {"error": f"Unsupported attachment type: {ext}", "filename": base}
        except Exception as e:
            logger.error(f"attachment-as-doc {uid}/{index} failed: {e}")
            return {"error": "Mail operation failed"}

    @router.post("/attachment-path/{uid}/{index}")
    async def get_attachment_path(uid: str, index: int, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Extract attachment to local disk and return the path (for AI to read via read_file)."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)

            target_dir = attachment_extract_dir(folder, uid)
            filepath = _extract_attachment_to_disk(msg, index, target_dir)
            if not filepath:
                return {"error": f"Attachment index {index} not found"}

            return {"path": str(filepath), "filename": filepath.name, "size": filepath.stat().st_size}
        except Exception as e:
            logger.error(f"Failed to get attachment path {uid}/{index}: {e}")
            return {"error": "Mail operation failed"}

    @router.post("/mark-unread/{uid}")
    async def mark_unread(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Mark an email as unread (clear \\Seen flag)."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _store_email_flag(conn, uid, "\\Seen", add=False):
                    return {"success": False, "error": "Email not found"}
            _invalidate_list_cache(account_id, folder)
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to mark unread {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/flag/{uid}")
    async def flag_email(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None),
                         on: bool = Query(True), owner: str = Depends(require_owner)):
        """Toggle the \\Flagged flag (a.k.a. favorite / star) on an email.
        Pass `on=true` to favorite, `on=false` to unfavorite."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _store_email_flag(conn, uid, "\\Flagged", add=bool(on)):
                    return {"success": False, "error": "Email not found"}
            _invalidate_list_cache(account_id, folder)
            return {"success": True, "flagged": bool(on)}
        except Exception as e:
            logger.error(f"Failed to flag {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/mark-read/{uid}")
    async def mark_read(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Mark an email as read (set \\Seen flag)."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _store_email_flag(conn, uid, "\\Seen", add=True):
                    return {"success": False, "error": "Email not found"}
            _invalidate_list_cache(account_id, folder)
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to mark read {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/archive/{uid}")
    # Sync def: blocking IMAP I/O with no awaits — see search_emails above. Runs in a
    # threadpool instead of blocking the event loop.
    def archive_email(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Move email to Archive folder."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _move_email_message(conn, uid, "Archive", role="archive"):
                    return {"success": False, "error": "Email not found"}
            _invalidate_list_cache(account_id)
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to archive email {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.delete("/delete/{uid}")
    async def delete_email(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Move email to Trash."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _move_email_message(conn, uid, "Trash", role="trash"):
                    return {"success": False, "error": "Email not found"}
            _invalidate_list_cache(account_id)
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to delete email {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.delete("/delete-permanent/{uid}")
    async def delete_email_permanent(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Permanently delete an email (no Trash)."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _store_email_flag(conn, uid, "\\Deleted", add=True):
                    return {"success": False, "error": "Email not found"}
                conn.expunge()
            _invalidate_list_cache(account_id, folder)
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to permanently delete email {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.delete("/odysseus/reminders")
    async def delete_odysseus_reminder_emails(
        account_id: str | None = Query(None),
        permanent: bool = Query(False),
        owner: str = Depends(require_owner),
    ):
        """Delete email messages stamped as Odysseus reminders."""
        if account_id:
            _assert_owns_account(account_id, owner)
        deleted = 0
        folders_checked = []
        try:
            cfg = _get_email_config(account_id, owner=owner)
            own_addrs = [
                (cfg.get("from_address") or "").strip(),
                (cfg.get("smtp_user") or "").strip(),
                (cfg.get("imap_user") or "").strip(),
            ]
            own_addrs = [a for i, a in enumerate(own_addrs) if a and a not in own_addrs[:i]]

            def _search_quote(value: str) -> str:
                return '"' + (value or "").replace("\\", "\\\\").replace('"', '\\"') + '"'

            def _search_uids(conn, criteria: str):
                st, data = conn.uid("SEARCH", None, criteria)
                return set(data[0].split()) if st == "OK" and data and data[0] else set()

            with _imap(account_id, owner=owner) as conn:
                sent_folder = _detect_sent_folder(conn)
                candidates = ["INBOX", sent_folder, "All Mail", "[Gmail]/All Mail"]
                seen = set()
                for folder_name in candidates:
                    if not folder_name or folder_name in seen:
                        continue
                    seen.add(folder_name)
                    try:
                        st, _ = conn.select(_q(folder_name))
                        if st != "OK":
                            continue
                        folders_checked.append(folder_name)
                        uids = set()
                        # Match the Reminders filter: new messages have the
                        # explicit kind header, and subject fallback catches
                        # clients/providers that stripped custom headers.
                        uids.update(_search_uids(conn, f'(HEADER X-Odysseus-Kind {_search_quote("reminder")})'))
                        uids.update(_search_uids(conn, f'(SUBJECT {_search_quote("Reminder (Odysseus):")})'))
                        for addr in own_addrs:
                            addr_q = _search_quote(addr)
                            uids.update(_search_uids(conn, f'(FROM {addr_q} SUBJECT {_search_quote("Reminder (Odysseus):")})'))
                            # Legacy reminders created before the Odysseus
                            # prefix still came from this mailbox as
                            # "Reminder: ..."; include them in Clear without
                            # sweeping unrelated external reminder emails.
                            uids.update(_search_uids(conn, f'(FROM {addr_q} SUBJECT {_search_quote("Reminder:")})'))
                        if not uids:
                            continue
                        for uid in sorted(uids, key=lambda b: int(b)):
                            if permanent:
                                conn.uid("STORE", uid, "+FLAGS", "\\Deleted")
                            else:
                                copy_st, _ = conn.uid("COPY", uid, _q("Trash"))
                                if copy_st == "OK":
                                    conn.uid("STORE", uid, "+FLAGS", "\\Deleted")
                                else:
                                    conn.uid("STORE", uid, "+FLAGS", "\\Deleted")
                            deleted += 1
                        conn.expunge()
                    except Exception as e:
                        logger.warning(f"Skipped reminder cleanup in {folder_name!r}: {e}")
            _invalidate_list_cache(account_id)
            return {"success": True, "deleted": deleted, "folders_checked": folders_checked}
        except Exception as e:
            logger.error(f"delete_odysseus_reminder_emails failed: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/move/{uid}")
    async def move_email(uid: str, folder: str = Query("INBOX"), dest: str = Query(...), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Move an email to another folder."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _move_email_message(conn, uid, dest):
                    return {"success": False, "error": f"Failed to move to {dest}"}
            _invalidate_list_cache(account_id)
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to move email {uid} to {dest}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.get("/folders")
    async def list_folders(account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """List IMAP folders."""
        try:
            with _imap(account_id, owner=owner) as conn:
                status, folders = conn.list()
            result = []
            for f in folders:
                decoded = f.decode() if isinstance(f, bytes) else f
                match = re.search(r'"([^"]*)"$|(\S+)$', decoded)
                if match:
                    name = match.group(1) or match.group(2)
                    result.append(name)
            return {"folders": result}
        except Exception as e:
            logger.error(f"list_folders failed: {e}")
            return {"folders": [], "error": "Mail operation failed"}

    @router.post("/mark-answered/{uid}")
    async def mark_answered(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Mark an email as answered (set \\Answered flag)."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _store_email_flag(conn, uid, "\\Answered", add=True):
                    return {"success": False, "error": "Email not found"}
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to mark answered {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/clear-answered/{uid}")
    async def clear_answered(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Clear the \\Answered flag from an email."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder))
                if not _store_email_flag(conn, uid, "\\Answered", add=False):
                    return {"success": False, "error": "Email not found"}
            return {"success": True}
        except Exception as e:
            logger.error(f"Failed to clear answered {uid}: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/compose-upload")
    async def compose_upload(file: UploadFile = File(...), owner: str = Depends(require_owner)):
        """Upload a file for attaching to a compose email. Returns a token."""
        try:
            # Sanitize filename and generate a unique token
            safe_name = re.sub(r"[^\w\s\-.]", "_", file.filename or "file").strip()
            token = f"{uuid.uuid4().hex}_{safe_name}"
            filepath = COMPOSE_UPLOADS_DIR / token
            content = await read_upload_limited(file, EMAIL_COMPOSE_UPLOAD_MAX_BYTES, "Attachment")
            with open(filepath, "wb") as f:
                f.write(content)
            return {
                "success": True,
                "token": token,
                "filename": safe_name,
                "size": len(content),
            }
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to upload attachment: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.delete("/compose-upload/{token}")
    async def delete_compose_upload(token: str, owner: str = Depends(require_owner)):
        """Delete a staged compose upload."""
        try:
            # Prevent path traversal
            safe_token = Path(token).name
            filepath = COMPOSE_UPLOADS_DIR / safe_token
            if filepath.exists():
                filepath.unlink()
            return {"success": True}
        except Exception as e:
            logger.error(f"delete_compose_upload {token!r} failed: {e}")
            return {"success": False, "error": "Mail operation failed"}

    async def _send_email_sync(
        to, cc, bcc, subject, body, in_reply_to, references, attachments,
        account_id=None, owner="", odysseus_kind=None, odysseus_ref=None,
    ):
        """Shared send logic used by both /send and scheduled delivery.

        SECURITY: callers MUST pass `owner` (the authed user) so the config
        lookup is scoped — otherwise the fallback picks whichever account
        happens to be is_default globally and the message ships through
        someone else's SMTP creds + From-address.
        """
        cfg = _resolve_send_config(account_id, owner=owner)
        has_atts = bool(attachments)
        if has_atts:
            outer = MIMEMultipart("mixed")
            body_container = MIMEMultipart("alternative")
        else:
            outer = MIMEMultipart("alternative")
            body_container = outer

        outer["From"] = email.utils.formataddr((cfg.get("display_name") or "", cfg["from_address"]))
        outer["To"] = to
        if cc:
            outer["Cc"] = cc
        outer["Subject"] = subject or ""
        outer["Date"] = datetime.utcnow().strftime("%a, %d %b %Y %H:%M:%S +0000")
        _apply_odysseus_headers(outer, odysseus_kind or "scheduled", odysseus_ref)
        if in_reply_to:
            outer["In-Reply-To"] = in_reply_to
        if references:
            outer["References"] = references

        body_container.attach(MIMEText(body or "", "plain", "utf-8"))
        body_container.attach(MIMEText(_md_to_email_html(body or ""), "html", "utf-8"))

        if has_atts:
            outer.attach(body_container)
            _attach_compose_uploads(outer, attachments)

        recipients = _envelope_recipients(to, cc, bcc)

        _send_smtp_message(cfg, cfg["from_address"], recipients, outer.as_string())

        _cleanup_compose_uploads(attachments)

    @router.post("/schedule")
    async def schedule_email(req: dict, owner: str = Depends(require_owner)):
        """Schedule an email to be sent at a specific time. ISO8601 UTC."""
        import sqlite3
        import uuid as _uuid
        try:
            send_at = req.get("send_at")
            if not send_at:
                return {"success": False, "error": "send_at required (ISO8601 UTC)"}
            # Body-based account_id — dep can't see it, check here.
            _acct = req.get("account_id")
            if _acct:
                _assert_owns_account(_acct, owner)
            # Validate parseable + reject past times (the poller fires
            # anything in the past immediately on the next tick — a
            # 1970-dated schedule would deliver right now).
            from datetime import datetime as _dt, timezone as _tz
            try:
                parsed_at = _dt.fromisoformat(send_at.replace("Z", "+00:00"))
            except ValueError:
                return {"success": False, "error": "send_at must be ISO8601"}
            now_utc = _dt.now(_tz.utc) if parsed_at.tzinfo else _dt.utcnow()
            # Tiny 30s grace so a user clicking Send right at the chosen
            # minute doesn't trip the past-time guard.
            if parsed_at < now_utc:
                return {"success": False, "error": "send_at must be in the future"}
            # Normalize to naive UTC before storing: the poller selects due
            # rows with a lexicographic string compare against a naive
            # datetime.utcnow().isoformat(), so storing the raw client string
            # makes "+02:00" schedules fire hours late, negative offsets fire
            # hours early, and a "Z" suffix compares after the fractional
            # seconds of the poller timestamp.
            if parsed_at.tzinfo:
                parsed_at = parsed_at.astimezone(_tz.utc).replace(tzinfo=None)
            send_at = parsed_at.isoformat()

            sid = _uuid.uuid4().hex[:16]
            conn = sqlite3.connect(SCHEDULED_DB)
            conn.execute("""
                INSERT INTO scheduled_emails
                (id, to_addr, cc, bcc, subject, body, in_reply_to, references_hdr, attachments, send_at, created_at, status, account_id, odysseus_kind, owner)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?, ?, ?)
            """, (
                sid,
                req.get("to", ""),
                req.get("cc") or None,
                req.get("bcc") or None,
                req.get("subject") or "",
                req.get("body") or "",
                req.get("in_reply_to") or None,
                req.get("references") or None,
                json.dumps(req.get("attachments") or []),
                send_at,
                datetime.utcnow().isoformat(),
                req.get("account_id") or None,
                req.get("odysseus_kind") or "scheduled",
                owner or "",
            ))
            conn.commit()
            conn.close()
            logger.info(f"Scheduled email {sid} for {send_at}")
            return {"success": True, "id": sid, "send_at": send_at}
        except Exception as e:
            logger.error(f"Failed to schedule email: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.get("/scheduled")
    async def list_scheduled(owner: str = Depends(require_owner)):
        """List all scheduled (pending) emails."""
        import sqlite3
        try:
            conn = sqlite3.connect(SCHEDULED_DB)
            rows = conn.execute("""
                SELECT id, to_addr, cc, subject, send_at, created_at, status, error
                FROM scheduled_emails
                WHERE status IN ('pending', 'failed') AND owner = ?
                ORDER BY send_at ASC
            """, (owner or "",)).fetchall()
            conn.close()
            return {"scheduled": [
                {
                    "id": r[0], "to": r[1], "cc": r[2], "subject": r[3],
                    "send_at": r[4], "created_at": r[5], "status": r[6], "error": r[7],
                } for r in rows
            ]}
        except Exception as e:
            logger.error(f"list_scheduled failed: {e}")
            return {"scheduled": [], "error": "Mail operation failed"}

    @router.delete("/scheduled/{sid}")
    async def cancel_scheduled(sid: str, owner: str = Depends(require_owner)):
        """Cancel a scheduled email."""
        import sqlite3
        try:
            conn = sqlite3.connect(SCHEDULED_DB)
            conn.execute(
                "DELETE FROM scheduled_emails WHERE id = ? AND status = 'pending' AND owner = ?",
                (sid, owner or ""),
            )
            conn.commit()
            conn.close()
            return {"success": True}
        except Exception as e:
            logger.error(f"cancel_scheduled {sid!r} failed: {e}")
            return {"success": False, "error": "Mail operation failed"}

    # ── Agent send-confirm: list/approve/cancel ──────────────────────────
    # When `agent_email_confirm` is on, the MCP send_email tool drops the
    # composed email into scheduled_emails with status='agent_draft' (a
    # far-future send_at so the poller never picks it up). These endpoints
    # let the chat UI surface them for the user and either approve (flip
    # to status='pending' with send_at=now so the poller delivers it) or
    # cancel (status='cancelled').
    @router.get("/pending")
    async def list_pending_agent_drafts(owner: str = Depends(require_owner)):
        import sqlite3
        try:
            conn = sqlite3.connect(SCHEDULED_DB)
            conn.row_factory = sqlite3.Row
            rows = conn.execute(
                """SELECT id, to_addr, subject, body, created_at, account_id
                   FROM scheduled_emails
                   WHERE status = 'agent_draft' AND owner = ?
                   ORDER BY created_at DESC""",
                (owner or "",),
            ).fetchall()
            conn.close()
            return {"pending": [dict(r) for r in rows]}
        except Exception as e:
            logger.error(f"list_pending_agent_drafts failed: {e}")
            return {"pending": [], "error": "Mail operation failed"}

    @router.post("/pending/{sid}/approve")
    async def approve_agent_draft(sid: str, owner: str = Depends(require_owner)):
        """Approve a draft staged by the agent: flip status → pending and
        backdate send_at so the scheduled-send poller picks it up
        immediately."""
        import sqlite3
        try:
            conn = sqlite3.connect(SCHEDULED_DB)
            cur = conn.execute(
                """UPDATE scheduled_emails
                   SET status = 'pending', send_at = ?
                   WHERE id = ? AND status = 'agent_draft' AND owner = ?""",
                (datetime.utcnow().isoformat(), sid, owner or ""),
            )
            conn.commit()
            affected = cur.rowcount
            conn.close()
            if not affected:
                return {"success": False, "error": "Draft not found or already handled"}
            return {"success": True}
        except Exception as e:
            logger.error(f"approve_agent_draft {sid!r} failed: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.delete("/pending/{sid}")
    async def cancel_agent_draft(sid: str, owner: str = Depends(require_owner)):
        """Discard a draft the agent staged for approval."""
        import sqlite3
        try:
            conn = sqlite3.connect(SCHEDULED_DB)
            cur = conn.execute(
                """UPDATE scheduled_emails SET status = 'cancelled'
                   WHERE id = ? AND status = 'agent_draft' AND owner = ?""",
                (sid, owner or ""),
            )
            conn.commit()
            affected = cur.rowcount
            conn.close()
            if not affected:
                return {"success": False, "error": "Draft not found or already handled"}
            return {"success": True}
        except Exception as e:
            logger.error(f"cancel_agent_draft {sid!r} failed: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.get("/resolve-contact")
    async def resolve_contact(name: str = Query(..., description="Name to search for"), owner: str = Depends(require_owner)):
        """Search Sent folder for a contact by name. Returns matching email addresses."""
        try:
            with _imap(owner=owner) as conn:
                matches = {}
                for folder in ["Sent", "INBOX", "Drafts"]:
                    try:
                        st, _ = conn.select(_q(folder), readonly=True)
                        if st != "OK":
                            continue
                        st, data = conn.search(None, "ALL")
                        if st != "OK" or not data[0]:
                            continue
                        uids = data[0].split()[-200:]
                        for uid in reversed(uids):
                            try:
                                st2, msg_data = conn.fetch(uid, "(BODY.PEEK[HEADER.FIELDS (FROM TO CC)])")
                                if st2 != "OK":
                                    continue
                                raw = msg_data[0][1] if msg_data[0] and len(msg_data[0]) > 1 else b""
                                hdr = email_mod.message_from_bytes(raw)
                                for field in ["From", "To", "Cc"]:
                                    val = _decode_header(hdr.get(field, ""))
                                    if not val:
                                        continue
                                    for part in val.split(","):
                                        part = part.strip()
                                        if name.lower() in part.lower():
                                            addr_match = re.search(r'<([^>]+)>', part)
                                            addr = addr_match.group(1) if addr_match else part
                                            addr = addr.strip().lower()
                                            if addr and "@" in addr:
                                                display = part.split("<")[0].strip().strip('"') or addr
                                                if addr not in matches:
                                                    matches[addr] = display
                            except Exception:
                                continue
                    except Exception:
                        continue
                    if len(matches) >= 10:
                        break
                results = [{"email": addr, "name": display} for addr, display in matches.items()]
                return {"contacts": results[:10], "query": name}
        except Exception as e:
            logger.error(f"resolve_contact {name!r} failed: {e}")
            return {"contacts": [], "error": "Mail operation failed"}

    @router.post("/send")
    async def send_email(req: SendEmailRequest, background_tasks: BackgroundTasks, owner: str = Depends(require_owner)):
        """Queue an email for SMTP delivery. Returns immediately; send runs in background.

        Uses req.account_id to pick the sending account (falls back to default)."""
        # Body-based account_id — dep can't see it, check here.
        if req.account_id:
            _assert_owns_account(req.account_id, owner)

        try:
            cfg = _resolve_send_config(req.account_id, owner=owner)
        except Exception as e:
            logger.warning(f"No SMTP-capable account resolved: {e}")
            return {"success": False, "error": str(e) or "No SMTP-capable email account configured"}

        # Use 'mixed' if we have attachments, 'alternative' otherwise
        has_attachments = bool(req.attachments)
        logger.info(f"Sending email to {req.to}: subject={req.subject!r}, attachments={req.attachments}")
        if has_attachments:
            outer = MIMEMultipart("mixed")
            body_container = MIMEMultipart("alternative")
        else:
            outer = MIMEMultipart("alternative")
            body_container = outer

        outer["From"] = email.utils.formataddr((cfg.get("display_name") or "", cfg["from_address"]))
        outer["To"] = req.to
        if req.cc:
            outer["Cc"] = req.cc
        outer["Subject"] = req.subject
        outer["Date"] = datetime.utcnow().strftime("%a, %d %b %Y %H:%M:%S +0000")
        outer["Message-ID"] = email.utils.make_msgid(domain="odysseus.local")

        if req.in_reply_to:
            outer["In-Reply-To"] = req.in_reply_to
        if req.references:
            outer["References"] = req.references
        if req.odysseus_kind:
            _apply_odysseus_headers(outer, req.odysseus_kind)

        # Plain + HTML body. Escape user content so a `<script>` or
        # `<img onerror=...>` paste in compose doesn't end up as live HTML
        # in the recipient's MUA.
        body_container.attach(MIMEText(req.body, "plain", "utf-8"))
        # HTML part: prefer the WYSIWYG composer's HTML (sanitized via allowlist);
        # otherwise render the markdown body. Both routes escape untrusted text,
        # so neither can introduce live script/handlers.
        _html_part = (_sanitize_email_html(req.body_html) if req.body_html else None) \
            or _md_to_email_html(req.body)
        body_container.attach(MIMEText(_html_part, "html", "utf-8"))

        if has_attachments:
            outer.attach(body_container)
            _attach_compose_uploads(outer, req.attachments)

        # Build recipient list (parse the address grammar so display names with
        # commas don't get split into broken envelope addresses)
        recipients = _envelope_recipients(req.to, req.cc, req.bcc)

        # Serialize what the background task needs so the request object can be GC'd
        outer_bytes = outer.as_bytes()
        outer_str = outer.as_string()
        _from = cfg["from_address"]
        _smtp_host = cfg["smtp_host"]
        _smtp_port = cfg["smtp_port"]
        _smtp_security = cfg.get("smtp_security")
        _smtp_user = cfg["smtp_user"]
        _smtp_pw = cfg["smtp_password"]
        _recipients = list(recipients)
        _to_label = req.to
        _subject = req.subject
        _atts = list(req.attachments or [])
        _message_id = outer["Message-ID"]

        _account_id = cfg.get("account_id") or req.account_id  # capture for the IMAP append in the closure
        _in_reply_to = (req.in_reply_to or "").strip()
        _oauth_provider = cfg.get("oauth_provider") or ""
        _oauth_access_token = cfg.get("oauth_access_token") or ""
        _oauth_refresh_token = cfg.get("oauth_refresh_token") or ""
        _oauth_token_expiry = cfg.get("oauth_token_expiry") or ""

        def _deliver():
            try:
                _send_smtp_message(
                    {
                        "smtp_host": _smtp_host,
                        "smtp_port": _smtp_port,
                        "smtp_security": _smtp_security,
                        "smtp_user": _smtp_user,
                        "smtp_password": _smtp_pw,
                        "account_id": _account_id,
                        "oauth_provider": _oauth_provider,
                        "oauth_access_token": _oauth_access_token,
                        "oauth_refresh_token": _oauth_refresh_token,
                        "oauth_token_expiry": _oauth_token_expiry,
                    },
                    _from,
                    _recipients,
                    outer_str,
                )
                logger.info(f"Email sent to {_to_label}: {_subject}")
                delivery_result = {
                    "success": True,
                    "account_id": cfg.get("account_id") or _account_id,
                    "sent_folder": None,
                    "sent_uid": None,
                    "message_id": _message_id,
                }
                try:
                    with _imap(_account_id, owner=owner) as imap:
                        sent_folder = _detect_sent_folder(imap)
                        sent_uid = None
                        append_st, append_data = imap.append(sent_folder, "\\Seen", None, outer_bytes)
                        if append_st == "OK" and append_data:
                            m = re.search(rb"APPENDUID\s+\d+\s+(\d+)", append_data[0] or b"")
                            if m:
                                sent_uid = m.group(1).decode("ascii", errors="ignore")
                        if not sent_uid:
                            try:
                                st_sel, _ = imap.select(_q(sent_folder), readonly=True)
                                if st_sel == "OK":
                                    mid = (_message_id or "").strip().lstrip("<").rstrip(">").replace('"', '\\"')
                                    st_uid, uid_data = imap.uid("SEARCH", None, f'HEADER Message-ID "{mid}"')
                                    if st_uid == "OK" and uid_data and uid_data[0]:
                                        sent_uid = uid_data[0].split()[-1].decode("ascii", errors="ignore")
                            except Exception:
                                pass
                        # Auto-mark the source email as Answered/done so it
                        # disappears from "undone" filters.
                        if _in_reply_to:
                            try:
                                # Strip any angle brackets and quote for IMAP
                                mid = _in_reply_to.strip().lstrip("<").rstrip(">").replace('"', '\\"')
                                # Search common folders for the source message.
                                folder_candidates = (
                                    "INBOX",
                                    sent_folder,
                                    "Sent",
                                    "[Gmail]/Sent Mail",
                                    "Archive",
                                    "All Mail",
                                    "[Gmail]/All Mail",
                                )
                                for folder_name in dict.fromkeys(folder_candidates):
                                    try:
                                        st, _sel = imap.select(_q(folder_name), readonly=False)
                                        if st != "OK":
                                            continue
                                        st2, sd = imap.search(None, f'HEADER Message-ID "{mid}"')
                                        if st2 == "OK" and sd and sd[0]:
                                            for u in sd[0].split():
                                                imap.store(u, "+FLAGS", "\\Answered")
                                            logger.info(f"Marked source {mid[:60]!r} as \\Answered in {folder_name}")
                                            break
                                    except Exception:
                                        continue
                            except Exception as e:
                                logger.warning(f"Failed to auto-mark source as answered: {e}")
                        delivery_result = {
                            "success": True,
                            "account_id": cfg.get("account_id") or _account_id,
                            "sent_folder": sent_folder,
                            "sent_uid": sent_uid,
                            "message_id": _message_id,
                        }
                except Exception as e:
                    logger.warning(f"Failed to append to Sent: {e}")
                _cleanup_compose_uploads(_atts)
                return delivery_result
            except Exception as e:
                logger.error(f"Failed to send email to {_to_label}: {e}")
                return {"success": False, "error": str(e) or "Failed to send email"}

        if req.wait_for_delivery:
            result = await asyncio.to_thread(_deliver)
            if result.get("success"):
                return {"success": True, "queued": False, "message": f"Email sent to {req.to}", **result}
            return result

        background_tasks.add_task(_deliver)
        return {
            "success": True,
            "queued": True,
            "account_id": cfg.get("account_id") or req.account_id,
            "message": f"Email queued for {req.to}",
        }

    @router.post("/draft")
    async def save_draft(req: SendEmailRequest, owner: str = Depends(require_owner)):
        """Save email as draft in IMAP Drafts folder.

        IMAP append is sync; offload via asyncio.to_thread so the event loop
        stays responsive on slow remote IMAP servers.
        """
        if req.account_id:
            _assert_owns_account(req.account_id, owner)
        cfg = _get_email_config(req.account_id, owner=owner)

        # Multipart plain+HTML when the WYSIWYG composer supplied HTML, so a
        # reopened draft keeps its formatting; plain MIMEText otherwise.
        _draft_html = _sanitize_email_html(req.body_html) if req.body_html else None
        if _draft_html:
            msg = MIMEMultipart("alternative")
            msg.attach(MIMEText(req.body, "plain", "utf-8"))
            msg.attach(MIMEText(_draft_html, "html", "utf-8"))
        else:
            msg = MIMEText(req.body, "plain", "utf-8")
        msg["From"] = email.utils.formataddr((cfg.get("display_name") or "", cfg["from_address"]))
        msg["To"] = req.to
        if req.cc:
            msg["Cc"] = req.cc
        if req.bcc:
            msg["Bcc"] = req.bcc
        msg["Subject"] = req.subject
        msg["Date"] = datetime.utcnow().strftime("%a, %d %b %Y %H:%M:%S +0000")

        if req.in_reply_to:
            msg["In-Reply-To"] = req.in_reply_to
        if req.references:
            msg["References"] = req.references

        _draft_acct = req.account_id

        def _do_append():
            try:
                with _imap(_draft_acct, owner=owner) as imap:
                    drafts_folder = _detect_drafts_folder(imap)
                    imap.append(drafts_folder, "\\Draft", None, msg.as_bytes())
                return None
            except Exception as e:
                return str(e)

        err = await asyncio.to_thread(_do_append)
        if err:
            logger.error(f"Failed to save draft: {err}")
            return {"success": False, "error": err}
        logger.info(f"Draft saved: {req.subject}")
        return {"success": True, "message": "Draft saved"}

    @router.post("/extract-style")
    async def extract_writing_style(req: ExtractStyleRequest, owner: str = Depends(require_owner)):
        """Extract writing style from sent emails using LLM.

        IMAP fetch is offloaded to a worker thread; the LLM call uses the
        async client. Otherwise this handler froze the event loop for ~5s
        on the IMAP step alone with a remote server.
        """

        def _gather_samples() -> tuple[list[str], str | None]:
            try:
                with _imap(owner=owner) as imap:
                    imap.select(_q(_detect_sent_folder(imap)), readonly=True)
                    status, data = imap.search(None, "ALL")
                    if status != "OK" or not data[0]:
                        return [], "No sent emails found"
                    uid_list = data[0].split()[-req.sample_count:]

                    out = []
                    for uid in uid_list:
                        try:
                            status, msg_data = imap.fetch(uid, "(RFC822)")
                            if status != "OK":
                                continue
                            raw = msg_data[0][1]
                            msg = email_mod.message_from_bytes(raw)
                            body = ""
                            if msg.is_multipart():
                                for part in msg.walk():
                                    if part.get_content_type() == "text/plain":
                                        payload = part.get_payload(decode=True)
                                        if payload:
                                            charset = part.get_content_charset() or "utf-8"
                                            body = payload.decode(charset, errors="replace")
                                            break
                            else:
                                payload = msg.get_payload(decode=True)
                                if payload:
                                    charset = msg.get_content_charset() or "utf-8"
                                    body = payload.decode(charset, errors="replace")
                            if body.strip() and len(body) > 20:
                                out.append(body[:1000])
                        except Exception:
                            continue
                    return out, None
            except Exception as e:
                return [], str(e)

        try:
            samples, err = await asyncio.to_thread(_gather_samples)
            if err and not samples:
                return {"success": False, "error": err}

            if len(samples) < 3:
                return {"success": False, "error": f"Only found {len(samples)} usable sent emails, need at least 3"}

            # Call LLM to analyze writing style. Prefer the utility model;
            # fall back to the default chat model when utility isn't set
            # (matches how the background email tasks behave).
            from src.endpoint_resolver import resolve_endpoint

            url, model, headers = resolve_endpoint("utility", owner=owner)
            if not url or not model:
                url, model, headers = resolve_endpoint("default", owner=owner)
            if not url or not model:
                return {"success": False, "error": "No LLM endpoint configured — set a Utility or Default Chat model in Settings → AI Defaults."}

            sample_text = "\n\n---EMAIL---\n\n".join(samples[:15])
            messages = [
                {
                    "role": "system",
                    "content": (
                        "You are analyzing a user's email writing style. Based on the sample emails below, "
                        "describe their writing style in 3-5 concise sentences. Cover: tone (formal/informal), "
                        "typical greeting and sign-off patterns, sentence structure (short/long), "
                        "any distinctive phrases or habits, and overall communication approach. "
                        "Write this as instructions for an AI to mimic this style. "
                        "Start with 'Write emails in this style:'"
                    ),
                },
                {
                    "role": "user",
                    "content": f"Here are {len(samples)} recently sent emails:\n\n{sample_text}",
                },
            ]

            style = await llm_call_async(url, model, messages, headers=headers, max_tokens=2048)
            style = _strip_think(style or "")
            if not style:
                return {"success": False, "error": "LLM failed to generate style description"}

            # Save to settings
            settings = _load_settings()
            settings["email_writing_style"] = style
            _save_settings(settings)

            logger.info("Writing style extracted and saved")
            return {"success": True, "style": style}

        except Exception as e:
            logger.error(f"Failed to extract writing style: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/summarize")
    async def summarize_email(data: dict, owner: str = Depends(require_owner)):
        """Generate a quick AI summary of an email body."""
        try:
            from src.endpoint_resolver import resolve_endpoint
            from src.llm_core import _uses_max_completion_tokens, _restricts_temperature
            import requests as _req

            body = data.get("body", "")
            subject = data.get("subject", "")
            sender = data.get("from", "")
            uid = data.get("uid", "")
            folder = data.get("folder", "INBOX") or "INBOX"
            account_id = data.get("account_id")
            if account_id:
                _assert_owns_account(account_id, owner)
            if not body:
                return {"success": False, "error": "No body provided"}

            # If we know which UID this is, fetch the raw message and pull
            # attachment text so the summary can reference invoice totals,
            # contract clauses, etc. — not just the body.
            att_text = ""
            if uid:
                try:
                    def _fetch_atts():
                        with _imap(account_id, owner=owner) as conn:
                            conn.select(_q(folder), readonly=True)
                            status, msg_data = _imap_uid_fetch(conn, str(uid), "(BODY.PEEK[])")
                            if status != "OK" or not msg_data or not msg_data[0]:
                                return ""
                            raw = msg_data[0][1]
                            msg_obj = email_mod.message_from_bytes(raw)
                            return _extract_attachment_text(msg_obj, max_chars=6000)
                    att_text = await asyncio.to_thread(_fetch_atts)
                except Exception as _ae:
                    logger.debug(f"on-demand summarize attachment fetch failed for uid={uid}: {_ae}")

            body_for_llm = body
            if att_text:
                body_for_llm = body + "\n\n--- ATTACHMENTS ---\n\n" + att_text

            url, model, headers = resolve_endpoint("utility", owner=owner)
            if not url:
                url, model, headers = resolve_endpoint("default", owner=owner)
            if not url or not model:
                return {"success": False, "error": "No LLM endpoint configured"}

            req_headers = {"Content-Type": "application/json"}
            if headers:
                req_headers.update(headers)
            tok_key = "max_completion_tokens" if _uses_max_completion_tokens(model) else "max_tokens"
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": "You are an email summarizer. Format: 1-3 short bullet points (use '- '). Cover: main point, action items, deadlines. If the email has attachments (marked '--- ATTACHMENTS ---'), USE THEIR CONTENTS — pull invoice totals, deadlines, key clauses, concrete numbers/dates from PDFs/docs into the bullets. Be terse.\n\nOUTPUT FORMAT: Put ONLY the bullet points between these exact markers, each on its own line:\n<<<SUMMARY>>>\n- ...\n<<<END>>>\nAny reasoning must come BEFORE <<<SUMMARY>>> (ideally inside <think>...</think>). Only the text between the markers is kept."},
                    {"role": "user", "content": f"From: {sender}\nSubject: {subject}\n\n{body_for_llm[:12000]}\n\n---\n\nSummarize the email. Output the bullets between <<<SUMMARY>>> and <<<END>>>."},
                ],
                tok_key: 8192,
                "temperature": 0.3,
                "stream": False,
            }
            # Reasoning models (o1/o3/o4/gpt-5) reject an explicit temperature.
            if _restricts_temperature(model):
                payload.pop("temperature", None)
            resp = await asyncio.to_thread(
                _req.post, url, json=payload, headers=req_headers, timeout=180
            )
            if not resp.ok:
                return {"success": False, "error": f"LLM HTTP {resp.status_code}"}
            rdata = resp.json()
            msg = (rdata.get("choices") or [{}])[0].get("message", {})
            content = (msg.get("content") or "").strip()
            content = _extract_reply(content)

            if not content:
                # Model put everything in reasoning_content — extract bullet points
                rc = (msg.get("reasoning_content") or "").strip()
                # Find bullet-point style output (lines starting with -, •, *, or numbered)
                bullet_lines = []
                for line in rc.split("\n"):
                    stripped = line.strip()
                    if re.match(r"^[-•*]\s+|^\d+[.)]\s+", stripped):
                        bullet_lines.append(stripped)
                if bullet_lines:
                    content = "\n".join(bullet_lines)
                else:
                    # Last resort: take the last paragraph
                    paragraphs = [p.strip() for p in rc.split("\n\n") if p.strip()]
                    content = paragraphs[-1] if paragraphs else rc[:500]

            if not content:
                return {"success": False, "error": "Empty response from model"}

            # Cache the summary if we have a message_id
            mid = data.get("message_id", "")
            if mid:
                try:
                    import sqlite3 as _sql3
                    _c = _sql3.connect(SCHEDULED_DB)
                    _c.execute("""
                        INSERT OR REPLACE INTO email_summaries
                        (message_id, owner, uid, folder, subject, sender, summary, model_used, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        mid, owner, data.get("uid", ""), data.get("folder", ""),
                        subject, sender, content, model, datetime.utcnow().isoformat(),
                    ))
                    _c.commit()
                    _c.close()
                except Exception as e:
                    logger.warning(f"Failed to cache summary: {e}")

            return {"success": True, "summary": content, "model_used": model}
        except Exception as e:
            logger.error(f"Failed to summarize: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.post("/ai-reply")
    async def ai_reply(data: dict, owner: str = Depends(require_owner)):
        """Generate an AI-drafted reply to an email using the user's writing style."""
        try:
            from src.endpoint_resolver import resolve_endpoint

            to = data.get("to", "")
            subject = data.get("subject", "")
            original_body = data.get("original_body", "")
            requested_model = data.get("model", "").strip()
            session_id = data.get("session_id", "").strip()
            message_id = (data.get("message_id") or "").strip()
            source_uid = (data.get("uid") or "").strip()
            source_folder = (data.get("folder") or "INBOX").strip()
            fast_reply = bool(data.get("fast", False))
            user_hint = (data.get("user_hint") or "").strip()

            if not original_body:
                return {"success": False, "error": "No email body provided"}

            # Skip cache lookup when the caller supplied a user_hint — the
            # cached generic reply doesn't reflect the instructions and
            # would silently override them.
            if message_id and not user_hint:
                try:
                    _c = _sql3.connect(SCHEDULED_DB)
                    owner_clause, owner_params = _email_cache_owner_clause(owner)
                    _row = _c.execute(
                        f"SELECT reply, model_used FROM email_ai_replies WHERE message_id = ? AND {owner_clause}",
                        (message_id, *owner_params),
                    ).fetchone()
                    _c.close()
                    if _row and _row[0]:
                        cached_reply = _apply_email_style_mechanics(_extract_reply(_row[0] or ""))
                        if cached_reply:
                            return {
                                "success": True,
                                "reply": cached_reply,
                                "model_used": _row[1] or "cached",
                                "cached": True,
                            }
                except Exception as e:
                    logger.warning(f"AI reply cache lookup failed: {e}")

            settings = _load_settings()
            style = settings.get("email_writing_style", "")

            # Try session's endpoint first if session_id provided
            url = None
            model = requested_model
            headers = None
            if session_id:
                try:
                    # The chat-session ORM model is `Session`, not `ChatSession`
                    # — the old import threw ImportError, was swallowed by the
                    # except, and left url=None so EVERY reply silently fell back
                    # to the "default" endpoint (wrong model). Its auth lives in
                    # `headers` (JSON), and `endpoint_url` is already the full
                    # chat-completions URL the chat path uses verbatim — so use
                    # those directly rather than rebuilding via a nonexistent
                    # `api_key` field.
                    from core.database import SessionLocal as _SL, Session as _CS
                    _db = _SL()
                    sess = _db.query(_CS).filter(_CS.id == session_id, _CS.owner == owner).first()
                    if sess and sess.endpoint_url:
                        url = sess.endpoint_url
                        # Some sessions stored headers double-encoded (a JSON
                        # string inside the JSON column), so the ORM hands back
                        # a str, not a dict — and llm_call_async's h.update()
                        # then throws "dictionary update sequence element…".
                        # Unwrap until we have a dict (or give up → no headers).
                        _h = sess.headers
                        for _ in range(3):
                            if isinstance(_h, str):
                                try:
                                    _h = json.loads(_h)
                                except Exception:
                                    _h = None
                                    break
                            else:
                                break
                        headers = _h if isinstance(_h, dict) and _h else None
                        if not requested_model:
                            model = sess.model
                    _db.close()
                except Exception as e:
                    logger.warning(f"Failed to read session endpoint: {e}")

            if not url:
                # Match the rest of email AI: prefer the caller's Utility
                # model, then fall back to their Default chat model. Using the
                # global default here could hit a stale provider/key even when
                # chat and summaries worked for the current user.
                url, fallback_model, headers = resolve_endpoint("utility", owner=owner)
                if not url:
                    url, fallback_model, headers = resolve_endpoint("default", owner=owner)
                if not model:
                    model = fallback_model

            if not url or not model:
                return {"success": False, "error": "No LLM endpoint configured"}

            # Resolve the model against what the endpoint actually serves. A
            # stored session model can drift from the server's
            # --served-model-name, giving a 404 "model does not exist". Match
            # by exact id, then basename; fall back to the first served model.
            try:
                from src.llm_core import list_model_ids
                _avail = list_model_ids(url, headers=headers)
                if _avail and model not in _avail:
                    import os as _os
                    _base = _os.path.basename((model or "").rstrip("/"))
                    _match = next((a for a in _avail if _os.path.basename(a.rstrip("/")) == _base), None)
                    model = _match or _avail[0]
            except Exception as _e:
                logger.warning(f"AI reply model resolve failed: {_e}")

            logger.info(f"AI reply using model={model} url={url}")

            # Manual AI Reply should feel immediate. The heavier context mining
            # can involve multiple IMAP folder searches and attachment parsing;
            # reserve that for callers that explicitly opt out of fast mode.
            # Owner-scoped so pre-retrieval never crosses tenants.
            context_snippets, _terms = ([], [])
            if not fast_reply:
                context_snippets, _terms = _pre_retrieve_context(original_body, to, owner=owner)

            # NEW: also pull the last few emails from the original sender +
            # their attachments. The "to" field on this endpoint is the
            # recipient of the *outgoing* reply — that is, the original
            # sender we're answering. So `to` doubles as the address we want
            # the thread context for.
            referenced = ""
            if not fast_reply:
                try:
                    from_addr_for_ctx = email.utils.parseaddr(to or "")[1]
                    referenced = _fetch_sender_thread_context(
                        sender_addr=from_addr_for_ctx,
                        exclude_uid=source_uid,
                        exclude_folder=source_folder,
                        limit=3,
                        owner=owner,
                    )
                except Exception as _e:
                    logger.warning(f"sender-thread-context failed: {_e}")

            system_prompt = _EMAIL_REPLY_SYS_PROMPT_BASE
            if style:
                system_prompt += f"\n\nWRITING STYLE TO MATCH:\n{style}"
            if context_snippets:
                system_prompt += "\n\nRELEVANT CONTEXT FROM PAST EMAILS AND CONTACTS:\n" + "\n\n---\n\n".join(context_snippets[:5])
            if referenced:
                system_prompt += (
                    "\n\nREFERENCED MATERIAL — the last few emails from this sender, "
                    "plus any text extracted from their attachments. Use this to "
                    "answer numbered questions or refer to documents they previously "
                    "sent. Do NOT cite this material verbatim unless the sender "
                    "directly asked about something in it.\n\n" + referenced[:18000]
                )

            user_msg = (
                f"Recipient: {to}\nSubject: {subject}\n\n"
                f"Original email and any current draft:\n{original_body[:6000]}\n\n"
            )
            if user_hint:
                user_msg += (
                    f"User's instructions for THIS reply (follow these — they override "
                    f"defaults like length/tone):\n{user_hint[:2000]}\n\n"
                )
            user_msg += "Draft a reply. Return only the reply body text."

            # Build a candidate chain so a stale session-stored API key
            # (the most common cause of "authentication failed" here)
            # doesn't kill AI Reply outright — fall through to the
            # user's Utility / Default endpoints AND their configured
            # fallback chains. Dedupe by url+model so we don't retry
            # the same broken endpoint.
            from src.llm_core import llm_call_async_with_fallback
            from src.endpoint_resolver import (
                resolve_utility_fallback_candidates,
                resolve_chat_fallback_candidates,
            )
            _seen = set()
            _candidates = []
            def _add(_url, _model, _headers):
                key = (_url or "", _model or "")
                if not _url or not _model or key in _seen:
                    return
                _seen.add(key)
                _candidates.append((_url, _model, _headers))
            # Session endpoint first (may be the broken one).
            _add(url, model, headers)
            # Primary utility endpoint — this is what the user has actually
            # configured as their background-task model, with fresh creds.
            try:
                _u_url, _u_model, _u_headers = resolve_endpoint("utility", owner=owner)
                _add(_u_url, _u_model, _u_headers)
            except Exception:
                pass
            # Primary default chat endpoint — last working chat config.
            try:
                _d_url, _d_model, _d_headers = resolve_endpoint("default", owner=owner)
                _add(_d_url, _d_model, _d_headers)
            except Exception:
                pass
            # Configured fallback chains last.
            for cand in resolve_utility_fallback_candidates(owner=owner) or []:
                _add(*cand)
            for cand in resolve_chat_fallback_candidates(owner=owner) or []:
                _add(*cand)
            try:
                reply = await llm_call_async_with_fallback(
                    _candidates,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_msg},
                    ],
                    temperature=0.7,
                    max_tokens=1024 if fast_reply else 6144,
                    timeout=60 if fast_reply else 180,
                )
            except Exception as e:
                detail = getattr(e, "detail", None) or str(e)
                _attempted = ", ".join(f"{m}@{u.split('/')[2] if '/' in u else u}" for u, m, _ in _candidates) or "no candidates"
                return {"success": False, "error": f"All endpoints failed ({_attempted}): {detail}. Check your API keys in Settings → Services."}

            reply = _apply_email_style_mechanics(_extract_reply(reply or ""))
            if not reply:
                return {"success": False, "error": "LLM returned empty response"}

            # Cache so next click is instant
            if message_id:
                try:
                    _c = _sql3.connect(SCHEDULED_DB)
                    _c.execute("""
                        INSERT OR REPLACE INTO email_ai_replies
                        (message_id, owner, uid, folder, reply, model_used, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    """, (message_id, owner, source_uid, source_folder, reply, model, datetime.utcnow().isoformat()))
                    _c.commit()
                    _c.close()
                except Exception as e:
                    logger.warning(f"Failed to cache ai_reply: {e}")

            return {"success": True, "reply": reply, "model_used": model}
        except Exception as e:
            logger.error(f"Failed to generate AI reply: {e}")
            return {"success": False, "error": "Mail operation failed"}

    @router.get("/style")
    async def get_writing_style(owner: str = Depends(require_user)):
        """Get the current writing style prompt."""
        settings = _load_settings()
        return {"style": settings.get("email_writing_style", "")}

    @router.put("/style")
    async def update_writing_style(data: dict, owner: str = Depends(require_user)):
        """Manually update the writing style prompt."""
        settings = _load_settings()
        settings["email_writing_style"] = data.get("style", "")
        _save_settings(settings)
        return {"success": True}

    @router.get("/config")
    async def get_email_config(owner: str = Depends(require_user)):
        """Get email configuration (passwords masked)."""
        cfg = _get_email_config(owner=owner)
        cfg["smtp_password"] = "***" if cfg["smtp_password"] else ""
        cfg["imap_password"] = "***" if cfg["imap_password"] else ""
        # Include preferences from settings.json
        settings = _load_settings()
        cfg["email_auto_summarize"] = bool(settings.get("email_auto_summarize", False))
        cfg["email_auto_reply"] = bool(settings.get("email_auto_reply", False))
        cfg["email_auto_tag"] = bool(settings.get("email_auto_tag", False))
        cfg["email_auto_spam"] = bool(settings.get("email_auto_spam", False))
        cfg["email_auto_calendar"] = bool(settings.get("email_auto_calendar", False))
        return cfg

    @router.put("/config")
    async def update_email_config(data: dict, owner: str = Depends(require_owner)):
        """Update email configuration.

        Automation flags (email_auto_*) still live in settings.json. Credentials
        are written to the default EmailAccount row. Passwords are only
        overwritten when a non-empty value is provided, so saving the form
        without retyping the password no longer wipes it.
        """
        # Automation flags stay in settings.json (they're global, not per-account)
        settings = _load_settings()
        for key in ["email_auto_summarize", "email_auto_reply", "email_auto_tag", "email_auto_spam", "email_auto_calendar"]:
            if key in data:
                settings[key] = data[key]
        _save_settings(settings)

        # Credentials go into the default account row
        from core.database import SessionLocal, EmailAccount
        import uuid as _uuid
        db = SessionLocal()
        try:
            q = db.query(EmailAccount).filter(EmailAccount.is_default == True)  # noqa: E712
            if owner:
                q = q.filter(EmailAccount.owner == owner)
            row = q.first()
            if row is None:
                row = EmailAccount(id=_uuid.uuid4().hex, owner=owner, name="Default", is_default=True, enabled=True)
                db.add(row)
            field_map = {
                "smtp_host": "smtp_host", "smtp_port": "smtp_port", "smtp_user": "smtp_user",
                "smtp_security": "smtp_security", "imap_host": "imap_host", "imap_port": "imap_port", "imap_user": "imap_user",
                "imap_starttls": "imap_starttls", "email_from": "from_address",
            }
            for in_key, col_name in field_map.items():
                if in_key in data:
                    val = data[in_key]
                    if col_name.endswith("_port") and val in (None, ""):
                        continue
                    if col_name.endswith("_port"):
                        val = int(val)
                    setattr(row, col_name, val)
            # Passwords: only update when a non-empty value is given.
            # Stored encrypted; see src/secret_storage.py.
            from src.secret_storage import encrypt as _enc
            if data.get("imap_password"):
                row.imap_password = _enc(data["imap_password"])
            if data.get("smtp_password"):
                row.smtp_password = _enc(data["smtp_password"])
            clear_q = db.query(EmailAccount).filter(EmailAccount.id != row.id)
            if owner:
                clear_q = clear_q.filter(EmailAccount.owner == owner)
            clear_q.update({EmailAccount.is_default: False})
            db.commit()
        finally:
            db.close()
        return {"success": True}

    # ═══════════════ Urgency state ═══════════════
    # Read-only state file written by `action_check_email_urgency`. The UI
    # uses this to color the unread email dot by urgency tier (3=red,
    # 2=orange, otherwise default blue) and per-row dots in the inbox list.

    @router.get("/urgency-state")
    async def get_email_urgency_state(owner: str = Depends(require_user)):
        from pathlib import Path as _P
        import json as _json
        _slug = "".join(c if (c.isalnum() or c in "-_.@") else "_" for c in (owner or "default"))
        path = _P(DATA_DIR) / f"email_urgency_state_{_slug}.json"
        if not path.exists():
            return {"total_unread": 0, "total_urgent": 0, "max_score": 0, "per_uid": {}}
        try:
            data = _json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return {"total_unread": 0, "total_urgent": 0, "max_score": 0, "per_uid": {}}
        # Drop `notified_uids` from the payload — it's an internal scheduler
        # debounce, not UI-relevant.
        data.pop("notified_uids", None)
        return data

    # ═══════════════ Email Accounts CRUD ═══════════════
    # Multi-account support. Each row is an independent IMAP/SMTP config.
    # Exactly one row has is_default=True; that account is used when callers
    # don't specify an account_id.

    @router.get("/accounts")
    async def list_email_accounts(owner: str = Depends(require_user)):
        """List all email accounts with credentials masked."""
        from core.database import SessionLocal, EmailAccount
        from sqlalchemy import and_, or_
        db = SessionLocal()
        try:
            out = []
            # SECURITY: scope to this user's accounts. Previously returned
            # every row in the EmailAccount table, leaking IMAP/SMTP hosts +
            # usernames across users. Also show legacy unowned rows that match
            # the logged-in mailbox; _get_email_config already accepts those,
            # so Settings should not hide the active account.
            q = db.query(EmailAccount)
            if owner:
                unowned = or_(EmailAccount.owner == None, EmailAccount.owner == "")  # noqa: E711
                same_mailbox = or_(EmailAccount.imap_user == owner, EmailAccount.from_address == owner)
                q = q.filter(or_(EmailAccount.owner == owner, and_(unowned, same_mailbox)))
            for r in q.order_by(
                EmailAccount.is_default.desc(), EmailAccount.created_at.asc()
            ).all():
                out.append({
                    "id": r.id,
                    "name": r.name,
                    "is_default": bool(r.is_default),
                    "enabled": bool(r.enabled),
                    "imap_host": r.imap_host or "",
                    "imap_port": int(r.imap_port or 993),
                    "imap_user": r.imap_user or "",
                    "imap_starttls": bool(r.imap_starttls),
                    "smtp_host": r.smtp_host or "",
                    "smtp_port": int(r.smtp_port or 465),
                    "smtp_security": _smtp_security_mode({"smtp_security": getattr(r, "smtp_security", ""), "smtp_port": r.smtp_port}),
                    "smtp_user": r.smtp_user or "",
                    "from_address": r.from_address or "",
                    "has_imap_password": bool(r.imap_password),
                    "has_smtp_password": bool(r.smtp_password),
                    "oauth_provider": r.oauth_provider or "",
                    "display_name": r.display_name or "",
                })
            return {"accounts": out}
        finally:
            db.close()

    @router.post("/accounts")
    async def create_email_account(data: dict, owner: str = Depends(require_owner)):
        """Create a new email account."""
        from core.database import SessionLocal, EmailAccount
        from src.secret_storage import encrypt as _enc
        import uuid as _uuid
        name = (data.get("name") or "").strip()
        if not name:
            return {"ok": False, "error": "name required"}
        db = SessionLocal()
        try:
            row = EmailAccount(
                id=_uuid.uuid4().hex,
                name=name,
                is_default=bool(data.get("is_default", False)),
                enabled=bool(data.get("enabled", True)),
                imap_host=(data.get("imap_host") or "").strip(),
                imap_port=int(data.get("imap_port") or 993),
                imap_user=(data.get("imap_user") or "").strip(),
                imap_password=_enc(data.get("imap_password") or ""),
                imap_starttls=bool(data.get("imap_starttls", True)),
                smtp_host=(data.get("smtp_host") or "").strip(),
                smtp_port=int(data.get("smtp_port") or 465),
                smtp_security=_smtp_security_mode({"smtp_security": data.get("smtp_security"), "smtp_port": data.get("smtp_port") or 465}),
                smtp_user=(data.get("smtp_user") or "").strip(),
                smtp_password=_enc(data.get("smtp_password") or ""),
                from_address=(data.get("from_address") or "").strip(),
                display_name=(data.get("display_name") or "").strip(),
                # SECURITY: stamp the creator so all subsequent reads / mutations
                # can filter by user. Without this every new account leaks to
                # every other user.
                owner=owner,
            )
            # If there are no accounts yet OR caller asked for default, enforce
            # the one-default invariant — but scope it to THIS user's accounts,
            # otherwise creating a default would clear every other user's
            # default flag too.
            scope_q = db.query(EmailAccount)
            if owner:
                scope_q = scope_q.filter(EmailAccount.owner == owner)
            existing_count = scope_q.count()
            if row.is_default or existing_count == 0:
                scope_q.update({EmailAccount.is_default: False})
                row.is_default = True
            db.add(row)
            db.commit()
            return {"ok": True, "id": row.id}
        finally:
            db.close()

    @router.put("/accounts/{account_id}")
    async def update_email_account(account_id: str, data: dict, owner: str = Depends(require_user)):
        """Update an email account. Passwords only overwrite if non-empty."""
        # Path param account_id — dep validated via Query, re-check the path-param value.
        _assert_owns_account(account_id, owner)
        from core.database import SessionLocal, EmailAccount
        db = SessionLocal()
        try:
            row = db.get(EmailAccount, account_id)
            if not row:
                return {"ok": False, "error": "Account not found"}
            # Simple fields
            for key in ("name", "imap_host", "imap_user", "smtp_host", "smtp_user", "from_address", "display_name"):
                if key in data:
                    setattr(row, key, (data[key] or "").strip())
            for key in ("imap_port", "smtp_port"):
                if data.get(key) not in (None, ""):
                    setattr(row, key, int(data[key]))
            if "smtp_security" in data:
                row.smtp_security = _smtp_security_mode({"smtp_security": data.get("smtp_security"), "smtp_port": data.get("smtp_port") or row.smtp_port})
            for key in ("imap_starttls", "enabled"):
                if key in data:
                    setattr(row, key, bool(data[key]))
            # Passwords — only overwrite when a non-empty value is
            # provided. Stored encrypted; see src/secret_storage.py.
            from src.secret_storage import encrypt as _enc
            if data.get("imap_password"):
                row.imap_password = _enc(data["imap_password"])
            if data.get("smtp_password"):
                row.smtp_password = _enc(data["smtp_password"])
            db.commit()
            return {"ok": True, "id": row.id}
        finally:
            db.close()

    @router.delete("/accounts/{account_id}")
    async def delete_email_account(account_id: str, owner: str = Depends(require_user)):
        _assert_owns_account(account_id, owner)
        from core.database import SessionLocal, EmailAccount
        db = SessionLocal()
        try:
            row = db.get(EmailAccount, account_id)
            if not row:
                return {"ok": False, "error": "Account not found"}
            was_default = bool(row.is_default)
            db.delete(row)
            db.commit()
            # If the deleted row was default, promote the next-oldest enabled
            # row owned by THIS user. Without the owner filter we'd promote
            # another user's account and the deleter would silently inherit
            # it as their default.
            if was_default:
                promote_q = db.query(EmailAccount).filter(EmailAccount.enabled == True)  # noqa: E712
                if owner:
                    promote_q = promote_q.filter(EmailAccount.owner == owner)
                promote = promote_q.order_by(EmailAccount.created_at.asc()).first()
                if promote:
                    promote.is_default = True
                    db.commit()
            return {"ok": True}
        finally:
            db.close()

    @router.post("/accounts/test")
    async def test_account_config(req: Request, owner: str = Depends(require_user)):
        """Try to actually connect to the provided IMAP (and optionally SMTP)
        server with the given credentials. Lets the user verify a config
        BEFORE saving it. Returns per-protocol status so the UI can show
        which half failed.

        If `account_id` is provided (instead of inline credentials), load
        the saved row's stored creds and test those — used by the
        clickable test-dot in the integrations list, where the form has
        no live values."""
        try:
            body = await req.json()
        except Exception:
            return {"ok": False, "imap": {"ok": False, "error": "invalid request body"}}

        # Saved-account shortcut — hydrate missing credentials from the DB row,
        # while keeping any edited form fields from the request. This lets the UI
        # test unsaved host/port changes without forcing the user to retype the
        # stored password.
        # `imap_password` / `smtp_password` are Fernet-encrypted at rest
        # (see _migrate_encrypt_email_passwords); decrypt before use so
        # the test actually sends the real password to the server.
        acc_id = body.get("account_id")
        if acc_id:
            _assert_owns_account(acc_id, owner)
            from core.database import SessionLocal, EmailAccount
            from src.secret_storage import decrypt as _decrypt
            db = SessionLocal()
            try:
                row = db.get(EmailAccount, acc_id)
                if not row:
                    return {"ok": False, "imap": {"ok": False, "error": "Account not found"}}
                saved_body = {
                    "imap_host": row.imap_host or "",
                    "imap_port": row.imap_port or 993,
                    "imap_user": row.imap_user or "",
                    "imap_password": _decrypt(row.imap_password or ""),
                    "imap_starttls": bool(row.imap_starttls),
                    "smtp_host": row.smtp_host or "",
                    "smtp_port": row.smtp_port or 465,
                    "smtp_security": _smtp_security_mode({"smtp_security": getattr(row, "smtp_security", ""), "smtp_port": row.smtp_port}),
                    "smtp_user": row.smtp_user or "",
                    "smtp_password": _decrypt(row.smtp_password or ""),
                }
                for key, value in body.items():
                    if key == "account_id":
                        continue
                    if value not in (None, ""):
                        saved_body[key] = value
                body = saved_body
            finally:
                db.close()

        imap_result = {"ok": False}
        smtp_result = None

        imap_host = (body.get("imap_host") or "").strip()
        imap_port = int(body.get("imap_port") or 993)
        imap_user = (body.get("imap_user") or "").strip()
        imap_pass = body.get("imap_password") or ""
        imap_starttls = bool(body.get("imap_starttls"))

        if not (imap_host and imap_user and imap_pass):
            imap_result = {"ok": False, "error": "Need IMAP host, username, and password"}
        else:
            # Connection mode resolution:
            #   STARTTLS on  → plain IMAP4 + .starttls() (upgrade)
            #   STARTTLS off + port 993 → IMAP4_SSL (implicit SSL, "IMAPS")
            #   STARTTLS off + any other port → plain IMAP4 (no encryption)
            # Without the last branch, local servers exposed on a non-993
            # port (Dovecot on 31143, etc.) would always fail the SSL
            # handshake because they're not actually wrapped in TLS.
            try:
                conn = _open_imap_connection(
                    imap_host,
                    imap_port,
                    starttls=imap_starttls,
                    timeout=_IMAP_TIMEOUT_SECONDS,
                )
                try:
                    conn.login(imap_user, imap_pass)
                    imap_result = {"ok": True}
                finally:
                    try: conn.logout()
                    except Exception: pass
            except Exception as e:
                imap_result = {"ok": False, "error": _friendly_email_auth_error("IMAP", imap_host, e)}

        smtp_host = (body.get("smtp_host") or "").strip()
        if smtp_host:
            smtp_port = int(body.get("smtp_port") or 465)
            smtp_security = _smtp_security_mode({"smtp_security": body.get("smtp_security"), "smtp_port": smtp_port})
            smtp_user = (body.get("smtp_user") or imap_user).strip()
            smtp_pass = body.get("smtp_password") or imap_pass
            try:
                if smtp_security == "ssl":
                    smtp = smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=10)
                else:
                    smtp = smtplib.SMTP(smtp_host, smtp_port, timeout=10)
                    if smtp_security == "starttls":
                        smtp.starttls()
                try:
                    smtp.login(smtp_user, smtp_pass)
                    smtp_result = {"ok": True}
                finally:
                    try: smtp.quit()
                    except Exception: pass
            except Exception as e:
                smtp_result = {"ok": False, "error": _friendly_email_auth_error("SMTP", smtp_host, e)}

        return {
            "ok": imap_result["ok"] and (smtp_result is None or smtp_result["ok"]),
            "imap": imap_result,
            "smtp": smtp_result,
        }

    @router.post("/accounts/{account_id}/set-default")
    async def set_default_account(account_id: str, owner: str = Depends(require_user)):
        _assert_owns_account(account_id, owner)
        from core.database import SessionLocal, EmailAccount
        db = SessionLocal()
        try:
            row = db.get(EmailAccount, account_id)
            if not row:
                return {"ok": False, "error": "Account not found"}
            # SECURITY: scope the "clear other defaults" sweep to this user's
            # accounts so we don't unset another user's default flag.
            clear_q = db.query(EmailAccount)
            if owner:
                clear_q = clear_q.filter(EmailAccount.owner == owner)
            clear_q.update({EmailAccount.is_default: False})
            row.is_default = True
            db.commit()
            return {"ok": True}
        finally:
            db.close()

    # ── Google OAuth2 routes ──

    @router.get("/oauth/google/authorize")
    async def google_oauth_authorize(account_id: str = Query(...), request: Request = None, owner: str = Depends(require_user)):
        import urllib.parse
        _assert_owns_account(account_id, owner)
        client_id = os.environ.get("GOOGLE_OAUTH_CLIENT_ID", "")
        if not client_id:
            raise HTTPException(400, "GOOGLE_OAUTH_CLIENT_ID not set — add it to .env")
        redirect_uri = (
            os.environ.get("GOOGLE_OAUTH_REDIRECT_URI")
            or f"http://{request.headers.get('host', 'localhost:7000')}/api/email/oauth/google/callback"
        )
        state = make_oauth_state(account_id, owner)
        params = urllib.parse.urlencode({
            "client_id": client_id,
            "redirect_uri": redirect_uri,
            "response_type": "code",
            "scope": "https://mail.google.com/ email",
            "access_type": "offline",
            "prompt": "consent",
            "state": state,
        })
        from fastapi.responses import RedirectResponse as _RR
        return _RR(f"https://accounts.google.com/o/oauth2/v2/auth?{params}")

    @router.get("/oauth/google/callback")
    async def google_oauth_callback(
        code: str = Query(None),
        state: str = Query(None),
        error: str = Query(None),
        request: Request = None,
    ):
        import urllib.parse
        from fastapi.responses import RedirectResponse as _RR
        if error:
            return _RR("/?section=integrations&email_oauth_error=google_error")
        if not code or not state:
            return _RR("/?section=integrations&email_oauth_error=missing_code")
        state_data = verify_oauth_state(state)
        if not state_data:
            return _RR("/?section=integrations&email_oauth_error=invalid_state")
        account_id = state_data.get("a", "")
        owner = state_data.get("o", "")
        client_id = os.environ.get("GOOGLE_OAUTH_CLIENT_ID", "")
        client_secret = os.environ.get("GOOGLE_OAUTH_CLIENT_SECRET", "")
        redirect_uri = (
            os.environ.get("GOOGLE_OAUTH_REDIRECT_URI")
            or f"http://{request.headers.get('host', 'localhost:7000')}/api/email/oauth/google/callback"
        )
        import httpx as _httpx
        try:
            resp = _httpx.post("https://oauth2.googleapis.com/token", data={
                "code": code,
                "client_id": client_id,
                "client_secret": client_secret,
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
            }, timeout=10)
            resp.raise_for_status()
            data = resp.json()
        except Exception:
            logger.warning("Google token exchange failed")
            return _RR("/?section=integrations&email_oauth_error=token_exchange_failed")
        access_token = data.get("access_token", "")
        refresh_token = data.get("refresh_token", "")
        expiry = str(int(time.time()) + data.get("expires_in", 3600))
        # Fetch the email address from userinfo so we can auto-fill imap_user.
        email_addr = ""
        display_name = ""
        try:
            ui = _httpx.get("https://www.googleapis.com/oauth2/v1/userinfo",
                            headers={"Authorization": f"Bearer {access_token}"}, timeout=10)
            if ui.is_success:
                ui_data = ui.json()
                email_addr = ui_data.get("email", "")
                display_name = ui_data.get("name", "")
        except Exception:
            pass
        from core.database import SessionLocal, EmailAccount
        from src.secret_storage import encrypt as _enc
        db = SessionLocal()
        try:
            row = db.query(EmailAccount).filter(EmailAccount.id == account_id).first()
            if not row:
                return _RR("/?section=integrations&email_oauth_error=account_not_found")
            # SECURITY: verify the account belongs to the initiating user.
            if owner and row.owner and row.owner != owner:
                logger.warning("OAuth callback owner mismatch — rejecting token write")
                return _RR("/?section=integrations&email_oauth_error=ownership_error")
            row.oauth_provider = "google"
            row.oauth_access_token = _enc(access_token)
            if refresh_token:
                row.oauth_refresh_token = _enc(refresh_token)
            row.oauth_token_expiry = expiry
            # Auto-fill Google IMAP/SMTP settings if not already configured.
            if not row.imap_host:
                row.imap_host = "imap.gmail.com"
                row.imap_port = 993
                row.imap_starttls = False
            if not row.smtp_host:
                row.smtp_host = "smtp.gmail.com"
                row.smtp_port = 587
            if email_addr:
                if not row.imap_user:
                    row.imap_user = email_addr
                if not row.smtp_user:
                    row.smtp_user = email_addr
                if not row.from_address:
                    row.from_address = email_addr
                if not row.name or row.name == row.id:
                    row.name = email_addr
            if display_name and not row.display_name:
                row.display_name = display_name
            db.commit()
        finally:
            db.close()
        return _RR("/?section=integrations&email_oauth_success=1")

    return router
