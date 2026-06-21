# backend/utils/enhanced_file_processor.py
# Enhanced File Processing System
# Supports multiple formats with future-ready architecture for multi-format generation

import logging
import io
import os
import tempfile
from typing import Dict, List, Optional, Union, Any, Tuple
from pathlib import Path
from dataclasses import dataclass
from abc import ABC, abstractmethod
from enum import Enum

logger = logging.getLogger(__name__)

class FileFormat(Enum):
    """Supported file formats for processing and generation"""
    CSV = "csv"
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XML = "xml"
    HTML = "html"
    TXT = "txt"
    MD = "md"
    JSON = "json"
    YAML = "yaml"
    RTF = "rtf"
    ODT = "odt"
    # Phase 2A: Added image formats
    JPG = "jpg"
    JPEG = "jpeg"
    PNG = "png"
    GIF = "gif"
    BMP = "bmp"
    WEBP = "webp"
    SVG = "svg"
    # Phase 2A.2: Added Excel formats
    XLSX = "xlsx"
    XLS = "xls"
    XLSM = "xlsm"
    XLSB = "xlsb"

@dataclass
class FileMetadata:
    """Metadata for processed files"""
    format: FileFormat
    size_bytes: int
    mime_type: str
    encoding: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    author: Optional[str] = None
    title: Optional[str] = None
    created_date: Optional[str] = None
    # Phase 2A: Added image-specific metadata
    image_dimensions: Optional[Tuple[int, int]] = None
    extraction_confidence: Optional[float] = None
    vision_model_used: Optional[str] = None

@dataclass
class ProcessedContent:
    """Container for processed file content and metadata"""
    text_content: str
    metadata: FileMetadata
    structured_data: Optional[Dict] = None
    # Phase 2A: Added extraction results
    extraction_results: Optional[Dict[str, Any]] = None
    # Tables extracted from documents (CSV, DOCX, Excel)
    tables: Optional[List[List[List[str]]]] = None

class FileProcessor(ABC):
    """Abstract base class for file processors"""
    
    @abstractmethod
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the given file"""
        pass
    
    @abstractmethod
    def process(self, file_path: str) -> ProcessedContent:
        """Process the file and return structured content"""
        pass
    
    def generate(self, content: str, output_path: str, **kwargs) -> bool:
        """Generate a file in this format (optional capability)"""
        return False

class CSVProcessor(FileProcessor):
    """Enhanced CSV processor with proper structure handling"""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.csv')
    
    def process(self, file_path: str) -> ProcessedContent:
        import csv
        
        try:
            with open(file_path, 'r', encoding='utf-8', newline='') as file:
                # Detect CSV dialect
                sample = file.read(8192)
                file.seek(0)
                
                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel
                
                reader = csv.reader(file, dialect)
                rows = list(reader)
                
                # Extract headers and data
                headers = rows[0] if rows else []
                data_rows = rows[1:] if len(rows) > 1 else []
                
                # Convert to structured data
                structured_data = {
                    "headers": headers,
                    "rows": data_rows,
                    "total_rows": len(data_rows),
                    "total_columns": len(headers)
                }
                
                # Create text representation
                text_content = "\n".join([",".join(row) for row in rows])
                
                metadata = FileMetadata(
                    format=FileFormat.CSV,
                    size_bytes=os.path.getsize(file_path),
                    word_count=len(text_content.split()),
                    encoding="utf-8",
                    mime_type="text/csv"
                )
                
                return ProcessedContent(
                    text_content=text_content,
                    metadata=metadata,
                    structured_data=structured_data,
                    tables=[rows]  # CSV is essentially one table
                )
                
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            raise
    
    def generate(self, content: str, output_path: str, **kwargs) -> bool:
        """Generate CSV file with proper formatting"""
        try:
            from backend.utils.csv_formatter import format_csv_content
            
            # Apply CSV formatting if not already formatted
            formatted_content = format_csv_content(content, kwargs.get('user_prompt', ''))
            
            with open(output_path, 'w', encoding='utf-8', newline='') as file:
                file.write(formatted_content)
            
            logger.info(f"Generated CSV file: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error generating CSV file {output_path}: {e}")
            return False

class PDFProcessor(FileProcessor):
    """Enhanced PDF processor with extraction capabilities"""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def process(self, file_path: str) -> ProcessedContent:
        try:
            # Try PyPDF2 first, then fallback to LlamaIndex
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata_dict = pdf_reader.metadata or {}
                page_count = len(pdf_reader.pages)
                
                # Extract text from all pages
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                metadata = FileMetadata(
                    format=FileFormat.PDF,
                    size_bytes=os.path.getsize(file_path),
                    page_count=page_count,
                    word_count=len(text_content.split()),
                    author=metadata_dict.get("/Author"),
                    title=metadata_dict.get("/Title"),
                    created_date=str(metadata_dict.get("/CreationDate", "")),
                    mime_type="application/pdf"
                )
                
                return ProcessedContent(
                    text_content=text_content.strip(),
                    metadata=metadata
                )
                
        except ImportError:
            logger.warning("PyPDF2 not available, falling back to LlamaIndex PDF reader")
            return self._process_with_llamaindex(file_path)
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {e}")
            raise
    
    def _process_with_llamaindex(self, file_path: str) -> ProcessedContent:
        """Fallback PDF processing using LlamaIndex"""
        try:
            from llama_index.readers.file import PDFReader
            
            pdf_reader = PDFReader()
            documents = pdf_reader.load_data(file=Path(file_path))
            
            text_content = "\n\n".join([doc.text for doc in documents])
            
            metadata = FileMetadata(
                format=FileFormat.PDF,
                size_bytes=os.path.getsize(file_path),
                page_count=len(documents),
                word_count=len(text_content.split()),
                mime_type="application/pdf"
            )
            
            return ProcessedContent(
                text_content=text_content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"LlamaIndex PDF processing failed for {file_path}: {e}")
            raise
    
    def generate(self, content: str, output_path: str, **kwargs) -> bool:
        """Generate PDF file (requires additional libraries)"""
        try:
            # For now, return False - PDF generation requires reportlab or similar
            logger.warning("PDF generation not implemented - requires additional libraries")
            return False
            
        except Exception as e:
            logger.error(f"Error generating PDF file {output_path}: {e}")
            return False

class DOCXProcessor(FileProcessor):
    """DOCX processor with table and formatting support"""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def process(self, file_path: str) -> ProcessedContent:
        try:
            import docx
            
            # Open the document
            doc = docx.Document(file_path)
            
            # Extract text content
            text_content = ""
            tables = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
                
                # Add table content to text
                for row in table_data:
                    text_content += "\t".join(row) + "\n"
            
            # Extract document properties
            core_props = doc.core_properties
            
            metadata = FileMetadata(
                format=FileFormat.DOCX,
                size_bytes=os.path.getsize(file_path),
                word_count=len(text_content.split()),
                author=core_props.author,
                title=core_props.title,
                created_date=str(core_props.created) if core_props.created else None,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            return ProcessedContent(
                text_content=text_content.strip(),
                metadata=metadata,
                tables=tables
            )
            
        except ImportError:
            logger.error("python-docx library not available for DOCX processing")
            raise ImportError("python-docx library required for DOCX processing")
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            raise
    
    def generate(self, content: str, output_path: str, **kwargs) -> bool:
        """Generate DOCX file with basic formatting"""
        try:
            import docx
            
            doc = docx.Document()
            
            # Add content as paragraphs
            paragraphs = content.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text)
            
            doc.save(output_path)
            logger.info(f"Generated DOCX file: {output_path}")
            return True
            
        except ImportError:
            logger.error("python-docx library not available for DOCX generation")
            return False
        except Exception as e:
            logger.error(f"Error generating DOCX file {output_path}: {e}")
            return False

class XMLProcessor(FileProcessor):
    """Enhanced XML processor with structured data extraction"""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.xml')
    
    def process(self, file_path: str) -> ProcessedContent:
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract text content
            text_content = ET.tostring(root, encoding='unicode', method='text')
            
            # Create structured representation
            structured_data = self._element_to_dict(root)
            
            metadata = FileMetadata(
                format=FileFormat.XML,
                size_bytes=os.path.getsize(file_path),
                word_count=len(text_content.split()),
                encoding="utf-8",
                mime_type="application/xml"
            )
            
            return ProcessedContent(
                text_content=text_content.strip(),
                metadata=metadata,
                structured_data=structured_data
            )
            
        except Exception as e:
            logger.error(f"Error processing XML file {file_path}: {e}")
            raise
    
    def _element_to_dict(self, element):
        """Convert XML element to dictionary"""
        result = {}
        
        # Add attributes
        if element.attrib:
            result['@attributes'] = element.attrib
        
        # Add text content
        if element.text and element.text.strip():
            if len(element) == 0:
                return element.text.strip()
            result['#text'] = element.text.strip()
        
        # Add child elements
        for child in element:
            child_data = self._element_to_dict(child)
            if child.tag in result:
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_data)
            else:
                result[child.tag] = child_data
        
        return result
    
    def generate(self, content: str, output_path: str, **kwargs) -> bool:
        """Generate XML file with proper formatting"""
        try:
            # If content is already XML, save as-is
            # Otherwise, wrap in basic XML structure
            if not content.strip().startswith('<?xml'):
                content = f'<?xml version="1.0" encoding="UTF-8"?>\n<root>\n{content}\n</root>'
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(content)
            
            logger.info(f"Generated XML file: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error generating XML file {output_path}: {e}")
            return False

class ImageProcessor(FileProcessor):
    """Enhanced image processor with OCR capabilities - Phase 2A.1"""
    
    def __init__(self):
        # Import here to avoid circular dependencies
        self.service_available = False
        self.image_extractor = None
        
        try:
            from backend.services.image_content_service import image_extractor
            self.image_extractor = image_extractor
            self.service_available = True
            logger.info("Image content service loaded successfully")
        except ImportError as e:
            logger.warning(f"Image content service not available: {e}")
            logger.info("Image files will be processed with basic descriptions only")
    
    def can_process(self, file_path: str) -> bool:
        """Check if file is a supported image format"""
        # Basic format check even without service
        image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg'}
        return Path(file_path).suffix.lower() in image_extensions
    
    def process(self, file_path: str) -> ProcessedContent:
        """Process image file and extract text content using vision model"""
        try:
            # Get basic file info
            file_size = os.path.getsize(file_path)
            file_ext = Path(file_path).suffix.lower().replace('.', '')
            
            # Try to get image dimensions (basic check)
            image_dimensions = None
            try:
                from PIL import Image
                with Image.open(file_path) as img:
                    image_dimensions = img.size
            except Exception as e:
                logger.debug(f"Could not get image dimensions for {file_path}: {e}")
            
            # Determine file format enum
            format_map = {
                'jpg': FileFormat.JPG,
                'jpeg': FileFormat.JPEG,
                'png': FileFormat.PNG,
                'gif': FileFormat.GIF,
                'bmp': FileFormat.BMP,
                'webp': FileFormat.WEBP,
                'svg': FileFormat.SVG
            }
            file_format = format_map.get(file_ext, FileFormat.JPG)  # Default fallback
            
            # Initialize default metadata
            metadata = FileMetadata(
                format=file_format,
                size_bytes=file_size,
                mime_type=f"image/{file_ext}",
                image_dimensions=image_dimensions
            )
            
            text_content = ""
            extraction_result = None
            
            # Try OCR extraction if service is available
            if self.service_available and self.image_extractor:
                try:
                    extraction_result = self.image_extractor.extract_text_from_image(file_path)
                    
                    if extraction_result.get('success'):
                        text_content = extraction_result.get('text_content', '')
                        
                        # Update metadata with extraction info
                        metadata.extraction_confidence = extraction_result.get('confidence', 0.0)
                        metadata.vision_model_used = extraction_result.get('model_used')
                        metadata.word_count = len(text_content.split()) if text_content else 0
                        
                        # If no text was extracted, create a basic description
                        if not text_content:
                            text_content = f"Image file: {Path(file_path).name} (no text content detected through OCR)"
                        else:
                            logger.info(f"Successfully extracted {len(text_content)} characters from {file_path}")
                    else:
                        # OCR failed but service was available
                        error_msg = extraction_result.get('error', 'Unknown error')
                        text_content = f"Image file: {Path(file_path).name} (OCR extraction failed: {error_msg})"
                        metadata.word_count = len(text_content.split())
                        logger.warning(f"Image extraction failed for {file_path}: {error_msg}")
                        
                except Exception as e:
                    logger.error(f"Error during OCR extraction for {file_path}: {e}")
                    text_content = f"Image file: {Path(file_path).name} (OCR processing error: {str(e)})"
                    metadata.word_count = len(text_content.split())
            else:
                # Service not available - create basic description
                text_content = f"Image file: {Path(file_path).name} (OCR service not available - basic image file indexing)"
                metadata.word_count = len(text_content.split())
                logger.debug(f"Processed image {file_path} without OCR - service not available")
            
            return ProcessedContent(
                text_content=text_content,
                metadata=metadata,
                extraction_results=extraction_result
            )
            
        except Exception as e:
            logger.error(f"Error processing image file {file_path}: {e}")
            # Return basic content even on error
            text_content = f"Image file: {Path(file_path).name} (processing error: {str(e)})"
            metadata = FileMetadata(
                format=FileFormat.JPG,  # Default
                size_bytes=0,
                mime_type="image/unknown",
                word_count=len(text_content.split())
            )
            return ProcessedContent(
                text_content=text_content,
                metadata=metadata
            )

class ExcelProcessor(FileProcessor):
    """Enhanced Excel processor with structured data extraction and formatting"""
    
    def __init__(self):
        # Import here to avoid circular dependencies
        self.service_available = False
        self.excel_extractor = None
        
        try:
            from backend.services.excel_content_service import excel_extractor
            self.excel_extractor = excel_extractor
            self.service_available = True
            logger.info("Excel content service loaded successfully")
        except ImportError as e:
            logger.warning(f"Excel content service not available: {e}")
            logger.info("Excel files will be processed with basic descriptions only")
    
    def can_process(self, file_path: str) -> bool:
        """Check if file is a supported Excel format"""
        excel_extensions = {'.xlsx', '.xls', '.xlsm', '.xlsb'}
        return Path(file_path).suffix.lower() in excel_extensions
    
    def process(self, file_path: str) -> ProcessedContent:
        """Process Excel file and extract structured data"""
        try:
            if not self.service_available or not self.excel_extractor:
                raise ImportError("Excel content service not available for processing")
            
            # Get basic file info
            file_size = os.path.getsize(file_path)
            file_ext = Path(file_path).suffix.lower().replace('.', '')
            
            # Determine file format enum
            format_map = {
                'xlsx': FileFormat.XLSX,
                'xls': FileFormat.XLS,
                'xlsm': FileFormat.XLSM,
                'xlsb': FileFormat.XLSB
            }
            file_format = format_map.get(file_ext, FileFormat.XLSX) # Default fallback
            
            # Initialize default metadata
            metadata = FileMetadata(
                format=file_format,
                size_bytes=file_size,
                mime_type=f"application/{file_ext}",
                word_count=0 # Will be updated by service
            )
            
            # Extract structured data
            extraction_result = self.excel_extractor.extract_excel_content(file_path)
            
            if extraction_result.get('success'):
                structured_data = extraction_result.get('structured_data')
                text_content = extraction_result.get('text_content', '')
                
                # Update metadata with extraction info
                metadata.word_count = len(text_content.split()) if text_content else 0
                metadata.extraction_confidence = extraction_result.get('confidence', 0.0)
                metadata.vision_model_used = extraction_result.get('model_used')
                
                logger.info(f"Successfully extracted Excel content from {file_path}")
            else:
                # Excel extraction failed
                error_msg = extraction_result.get('error', 'Unknown error')
                text_content = f"Excel file: {Path(file_path).name} (Excel extraction failed: {error_msg})"
                metadata.word_count = len(text_content.split())
                logger.warning(f"Excel extraction failed for {file_path}: {error_msg}")
                
            return ProcessedContent(
                text_content=text_content,
                metadata=metadata,
                structured_data=structured_data,
                extraction_results=extraction_result
            )
            
        except ImportError:
            logger.error("Excel content service not available for Excel processing")
            raise ImportError("Excel content service required for Excel processing")
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {e}")
            raise
    
    def generate(self, content: str, output_path: str, **kwargs) -> bool:
        """Generate Excel file (requires additional libraries)"""
        try:
            # For now, return False - Excel generation requires openpyxl or similar
            logger.warning("Excel generation not implemented - requires additional libraries")
            return False
            
        except Exception as e:
            logger.error(f"Error generating Excel file {output_path}: {e}")
            return False

class EnhancedFileProcessor:
    """Main file processor with support for multiple formats"""
    
    def __init__(self):
        self.processors = {
            FileFormat.CSV: CSVProcessor(),
            FileFormat.PDF: PDFProcessor(),
            FileFormat.DOCX: DOCXProcessor(),
            FileFormat.XML: XMLProcessor(),
            # Phase 2A: Added image processors
            FileFormat.JPG: ImageProcessor(),
            FileFormat.JPEG: ImageProcessor(),
            FileFormat.PNG: ImageProcessor(),
            FileFormat.GIF: ImageProcessor(),
            FileFormat.BMP: ImageProcessor(),
            FileFormat.WEBP: ImageProcessor(),
            FileFormat.SVG: ImageProcessor(),
            # Phase 2A.2: Added Excel processors
            FileFormat.XLSX: ExcelProcessor(),
            FileFormat.XLS: ExcelProcessor(),
            FileFormat.XLSM: ExcelProcessor(),
            FileFormat.XLSB: ExcelProcessor(),
        }
        
        # Register additional simple processors
        self._register_simple_processors()
    
    def _register_simple_processors(self):
        """Register simple text-based processors"""
        # These will be implemented as needed
        pass
    
    def detect_format(self, file_path: str) -> Optional[FileFormat]:
        """Detect file format based on extension and content"""
        file_path_lower = file_path.lower()
        
        if file_path_lower.endswith('.csv'):
            return FileFormat.CSV
        elif file_path_lower.endswith('.pdf'):
            return FileFormat.PDF
        elif file_path_lower.endswith(('.docx', '.doc')):
            return FileFormat.DOCX
        elif file_path_lower.endswith('.xml'):
            return FileFormat.XML
        elif file_path_lower.endswith(('.html', '.htm')):
            return FileFormat.HTML
        elif file_path_lower.endswith('.json'):
            return FileFormat.JSON
        elif file_path_lower.endswith(('.yaml', '.yml')):
            return FileFormat.YAML
        elif file_path_lower.endswith('.md'):
            return FileFormat.MD
        elif file_path_lower.endswith('.txt'):
            return FileFormat.TXT
        # Phase 2A: Added image format detection
        elif file_path_lower.endswith(('.jpg', '.jpeg')):
            return FileFormat.JPEG
        elif file_path_lower.endswith('.png'):
            return FileFormat.PNG
        elif file_path_lower.endswith('.gif'):
            return FileFormat.GIF
        elif file_path_lower.endswith('.bmp'):
            return FileFormat.BMP
        elif file_path_lower.endswith('.webp'):
            return FileFormat.WEBP
        elif file_path_lower.endswith('.svg'):
            return FileFormat.SVG
        # Phase 2A.2: Added Excel format detection
        elif file_path_lower.endswith('.xlsx'):
            return FileFormat.XLSX
        elif file_path_lower.endswith('.xls'):
            return FileFormat.XLS
        elif file_path_lower.endswith('.xlsm'):
            return FileFormat.XLSM
        elif file_path_lower.endswith('.xlsb'):
            return FileFormat.XLSB
        
        return None
    
    def can_process(self, file_path: str) -> bool:
        """Check if file can be processed"""
        format_type = self.detect_format(file_path)
        return format_type is not None and format_type in self.processors
    
    def process_file(self, file_path: str) -> Optional[ProcessedContent]:
        """Process a file and return structured content"""
        try:
            format_type = self.detect_format(file_path)
            if not format_type or format_type not in self.processors:
                logger.warning(f"No processor available for file: {file_path}")
                return None
            
            processor = self.processors[format_type]
            return processor.process(file_path)
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None
    
    def generate_file(self, content: str, output_path: str, format_type: FileFormat, **kwargs) -> bool:
        """Generate a file in the specified format"""
        try:
            if format_type not in self.processors:
                logger.error(f"No generator available for format: {format_type}")
                return False
            
            processor = self.processors[format_type]
            return processor.generate(content, output_path, **kwargs)
            
        except Exception as e:
            logger.error(f"Error generating file {output_path}: {e}")
            return False
    
    def batch_process(self, file_paths: List[str]) -> List[ProcessedContent]:
        """Process multiple files in batch"""
        results = []
        
        for file_path in file_paths:
            try:
                result = self.process_file(file_path)
                if result:
                    results.append(result)
                else:
                    logger.warning(f"Failed to process: {file_path}")
            except Exception as e:
                logger.error(f"Error in batch processing {file_path}: {e}")
                continue
        
        return results
    
    def get_supported_formats(self) -> List[FileFormat]:
        """Get list of supported file formats"""
        return list(self.processors.keys())
    
    def get_format_info(self, format_type: FileFormat) -> Dict[str, Any]:
        """Get information about a specific format"""
        format_info = {
            FileFormat.CSV: {
                "name": "Comma Separated Values",
                "extensions": [".csv"],
                "mime_types": ["text/csv"],
                "can_generate": True,
                "can_process": True,
                "features": ["tables", "structured_data", "headers"]
            },
            FileFormat.PDF: {
                "name": "Portable Document Format",
                "extensions": [".pdf"],
                "mime_types": ["application/pdf"],
                "can_generate": False,  # Requires additional libraries
                "can_process": True,
                "features": ["text_extraction", "metadata", "page_count"]
            },
            FileFormat.DOCX: {
                "name": "Microsoft Word Document",
                "extensions": [".docx", ".doc"],
                "mime_types": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
                "can_generate": True,
                "can_process": True,
                "features": ["tables", "formatting", "metadata", "text_extraction"]
            },
            FileFormat.XML: {
                "name": "Extensible Markup Language",
                "extensions": [".xml"],
                "mime_types": ["application/xml", "text/xml"],
                "can_generate": True,
                "can_process": True,
                "features": ["structured_data", "attributes", "hierarchical"]
            },
            # Phase 2A: Added image format info
            FileFormat.JPEG: {
                "name": "JPEG Image",
                "extensions": [".jpg", ".jpeg"],
                "mime_types": ["image/jpeg"],
                "can_generate": False,
                "can_process": True,
                "features": ["ocr", "text_extraction", "vision_model"]
            },
            FileFormat.PNG: {
                "name": "PNG Image",
                "extensions": [".png"],
                "mime_types": ["image/png"],
                "can_generate": False,
                "can_process": True,
                "features": ["ocr", "text_extraction", "vision_model"]
            },
            FileFormat.GIF: {
                "name": "GIF Image",
                "extensions": [".gif"],
                "mime_types": ["image/gif"],
                "can_generate": False,
                "can_process": True,
                "features": ["ocr", "text_extraction", "vision_model"]
            },
            FileFormat.BMP: {
                "name": "Bitmap Image",
                "extensions": [".bmp"],
                "mime_types": ["image/bmp"],
                "can_generate": False,
                "can_process": True,
                "features": ["ocr", "text_extraction", "vision_model"]
            },
            FileFormat.WEBP: {
                "name": "WebP Image",
                "extensions": [".webp"],
                "mime_types": ["image/webp"],
                "can_generate": False,
                "can_process": True,
                "features": ["ocr", "text_extraction", "vision_model"]
            },
            FileFormat.SVG: {
                "name": "SVG Vector Image",
                "extensions": [".svg"],
                "mime_types": ["image/svg+xml"],
                "can_generate": False,
                "can_process": True,
                "features": ["ocr", "text_extraction", "vision_model"]
            },
            # Phase 2A.2: Added Excel format info
            FileFormat.XLSX: {
                "name": "Microsoft Excel Document (XLSX)",
                "extensions": [".xlsx"],
                "mime_types": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
                "can_generate": False,
                "can_process": True,
                "features": ["structured_data", "tables", "text_extraction"]
            },
            FileFormat.XLS: {
                "name": "Microsoft Excel Document (XLS)",
                "extensions": [".xls"],
                "mime_types": ["application/vnd.ms-excel"],
                "can_generate": False,
                "can_process": True,
                "features": ["structured_data", "tables", "text_extraction"]
            },
            FileFormat.XLSM: {
                "name": "Microsoft Excel Document (XLSM)",
                "extensions": [".xlsm"],
                "mime_types": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
                "can_generate": False,
                "can_process": True,
                "features": ["structured_data", "tables", "text_extraction"]
            },
            FileFormat.XLSB: {
                "name": "Microsoft Excel Document (XLSB)",
                "extensions": [".xlsb"],
                "mime_types": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
                "can_generate": False,
                "can_process": True,
                "features": ["structured_data", "tables", "text_extraction"]
            }
        }
        
        return format_info.get(format_type, {})

# Convenience functions for easy integration
def create_file_processor() -> EnhancedFileProcessor:
    """Create a new enhanced file processor instance"""
    return EnhancedFileProcessor()

def process_single_file(file_path: str) -> Optional[ProcessedContent]:
    """Process a single file using the enhanced file processor"""
    processor = create_file_processor()
    return processor.process_file(file_path)

def get_supported_file_formats() -> List[FileFormat]:
    """Get list of all supported file formats"""
    processor = create_file_processor()
    return processor.get_supported_formats()

# Phase 2A: Image-specific convenience functions
def is_image_file_supported(file_path: str) -> bool:
    """Check if an image file is supported for processing"""
    # Basic format check without requiring service imports
    image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg'}
    return Path(file_path).suffix.lower() in image_extensions

def get_image_service_status() -> Dict[str, Any]:
    """Get status of image processing capabilities"""
    try:
        from backend.services.image_content_service import get_image_service_status
        return get_image_service_status()
    except ImportError:
        return {
            'service_available': False,
            'vision_model_available': False,
            'error': 'Image content service not installed',
            'supported_formats': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg'],
            'max_image_size_mb': 10,
            'import_source': 'none',
            'fallback_mode': True
        }

# Phase 2A.2: Excel-specific convenience functions
def is_excel_file_supported(file_path: str) -> bool:
    """Check if an Excel file is supported for processing"""
    # Basic format check without requiring service imports
    excel_extensions = {'.xlsx', '.xls', '.xlsm', '.xlsb'}
    return Path(file_path).suffix.lower() in excel_extensions

def get_excel_service_status() -> Dict[str, Any]:
    """Get status of Excel processing capabilities"""
    try:
        from backend.services.excel_content_service import get_excel_service_status
        return get_excel_service_status()
    except ImportError:
        return {
            'service_available': False,
            'pandas_available': False,
            'openpyxl_available': False,
            'error': 'Excel content service not installed',
            'supported_formats': ['.xlsx', '.xls', '.xlsm', '.xlsb'],
            'max_file_size_mb': 100,
            'fallback_mode': True
        } 