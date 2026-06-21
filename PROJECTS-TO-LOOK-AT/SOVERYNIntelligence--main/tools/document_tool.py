"""
Document Tool for Scout
Creates Word (.docx) and Excel (.xlsx) documents with dealer lead lists,
outreach logs, or any structured data Scout needs to hand off.
"""
import os
from datetime import datetime
from core.tool_base import Tool
from typing import Any, Dict

DOCS_DIR = os.path.join(os.path.dirname(__file__), '..', 'soveryn_memory', 'scout_docs')


class CreateDocumentTool(Tool):
    """
    Create a Word (.docx) or Excel (.xlsx) document.
    Use this to produce dealer lead lists, outreach reports, and summary sheets
    that can be opened in Word or Excel.
    """

    @property
    def name(self) -> str:
        return "create_document"

    @property
    def description(self) -> str:
        return (
            "Create a Word (.docx) or Excel (.xlsx) document. "
            "Use for dealer lead lists, outreach reports, and summary sheets. "
            "For Excel: pass rows as a list of dicts with consistent keys. "
            "For Word: pass sections as a list of {heading, body} dicts. "
            "Returns the file path of the created document."
        )

    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Output filename without extension (e.g. 'NC_Old_Hickory_Leads')"
                },
                "format": {
                    "type": "string",
                    "enum": ["docx", "xlsx"],
                    "description": "Document format: 'docx' for Word, 'xlsx' for Excel"
                },
                "title": {
                    "type": "string",
                    "description": "Document title (shown at top of Word doc or as sheet name in Excel)"
                },
                "rows_json": {
                    "type": "string",
                    "description": "For Excel/Word table: JSON string of a list of dicts, one per row. Keys become column headers. Example: '[{\"name\":\"Dealer\",\"phone\":\"555-1234\",\"email\":\"a@b.com\"}]'"
                },
                "sections_json": {
                    "type": "string",
                    "description": "For Word sections: JSON string of a list of {heading, body} dicts. Example: '[{\"heading\":\"Summary\",\"body\":\"Text here\"}]'"
                }
            },
            "required": ["filename", "format"]
        }

    async def execute(self, filename: str = "", format: str = "xlsx",
                      title: str = "", rows_json: str = "", sections_json: str = "",
                      rows: list = None, sections: list = None, **kw) -> str:
        import json as _json
        # Accept rows/sections as JSON strings (from TOOL_CALL text parsing) or as native lists
        if rows_json and not rows:
            try:
                rows = _json.loads(rows_json)
            except Exception:
                rows = []
        if sections_json and not sections:
            try:
                sections = _json.loads(sections_json)
            except Exception:
                sections = []
        os.makedirs(DOCS_DIR, exist_ok=True)
        if not filename:
            filename = f"scout_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        # Sanitize filename — strip any existing extension then add the correct one
        filename = filename.replace('/', '_').replace('\\', '_')
        for ext in ('.xlsx', '.docx', '.txt', '.csv'):
            if filename.lower().endswith(ext):
                filename = filename[:-len(ext)]
                break
        filename = f"{filename}.{format}"

        out_path = os.path.join(DOCS_DIR, filename)

        try:
            if format == "xlsx":
                if not rows:
                    return "Error: cannot create Excel file — no rows provided. Pass rows_json with at least one row of data."
                return self._create_excel(out_path, title, rows)
            elif format == "docx":
                if not sections and not rows:
                    return "Error: cannot create Word file — no sections or rows provided. Pass sections_json with content."
                return self._create_word(out_path, title, sections or [], rows or [])
            else:
                return f"Unsupported format: {format}"
        except Exception as e:
            return f"Document creation error: {e}"

    def _create_excel(self, path: str, title: str, rows: list) -> str:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = title[:31] if title else "Leads"  # Excel sheet name limit

        if not rows:
            wb.save(path)
            return f"Created empty Excel: {path}"

        # Headers from first row's keys
        headers = list(rows[0].keys())
        header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header.replace('_', ' ').title())
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')

        # Data rows
        for row_idx, row in enumerate(rows, 2):
            for col_idx, key in enumerate(headers, 1):
                ws.cell(row=row_idx, column=col_idx, value=row.get(key, ''))

        # Auto-width columns
        for col in ws.columns:
            max_len = max((len(str(c.value or '')) for c in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

        # Freeze header row
        ws.freeze_panes = "A2"

        wb.save(path)
        return f"Created Excel with {len(rows)} rows: {path}"

    def _create_word(self, path: str, title: str, sections: list, rows: list) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # Sections
        for section in sections:
            h = section.get('heading', '')
            body = section.get('body', '')
            if h:
                doc.add_heading(h, level=2)
            if isinstance(body, list):
                for item in body:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(str(item))
            elif body:
                doc.add_paragraph(str(body))

        # If rows provided, add as a table
        if rows:
            headers = list(rows[0].keys())
            doc.add_heading("Dealer List", level=2)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h.replace('_', ' ').title()
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True

            # Data rows
            for row in rows:
                row_cells = table.add_row().cells
                for i, key in enumerate(headers):
                    row_cells[i].text = str(row.get(key, ''))

        doc.save(path)
        return f"Created Word doc with {len(sections)} sections, {len(rows)} table rows: {path}"
