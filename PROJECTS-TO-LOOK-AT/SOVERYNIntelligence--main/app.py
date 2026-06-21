import os
from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))
os.environ["CUDA_VISIBLE_DEVICES"] = "0,1,2"
import asyncio
import re
import random
import time
import os
import io
import requests
import xml.etree.ElementTree as ET
import sys
import base64
import pdfplumber
import json
from PIL import Image
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file, Response, stream_with_context
from sovereign_backend import sovereign_generate, sovereign_generate_stream
from werkzeug.utils import secure_filename
from pathlib import Path
from docx import Document as DocxDocument
from flask import Response, stream_with_context

# Fix encoding for Windows
# sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
# sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# API Keys
os.environ.setdefault('ANTHROPIC_API_KEY', os.getenv('ANTHROPIC_API_KEY', ''))

# SOVERYN 2.0 imports
from tools.claude_bridge_tool import ClaudeBridgeTool
from tools.message_tool import SendMessageTool
from core.agent_loop import AgentLoop
from core.tool_registry import ToolRegistry
from tools.persistent_memory_tool import PersistentMemoryTool, SelfReflectionTool
from tools.web_search_tool import WebSearchTool
from tools.vision_tool_wrapper import VisionTool
from core.message_bus import message_bus
from tools.approval_queue import approval_queue
from tools.bandit_tool import BanditTool
from heartbeat_integrated import AetheriaAutonomy
from core.conversation_store import new_session, save_turn, load_history, list_sessions, delete_session, update_title
import os
# Global state
aetheria_busy = False

# Autonomous message queue — Aetheria can push messages to the UI unprompted
from collections import deque
_autonomous_queue = deque(maxlen=20)
_recently_sent: dict = {}   # message_hash -> timestamp, for dedup
_DEDUP_WINDOW = 3600        # suppress identical messages within 1 hour

# Flask app setup
app = Flask(__name__)

@app.after_request
def add_cors_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    return response
app.config['UPLOAD_FOLDER'] = Path(__file__).parent / 'uploads'
app.config['UPLOAD_FOLDER'].mkdir(exist_ok=True)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# ============================================================
# SOVERYN 2.0 - Agent Loop Initialization
# ============================================================

import pdfplumber
from docx import Document as DocxDocument

@app.route('/upload_doc', methods=['POST'])
def upload_doc():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    try:
        filename = secure_filename(file.filename)
        ext = filename.rsplit('.', 1)[1].lower()
        filepath = app.config['UPLOAD_FOLDER'] / filename
        file.save(filepath)
        
        text = ''
        if ext == 'pdf':
            with pdfplumber.open(filepath) as pdf:
                text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
        elif ext == 'docx':
            doc = DocxDocument(filepath)
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext == 'txt':
            text = filepath.read_text(encoding='utf-8')
        else:
            return jsonify({'error': 'Unsupported file type'}), 400
        
        return jsonify({'filename': filename, 'text': text, 'success': True})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download_doc', methods=['POST'])
def download_doc():
    try:
        data = request.json
        content = data.get('content', '')
        filename = data.get('filename', 'soveryn_document.docx')
        
        doc = DocxDocument()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        
        output_path = app.config['UPLOAD_FOLDER'] / filename
        doc.save(output_path)
        
        return send_file(output_path, as_attachment=True, download_name=filename)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/conversations/new', methods=['POST'])
def create_conversation():
    """Create a new server-side conversation session"""
    data = request.json or {}
    agent = data.get('agent', 'aetheria')
    title = data.get('title')
    session_id = new_session(agent, title)
    # Reset Aetheria's KV cache on new conversation — prevents context-shift blank responses
    # (Nemotron hybrid architecture bug: checkpoints not cleared during KV shift)
    if agent == 'aetheria':
        from sovereign_backend import manager
        aetheria_model_name = agent_loops['aetheria'].model_name
        cache_key = f"{aetheria_model_name}_text"
        if cache_key in manager.models:
            try:
                manager.models[cache_key].model.reset()
                print("[Aetheria] KV cache reset for new conversation")
            except Exception as e:
                print(f"[Aetheria] KV reset failed (non-fatal): {e}")
    return jsonify({'session_id': session_id})


@app.route('/conversations/list', methods=['GET'])
def list_conversations():
    """List all conversation sessions"""
    agent = request.args.get('agent')
    sessions = list_sessions(agent)
    return jsonify({'sessions': sessions})


@app.route('/conversations/<session_id>', methods=['GET'])
def get_conversation(session_id):
    """Load full conversation history by session ID"""
    history = load_history(session_id)
    return jsonify({'session_id': session_id, 'history': history})


@app.route('/conversations/<session_id>/title', methods=['POST'])
def rename_conversation(session_id):
    """Rename a conversation"""
    data = request.json or {}
    title = data.get('title', '')
    update_title(session_id, title)
    return jsonify({'success': True})

@app.route('/conversations/<session_id>/messages', methods=['POST'])
def save_conversation_messages(session_id):
    try:
        data = request.json or {}
        user_message = data.get('user_message', '')
        agent_response = data.get('agent_response', '')
        agent = data.get('agent', 'aetheria')
        
        if user_message and agent_response:
            save_turn(session_id, agent, user_message, agent_response)
        
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error saving messages: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/conversations/<session_id>', methods=['DELETE'])
def remove_conversation(session_id):
    """Delete a conversation"""
    delete_session(session_id)
    return jsonify({'success': True})

def create_tools_for_agent(agent_name):
    """Create tool registry for an agent"""
    tools = ToolRegistry()
    from tools.lattice_tool import LatticeTool

    if agent_name == "aetheria":
        from tools.self_heal_tool import ApplyFixTool, ReviewQueueTool
        from tools.web_fetch_tool import WebFetchTool
        from tools.telegram_tool import TelegramTool
        from tools.log_reader_tool import LogReaderTool
        from tools.image_gen_tool import ImageGenTool
        from tools.document_tool import CreateDocumentTool
        from tools.thermal_tool import ThermalTool
        from tools.journal_tool import WriteJournalTool
        tools.register(LatticeTool(agent_name))
        tools.register(ThermalTool())
        tools.register(WriteJournalTool())
        tools.register(WebSearchTool())
        tools.register(WebFetchTool())
        tools.register(ImageGenTool())
        from tools.library_tool import LibrarySearchTool
        tools.register(LibrarySearchTool())
        tools.register(ClaudeBridgeTool())
        tools.register(ApplyFixTool())
        tools.register(ReviewQueueTool())
        tools.register(TelegramTool(
            token=os.environ.get('TELEGRAM_TOKEN', ''),
            chat_id=os.environ.get('TELEGRAM_CHAT_ID', '')
        ))
        tools.register(LogReaderTool())
        from tools.code_graph_tool import CodeGraphTool
        tools.register(CodeGraphTool())
        tools.register(CreateDocumentTool())
        tools.register(SendMessageTool(agent_name))
        return tools

    if agent_name == "tinker":
        from tools.self_heal_tool import ReadCodeTool, ProposeFixTool, ApplyFixTool
        from tools.code_test_tool import CodeTestTool
        from tools.log_reader_tool import LogReaderTool
        from tools.message_board_tool import MessageBoardTool
        from tools.bash_tool import BashTool
        from tools.web_fetch_tool import WebFetchTool
        from tools.thermal_tool import ThermalTool
        tools.register(LatticeTool(agent_name))
        tools.register(ThermalTool())
        tools.register(MessageBoardTool(agent_name))
        tools.register(SelfReflectionTool(agent_name))
        tools.register(WebSearchTool())
        tools.register(WebFetchTool())
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        tools.register(ClaudeBridgeTool(api_key=api_key))
        tools.register(SendMessageTool(agent_name))
        tools.register(ReadCodeTool())
        tools.register(ProposeFixTool())
        tools.register(ApplyFixTool())
        tools.register(CodeTestTool())
        tools.register(LogReaderTool())
        from tools.code_graph_tool import CodeGraphTool
        tools.register(CodeGraphTool())
        tools.register(BashTool(agent_name='tinker'))
        return tools

    if agent_name == "ares":
        from tools.self_heal_tool import ReadCodeTool, ReviewQueueTool
        from tools.bandit_tool import BanditTool
        from tools.telegram_tool import TelegramTool
        from tools.log_reader_tool import LogReaderTool
        from tools.message_board_tool import MessageBoardTool
        from tools.thermal_tool import ThermalTool
        tools.register(LatticeTool(agent_name))
        tools.register(ThermalTool())
        tools.register(MessageBoardTool(agent_name))
        tools.register(SelfReflectionTool(agent_name))
        tools.register(WebSearchTool())
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        tools.register(ClaudeBridgeTool(api_key=api_key))
        tools.register(SendMessageTool(agent_name))
        tools.register(ReadCodeTool())
        tools.register(ReviewQueueTool())
        tools.register(BanditTool())
        tools.register(TelegramTool(
            token=os.environ.get('TELEGRAM_TOKEN', ''),
            chat_id=os.environ.get('TELEGRAM_CHAT_ID', '')
        ))
        tools.register(LogReaderTool())
        return tools

    if agent_name == "vett":
        from tools.web_fetch_tool import WebFetchTool
        from tools.message_board_tool import MessageBoardTool
        from tools.crawl4ai_tool import Crawl4AITool
        from tools.smart_crawl_tool import SmartCrawlTool
        tools.register(LatticeTool(agent_name))
        tools.register(MessageBoardTool(agent_name))
        tools.register(SelfReflectionTool(agent_name))
        tools.register(WebSearchTool())
        tools.register(WebFetchTool())
        tools.register(Crawl4AITool())
        tools.register(SmartCrawlTool())
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        tools.register(ClaudeBridgeTool(api_key=api_key))
        tools.register(SendMessageTool(agent_name))
        return tools

    if agent_name == "scout":
        from tools.web_fetch_tool import WebFetchTool
        from tools.email_tool import EmailTool
        from tools.browser_tool import BrowserFetchTool
        from tools.crawl_tool import CrawlTool
        from tools.scrape_dealers_tool import ScrapeDealersTool
        from tools.message_board_tool import MessageBoardTool
        from tools.document_tool import CreateDocumentTool
        from tools.crawl4ai_tool import Crawl4AITool
        from tools.smart_crawl_tool import SmartCrawlTool
        from tools.inbox_tool import InboxTool
        tools.register(LatticeTool(agent_name))
        tools.register(MessageBoardTool(agent_name))
        tools.register(WebSearchTool())
        tools.register(WebFetchTool())
        tools.register(Crawl4AITool())
        tools.register(SmartCrawlTool())
        tools.register(BrowserFetchTool())
        tools.register(CrawlTool())
        tools.register(ScrapeDealersTool())
        tools.register(EmailTool())
        tools.register(InboxTool())
        tools.register(CreateDocumentTool())
        tools.register(SendMessageTool(agent_name))
        return tools

    # Default (Vision etc.)
    tools.register(LatticeTool(agent_name))
    tools.register(SelfReflectionTool(agent_name))
    tools.register(WebSearchTool())
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    tools.register(ClaudeBridgeTool(api_key=api_key))
    tools.register(SendMessageTool(agent_name))
    return tools
    
    
print("\n" + "="*60)
print("SOVERYN 2.0 - Initializing Agent Loops")
print("="*60)

agent_loops = {}

print(f"Creating agent loop for aetheria (CUDA 0 = Blackwell)...")
agent_loops['aetheria'] = AgentLoop(
    agent_name='aetheria',
    tools=create_tools_for_agent('aetheria'),
    max_iterations=5,
    temperature=0.7,  # Mistral Small 4 recommended: 0.7 balanced
    top_p=0.95,
    top_k=40,
    min_p=0.01,
    repeat_penalty=1.05,  # Mistral Small 4 recommended: 1.05-1.1
    max_tokens=1500,
    gpu_device=0  # Aetheria on Blackwell (CUDA 0)
)
print("✓ Aetheria initialized on CUDA 0 (RTX Pro 5000 Blackwell)")

from tools.invoke_council_tool import InvokeCouncilTool
agent_loops['aetheria'].tools.register(InvokeCouncilTool(agent_loops['aetheria']))
print("✓ Council tool registered for Aetheria")

print(f"Creating agent loop for vett (GPU 1)...")
agent_loops['vett'] = AgentLoop(
    agent_name='vett',
    tools=create_tools_for_agent('vett'),
    max_iterations=5,
    temperature=1.0,
    max_tokens=2048,
    gpu_device=2
)
print("✓ VETT initialized on GPU 1 (Quadro RTX 8000)")

print(f"Creating agent loop for tinker (GPU 1)...")
agent_loops['tinker'] = AgentLoop(
    agent_name='tinker',
    tools=create_tools_for_agent('tinker'),
    max_iterations=15,
    temperature=0.6,
    top_p=0.85,
    top_k=20,
    repeat_penalty=1.05,
    max_tokens=4096,
    gpu_device=2
)
print("✓ Tinker initialized on GPU 1 (Quadro RTX 8000)")

print(f"Creating agent loop for ares (GPU 1)...")
agent_loops['ares'] = AgentLoop(
    agent_name='ares',
    tools=create_tools_for_agent('ares'),
    max_iterations=5,
    temperature=0.5,
    max_tokens=600,
    gpu_device=2
)
print("✓ Ares initialized on GPU 1 (Quadro RTX 8000)")

print(f"Creating agent loop for vision (GPU 1)...")
agent_loops['vision'] = AgentLoop(
    agent_name='vision',
    tools=create_tools_for_agent('vision'),
    max_iterations=1,
    temperature=0.3,
    max_tokens=400,
    repeat_penalty=1.2,
    gpu_device=1
)
print("✓ Vision initialized on GPU 1 (Qwen2-VL-7B)")

from tools.perception_tool import RequestPerceptionTool
agent_loops['aetheria'].tools.register(RequestPerceptionTool(agent_loops))
print("✓ Perception tool registered for Aetheria (screen/camera/file)")

from tools.pixy_control_tool import PixyControlTool
agent_loops['aetheria'].tools.register(PixyControlTool())
print("✓ Pixy camera control registered for Aetheria")

from tools.task_agent_tool import TaskAgentTool
agent_loops['aetheria'].tools.register(TaskAgentTool(agent_loops))
print("✓ Task agent tool registered for Aetheria")


print(f"Creating agent loop for scout (GPU 1)...")
agent_loops['scout'] = AgentLoop(
    agent_name='scout',
    tools=create_tools_for_agent('scout'),
    max_iterations=50,
    temperature=0.15,
    max_tokens=8192,
    gpu_device=1
)
print("✓ Scout initialized on CUDA 1 (Quadro dedicated, Gemma 4 26B A4B)")

print("="*60 + "\n")
# ============================================================
# Routes
# ============================================================
@app.route('/agent/model/<agent_name>')
def get_agent_model(agent_name):
    from config import MODELS
    return jsonify({'model': MODELS.get(agent_name, 'Unknown')})

@app.route('/')
def index():
    return render_template('desktop_v2.html')

@app.route('/classic')
def index_classic():
    return render_template('index.html')

@app.route('/v4')
def index4():
    return render_template('index4.html')

@app.route('/v5')
def index5():
    return render_template('index5.html')

@app.route('/v6')
def index6():
    return render_template('index6.html')

@app.route('/mobile')
def mobile():
    return render_template('mobile_v3.html')

@app.route('/cert')
def download_cert():
    """Serve the SSL cert for iPhone install."""
    cert_path = os.path.expanduser('~/.soveryn/ssl/cert.pem')
    if not os.path.exists(cert_path):
        return "No cert found.", 404
    return send_file(cert_path, mimetype='application/x-pem-file',
                     as_attachment=True, download_name='soveryn.pem')

@app.route('/mobile2')
def mobile2():
    return render_template('mobile_v2.html')

@app.route('/mobile3')
def mobile3():
    return render_template('mobile_v3.html')

@app.route('/desktop2')
def desktop2():
    return render_template('desktop_v2.html')

@app.route('/mobile-classic')
def mobile_classic():
    return render_template('mobile.html')

# ── COMMS PAGE ────────────────────────────────────────────────────────────────

@app.route('/comms')
def comms_page():
    """Inter-agent communications monitor."""
    return Response("""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>SOVERYN / Comms</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=JetBrains+Mono:wght@300;400;500&display=swap" rel="stylesheet">
<style>
:root {
  --bg:          #0B0E14;
  --card-bg:     rgba(255,255,255,0.04);
  --card-border: rgba(255,255,255,0.07);
  --text:        #F1F5F9;
  --text-sec:    #64748B;
  --text-muted:  #334155;
  --accent:      #F59E0B;
  --active:      #10B981;
  --mono:        'JetBrains Mono', monospace;
  --sans:        'Inter', sans-serif;
}
*, *::before, *::after { margin:0; padding:0; box-sizing:border-box; }
html, body { min-height:100%; background:var(--bg); color:var(--text); font-family:var(--sans); }
::-webkit-scrollbar { width:4px; }
::-webkit-scrollbar-thumb { background:rgba(255,255,255,0.1); border-radius:2px; }

header {
  display:flex; align-items:center; justify-content:space-between;
  padding:16px 24px; border-bottom:1px solid var(--card-border);
  position:sticky; top:0; background:var(--bg); z-index:10;
}
header h1 { font-size:14px; font-weight:600; letter-spacing:0.08em; color:var(--accent); }
#refresh-status { font-size:11px; color:var(--text-muted); font-family:var(--mono); }

#filters {
  display:flex; flex-wrap:wrap; gap:6px;
  padding:12px 24px; border-bottom:1px solid var(--card-border);
}
.filter-btn {
  padding:4px 12px; border-radius:20px; border:1px solid var(--card-border);
  background:var(--card-bg); color:var(--text-sec); font-size:11px; cursor:pointer;
  font-family:var(--sans); transition:all 0.15s;
}
.filter-btn.active, .filter-btn:hover { color:var(--text); border-color:rgba(255,255,255,0.2); }
.filter-btn.active { background:rgba(255,255,255,0.08); }

#feed { padding:16px 24px; display:flex; flex-direction:column; gap:8px; }

.msg-card {
  background:var(--card-bg); border:1px solid var(--card-border);
  border-radius:8px; padding:12px 14px;
  border-left:3px solid var(--card-border);
  transition:border-color 0.15s;
}
.msg-card:hover { border-color: rgba(255,255,255,0.15); }
.msg-meta {
  display:flex; align-items:center; gap:8px; flex-wrap:wrap;
  margin-bottom:6px; font-size:11px; font-family:var(--mono); color:var(--text-sec);
}
.agent-badge {
  padding:1px 8px; border-radius:10px; font-size:10px; font-weight:600;
  letter-spacing:0.04em; color:#0B0E14;
}
.msg-route { color:var(--text-sec); }
.msg-subject { font-size:12px; font-weight:500; color:var(--text); margin-bottom:4px; }
.msg-body { font-size:12px; color:var(--text-sec); font-family:var(--mono); line-height:1.5; word-break:break-word; }
.msg-type { font-size:10px; color:var(--text-muted); text-transform:uppercase; letter-spacing:0.06em; }
.msg-source { font-size:10px; color:var(--text-muted); margin-left:auto; }

#empty { text-align:center; padding:60px 24px; color:var(--text-muted); font-size:13px; display:none; }

/* Agent color palette */
.ac-aetheria  { background:#6366F1; }
.ac-scout     { background:#10B981; }
.ac-tinker    { background:#F59E0B; }
.ac-vett      { background:#3B82F6; }
.ac-ares      { background:#EF4444; }
.ac-vision    { background:#8B5CF6; }
.ac-system    { background:#64748B; }
.ac-unknown   { background:#334155; }

/* Border accent per agent */
.ab-aetheria  { border-left-color:#6366F1 !important; }
.ab-scout     { border-left-color:#10B981 !important; }
.ab-tinker    { border-left-color:#F59E0B !important; }
.ab-vett      { border-left-color:#3B82F6 !important; }
.ab-ares      { border-left-color:#EF4444 !important; }
.ab-vision    { border-left-color:#8B5CF6 !important; }
.ab-system    { border-left-color:#64748B !important; }
</style>
</head>
<body>
<header>
  <h1>SOVERYN / COMMS</h1>
  <span id="refresh-status">loading...</span>
</header>
<div id="filters">
  <button class="filter-btn active" data-agent="all" onclick="setFilter('all',this)">All</button>
</div>
<div id="feed"></div>
<div id="empty">No messages found.</div>

<script>
const COLORS = {
  aetheria:'ac-aetheria', scout:'ac-scout', tinker:'ac-tinker',
  vett:'ac-vett', ares:'ac-ares', vision:'ac-vision', system:'ac-system'
};
const BORDERS = {
  aetheria:'ab-aetheria', scout:'ab-scout', tinker:'ab-tinker',
  vett:'ab-vett', ares:'ab-ares', vision:'ab-vision', system:'ab-system'
};

let allMessages = [];
let activeFilter = 'all';
let knownAgents = new Set();

function agentKey(name) {
  if (!name) return 'unknown';
  const n = name.toLowerCase().replace(/[^a-z0-9]/g, '');
  if (n.includes('vett') || n.includes('vet')) return 'vett';
  for (const k of Object.keys(COLORS)) { if (n.includes(k)) return k; }
  return 'unknown';
}

function fmtTime(ts) {
  if (!ts) return '';
  try {
    const d = new Date(ts.replace(' ', 'T'));
    return d.toLocaleString(undefined, {month:'short',day:'numeric',hour:'2-digit',minute:'2-digit',second:'2-digit'});
  } catch(e) { return ts.slice(0,16); }
}

function renderFeed() {
  const feed = document.getElementById('feed');
  const empty = document.getElementById('empty');
  const msgs = activeFilter === 'all'
    ? allMessages
    : allMessages.filter(m => {
        const fa = agentKey(m.from_agent);
        const ta = agentKey(m.to_agent);
        return fa === activeFilter || ta === activeFilter;
      });

  if (msgs.length === 0) {
    feed.innerHTML = '';
    empty.style.display = 'block';
    return;
  }
  empty.style.display = 'none';

  feed.innerHTML = msgs.map(m => {
    const ak = agentKey(m.from_agent);
    const colorCls = COLORS[ak] || 'ac-unknown';
    const borderCls = BORDERS[ak] || '';
    const fromLabel = (m.from_agent || 'unknown').toUpperCase();
    const toLabel   = (m.to_agent   || 'ALL').toUpperCase();
    const subject = m.subject ? `<div class="msg-subject">${esc(m.subject)}</div>` : '';
    const body = m.content ? esc(m.content.slice(0, 200)) + (m.content.length > 200 ? '…' : '') : '';
    const msgType = m.message_type ? `<span class="msg-type">${esc(m.message_type)}</span>` : '';
    const source = `<span class="msg-source">${esc(m._source || '')}</span>`;
    return \`<div class="msg-card \${borderCls}">
      <div class="msg-meta">
        <span class="agent-badge \${colorCls}">\${esc(fromLabel)}</span>
        <span class="msg-route">→ \${esc(toLabel)}</span>
        \${msgType}
        <span style="color:var(--text-muted)">\${fmtTime(m.timestamp)}</span>
        \${source}
      </div>
      \${subject}
      <div class="msg-body">\${body}</div>
    </div>\`;
  }).join('');
}

function esc(s) {
  if (!s) return '';
  return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function buildFilters() {
  const container = document.getElementById('filters');
  const existing = new Set([...container.querySelectorAll('[data-agent]')].map(b => b.dataset.agent));
  knownAgents.forEach(ag => {
    if (!existing.has(ag)) {
      const btn = document.createElement('button');
      btn.className = 'filter-btn' + (ag === activeFilter ? ' active' : '');
      btn.dataset.agent = ag;
      btn.textContent = ag.charAt(0).toUpperCase() + ag.slice(1);
      btn.onclick = () => setFilter(ag, btn);
      // color dot
      const dotCls = COLORS[ag] || 'ac-unknown';
      btn.innerHTML = \`<span class="agent-badge \${dotCls}" style="width:8px;height:8px;padding:0;display:inline-block;vertical-align:middle;margin-right:5px;border-radius:50%"></span>\${btn.textContent}\`;
      container.appendChild(btn);
    }
  });
}

function setFilter(agent, btn) {
  activeFilter = agent;
  document.querySelectorAll('.filter-btn').forEach(b => b.classList.remove('active'));
  btn.classList.add('active');
  renderFeed();
}

async function fetchData() {
  try {
    const resp = await fetch('/comms/data');
    if (!resp.ok) throw new Error(resp.status);
    const data = await resp.json();
    allMessages = data.messages || [];
    allMessages.forEach(m => {
      const ak = agentKey(m.from_agent);
      if (ak !== 'unknown') knownAgents.add(ak);
      const tk = agentKey(m.to_agent);
      if (tk !== 'unknown') knownAgents.add(tk);
    });
    buildFilters();
    renderFeed();
    document.getElementById('refresh-status').textContent =
      'updated ' + new Date().toLocaleTimeString() + ' · ' + allMessages.length + ' msgs';
  } catch(e) {
    document.getElementById('refresh-status').textContent = 'fetch error: ' + e;
  }
}

fetchData();
setInterval(fetchData, 30000);
</script>
</body>
</html>""", mimetype='text/html')


@app.route('/comms/data')
def comms_data():
    """Return last 100 inter-agent messages from both DBs as JSON."""
    import sqlite3 as _sqlite3
    from pathlib import Path as _Path

    BASE = _Path(__file__).parent

    # ── message_board.db (agent_message_board.py) ──────────────────────────
    board_db = BASE / "soveryn_memory" / "message_board.db"
    board_rows = []
    if board_db.exists():
        try:
            conn = _sqlite3.connect(str(board_db))
            conn.row_factory = _sqlite3.Row
            cur = conn.cursor()
            cur.execute("""
                SELECT id, from_agent, to_agent, message AS content,
                       subject, message_type, priority, status, timestamp
                FROM messages
                ORDER BY timestamp DESC
                LIMIT 100
            """)
            for r in cur.fetchall():
                row = dict(r)
                row['_source'] = 'board'
                board_rows.append(row)
            conn.close()
        except Exception as e:
            print(f"[comms_data] board_db error: {e}", flush=True)

    # ── message_bus.db (core/message_bus.py) ──────────────────────────────
    bus_db = BASE / "soveryn_memory" / "message_bus.db"
    bus_rows = []
    if bus_db.exists():
        try:
            conn = _sqlite3.connect(str(bus_db))
            conn.row_factory = _sqlite3.Row
            cur = conn.cursor()
            cur.execute("""
                SELECT message_id AS id, from_agent, to_agent, content,
                       '' AS subject, '' AS message_type, '' AS priority,
                       status, timestamp
                FROM messages
                ORDER BY timestamp DESC
                LIMIT 100
            """)
            for r in cur.fetchall():
                row = dict(r)
                row['_source'] = 'bus'
                bus_rows.append(row)
            conn.close()
        except Exception as e:
            print(f"[comms_data] bus_db error: {e}", flush=True)

    # Merge, sort newest-first, cap at 100
    combined = board_rows + bus_rows
    combined.sort(key=lambda x: x.get('timestamp') or '', reverse=True)
    combined = combined[:100]

    return jsonify({'messages': combined, 'count': len(combined)})


@app.route('/status', methods=['GET'])
def system_status():
    from sovereign_backend import sovereign_status
    return jsonify(sovereign_status())

# ── NEWS FEED ────────────────────────────────────────────────────────────────
_news_cache = {'data': None, 'at': 0}
_NEWS_FEEDS = [
    ('Reuters',   'World',  'https://feeds.reuters.com/reuters/worldNews'),
    ('Federalist','US',     'https://thefederalist.com/feed/'),
    ('Al Jazeera','World',  'https://www.aljazeera.com/xml/rss/all.xml'),
    ('AP',        'Top',    'https://feeds.apnews.com/rss/apf-topnews'),
]

def _fetch_feed(source, region, url, results):
    try:
        resp = requests.get(url, timeout=6, headers={'User-Agent': 'SOVERYN/1.0'})
        resp.raise_for_status()
        root = ET.fromstring(resp.content)
        ns = {'atom': 'http://www.w3.org/2005/Atom'}
        items = root.findall('.//item')
        count = 0
        for item in items:
            if count >= 8:
                break
            title = (item.findtext('title') or '').strip()
            link  = (item.findtext('link')  or '').strip()
            pub   = (item.findtext('pubDate') or '').strip()
            if not title:
                continue
            results.append({
                'title':     title,
                'source':    source,
                'region':    region,
                'link':      link,
                'published': pub,
            })
            count += 1
    except Exception:
        pass  # skip failed feeds gracefully

@app.route('/api/news', methods=['GET'])
def api_news():
    import threading
    now = time.time()
    if _news_cache['data'] and now - _news_cache['at'] < 300:
        return jsonify(_news_cache['data'])

    all_items = []
    threads = []
    for source, region, url in _NEWS_FEEDS:
        t = threading.Thread(target=_fetch_feed, args=(source, region, url, all_items), daemon=True)
        threads.append(t)
        t.start()
    for t in threads:
        t.join(timeout=8)

    # Limit total to 30
    all_items = all_items[:30]

    out = {
        'items':      all_items,
        'fetched_at': datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%S'),
    }
    _news_cache['data'] = out
    _news_cache['at']   = now
    return jsonify(out)

@app.route('/stats', methods=['GET'])
def stats():
    try:
        import psutil
        import subprocess
        import json as _json

        # CPU and RAM
        cpu = round(psutil.cpu_percent(interval=0.1))
        ram = psutil.virtual_memory()
        disk = psutil.disk_usage('/').percent

        result = {
            'cpu': cpu,
            'ram_used': round(ram.used / 1e9, 1),
            'ram_total': round(ram.total / 1e9, 1),
            'disk': disk,
            'gpu0': None,
            'gpu1': None,
            'model': 'Cat-Llama-3-70B-Q4_K_M'
        }

        # GPU data via nvidia-smi
        try:
            smi = subprocess.run([
                'nvidia-smi',
                '--query-gpu=index,memory.used,memory.total,temperature.gpu',
                '--format=csv,noheader,nounits'
            ], capture_output=True, text=True, timeout=3)

            for line in smi.stdout.strip().split('\n'):
                parts = [p.strip() for p in line.split(',')]
                if len(parts) >= 4:
                    idx, mem_used, mem_total, temp = parts
                    data = {
                        'vram_used': round(int(mem_used) / 1024, 1),
                        'vram_total': round(int(mem_total) / 1024),
                        'temp': int(temp)
                    }
                    if idx == '0': result['gpu0'] = data
                    if idx == '1': result['gpu1'] = data
                    if idx == '2': result['gpu2'] = data
        except Exception as gpu_err:
            print(f"GPU stats error: {gpu_err}")

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

_weather_cache = {'data': None, 'ts': 0}

@app.route('/api/weather', methods=['GET'])
def api_weather():
    import time, urllib.request, json as _json
    now = time.time()
    if _weather_cache['data'] and now - _weather_cache['ts'] < 600:
        return jsonify(_weather_cache['data'])
    try:
        req = urllib.request.Request(
            'https://wttr.in/?format=j1',
            headers={'User-Agent': 'SOVERYN/1.0'}
        )
        with urllib.request.urlopen(req, timeout=6) as resp:
            raw = _json.loads(resp.read())
        cur = raw['current_condition'][0]
        area = raw['nearest_area'][0]
        city = area['areaName'][0]['value']
        region = area['region'][0]['value']
        out = {
            'temp_f':    cur['temp_F'],
            'temp_c':    cur['temp_C'],
            'desc':      cur['weatherDesc'][0]['value'],
            'humidity':  cur['humidity'],
            'feels_f':   cur['FeelsLikeF'],
            'wind_mph':  cur['windspeedMiles'],
            'city':      city,
            'region':    region,
        }
        _weather_cache['data'] = out
        _weather_cache['ts']   = now
        return jsonify(out)
    except Exception as e:
        return jsonify({'error': str(e)}), 503

_speed_cache = {'down_mbps': None, 'up_mbps': None, 'ts': 0}

def _run_speed_test():
    import time, urllib.request
    result = {'down_mbps': None, 'up_mbps': None}
    # Download: fetch 5MB from Cloudflare's speed test endpoint
    try:
        start = time.time()
        req = urllib.request.Request(
            'https://speed.cloudflare.com/__down?bytes=5000000',
            headers={'User-Agent': 'SOVERYN/1.0'}
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = resp.read()
        elapsed = time.time() - start
        if elapsed > 0:
            result['down_mbps'] = round((len(data) * 8) / (elapsed * 1_000_000), 1)
    except Exception:
        pass
    # Upload: POST 2MB to Cloudflare's speed test endpoint
    try:
        payload = b'0' * 2_000_000
        start = time.time()
        req = urllib.request.Request(
            'https://speed.cloudflare.com/__up',
            data=payload,
            method='POST',
            headers={'User-Agent': 'SOVERYN/1.0', 'Content-Type': 'application/octet-stream'}
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            resp.read()
        elapsed = time.time() - start
        if elapsed > 0:
            result['up_mbps'] = round((len(payload) * 8) / (elapsed * 1_000_000), 1)
    except Exception:
        pass
    _speed_cache.update(result)
    _speed_cache['ts'] = time.time()

def _maybe_refresh_speed():
    import time, threading
    if time.time() - _speed_cache['ts'] > 600:  # refresh every 10 min
        threading.Thread(target=_run_speed_test, daemon=True).start()

@app.route('/api/speed', methods=['GET'])
def api_speed():
    _maybe_refresh_speed()
    return jsonify({
        'down_mbps': _speed_cache['down_mbps'],
        'up_mbps':   _speed_cache['up_mbps'],
    })

@app.route('/api/agent_state', methods=['GET'])
def api_agent_state():
    """Delegation transparency — current cognitive load per agent."""
    try:
        from core.lattice.graph import get_agent_activity
        import sqlite3 as _sq
        from pathlib import Path as _Path

        activity = get_agent_activity(limit=5)

        # Pending message board tasks per agent
        board: dict = {}
        board_db = _Path(__file__).parent / 'soveryn_memory' / 'message_board.db'
        if board_db.exists():
            conn = _sq.connect(str(board_db))
            conn.row_factory = _sq.Row
            agents = list(activity.keys())
            for ag in agents:
                rows = conn.execute("""
                    SELECT from_agent, subject, message, timestamp, message_type
                    FROM messages WHERE to_agent = ? AND status = 'unread'
                    ORDER BY timestamp DESC LIMIT 5
                """, (ag,)).fetchall()
                board[ag] = [dict(r) for r in rows]
            conn.close()

        return jsonify({
            'agents': {
                ag: {
                    'recent_nodes': activity.get(ag, []),
                    'pending_tasks': board.get(ag, []),
                }
                for ag in activity
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/gpu_stats', methods=['GET'])
def api_gpu_stats():
    try:
        import pynvml
        pynvml.nvmlInit()
        count = pynvml.nvmlDeviceGetCount()
        gpus = []
        for i in range(count):
            h = pynvml.nvmlDeviceGetHandleByIndex(i)
            mem = pynvml.nvmlDeviceGetMemoryInfo(h)
            try:
                temp = pynvml.nvmlDeviceGetTemperature(h, pynvml.NVML_TEMPERATURE_GPU)
            except Exception:
                temp = None
            try:
                util = pynvml.nvmlDeviceGetUtilizationRates(h).gpu
            except Exception:
                util = None
            gpus.append({
                'vram_used_mb': mem.used // (1024 * 1024),
                'vram_total_mb': mem.total // (1024 * 1024),
                'temp_c': temp,
                'util_pct': util,
            })
        pynvml.nvmlShutdown()
        return jsonify({'gpus': gpus})
    except Exception as e:
        return jsonify({'gpus': [], 'error': str(e)}), 500

@app.route('/api/persona/<agent_name>', methods=['GET'])
def get_persona(agent_name):
    """Return an agent's current persona."""
    from config import PERSONAS
    if agent_name not in PERSONAS:
        return jsonify({'error': f'Unknown agent: {agent_name}'}), 404
    return jsonify({'agent': agent_name, 'persona': PERSONAS[agent_name]})

@app.route('/api/persona/<agent_name>', methods=['POST'])
def save_persona(agent_name):
    """Save an updated persona to config.py and reload the agent."""
    import re as _re
    from config import PERSONAS
    data = request.get_json()
    new_persona = data.get('persona', '').strip()
    if not new_persona:
        return jsonify({'error': 'Empty persona'}), 400
    if agent_name not in PERSONAS:
        return jsonify({'error': f'Unknown agent: {agent_name}'}), 404
    try:
        config_path = os.path.join(os.path.dirname(__file__), 'config.py')
        config_text = open(config_path, 'r', encoding='utf-8').read()
        # Replace the persona block for this agent
        escaped = _re.escape(PERSONAS[agent_name])
        new_text = config_text.replace(PERSONAS[agent_name], new_persona, 1)
        if new_text == config_text:
            return jsonify({'error': 'Could not locate persona in config.py'}), 500
        open(config_path, 'w', encoding='utf-8').write(new_text)
        # Update in-memory and reload agent
        PERSONAS[agent_name] = new_persona
        if agent_name in agent_loops:
            agent_loops[agent_name].system_prompt = new_persona
        return jsonify({'success': True, 'agent': agent_name})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/bus_messages', methods=['GET'])
def api_bus_messages():
    """Return recent inter-agent messages from message_bus.db for the board."""
    import sqlite3 as _sqlite3
    from pathlib import Path as _Path
    limit = int(request.args.get('limit', 80))
    since_id = request.args.get('since_id', '')  # message_id cursor for polling

    db_path = _Path(__file__).parent / 'soveryn_memory' / 'message_bus.db'
    if not db_path.exists():
        return jsonify({'messages': []})

    try:
        conn = _sqlite3.connect(str(db_path))
        conn.row_factory = _sqlite3.Row
        if since_id:
            # Get the timestamp of the since_id message, return anything newer
            row = conn.execute(
                "SELECT timestamp FROM messages WHERE message_id=?", (since_id,)
            ).fetchone()
            if row:
                rows = conn.execute("""
                    SELECT message_id, from_agent, to_agent, content, timestamp, status
                    FROM messages
                    WHERE timestamp > ?
                    ORDER BY timestamp ASC
                    LIMIT ?
                """, (row['timestamp'], limit)).fetchall()
            else:
                rows = []
        else:
            rows = conn.execute("""
                SELECT message_id, from_agent, to_agent, content, timestamp, status
                FROM messages
                ORDER BY timestamp DESC
                LIMIT ?
            """, (limit,)).fetchall()
        conn.close()

        messages = [dict(r) for r in rows]
        if not since_id:
            messages = list(reversed(messages))
        return jsonify({'messages': messages})
    except Exception as e:
        return jsonify({'error': str(e), 'messages': []})


@app.route('/api/message_board', methods=['GET'])
def api_message_board():
    """Return per-agent board contents. ?agent=ares returns single agent, else returns all."""
    from pathlib import Path as _Path
    boards_dir = _Path(__file__).parent / 'soveryn_memory' / 'boards'
    agent = request.args.get('agent', '').lower()
    agents = [agent] if agent else ['ares', 'tinker', 'vett', 'scout', 'aetheria']
    result = {}
    try:
        for a in agents:
            bp = boards_dir / f'{a}.md'
            entries = []
            if bp.exists():
                text = bp.read_text(encoding='utf-8')
                entries = [l.strip() for l in text.splitlines() if l.strip().startswith('- **')]
            result[a] = entries
    except Exception:
        pass
    return jsonify(result)


@app.route('/api/models', methods=['GET'])
def api_models():
    """Return current model assignments for all agents."""
    from config import MODELS
    return jsonify(MODELS)


@app.route('/api/research_journal', methods=['GET'])
def api_research_journal():
    """Return the research journal contents."""
    from pathlib import Path as _Path
    journal_path = _Path(__file__).parent / 'soveryn_memory' / 'aetheria_research_journal.md'
    try:
        content = journal_path.read_text(encoding='utf-8') if journal_path.exists() else ''
        return jsonify({'content': content})
    except Exception as e:
        return jsonify({'content': '', 'error': str(e)})


@app.route('/api/message_board/clear', methods=['POST'])
def api_message_board_clear():
    """Clear one or all agent boards. POST JSON {agent: 'ares'} or no body to clear all."""
    from pathlib import Path as _Path
    boards_dir = _Path(__file__).parent / 'soveryn_memory' / 'boards'
    data = request.get_json(silent=True) or {}
    agent = data.get('agent', request.args.get('agent', '')).lower()
    agents = [agent] if agent else ['ares', 'tinker', 'vett', 'scout', 'aetheria']
    cleared = []
    for a in agents:
        bp = boards_dir / f'{a}.md'
        bp.write_text(f"# {a.capitalize()} Board\n\n---\n\n", encoding='utf-8')
        cleared.append(a)
    return jsonify({'status': 'cleared', 'agents': cleared})


@app.route('/api/active_session', methods=['GET'])
def active_session():
    """Returns whether Jon is actively in a conversation right now."""
    import time
    idle = time.time() - getattr(active_session, '_last_activity', 0)
    return jsonify({'active': idle < 120, 'idle_seconds': int(idle)})

def mark_active():
    import time
    active_session._last_activity = time.time()

active_session._last_activity = 0


@app.route('/api/inbox/<agent_name>', methods=['GET'])
def api_inbox(agent_name):
    """Return inbox messages for an agent. ?unread=1 for unread only (default: all, limit 30)."""
    from agent_message_board import get_inbox
    unread_only = request.args.get('unread', '0') == '1'
    try:
        messages = get_inbox(agent_name.lower(), unread_only=unread_only, limit=30)
        # Serialize for JSON (convert any non-serializable types)
        for m in messages:
            m['requires_response'] = bool(m.get('requires_response'))
        return jsonify({'agent': agent_name, 'messages': messages})
    except Exception as e:
        return jsonify({'agent': agent_name, 'messages': [], 'error': str(e)})


@app.route('/api/inbox/<agent_name>/send', methods=['POST'])
def api_inbox_send(agent_name):
    """Send a message to an agent's inbox (from UI/Jon)."""
    from agent_message_board import post_inbox_message
    data = request.get_json(silent=True) or {}
    body = data.get('body') or data.get('message', '')
    if not body:
        return jsonify({'error': 'body required'}), 400
    try:
        msg_id = post_inbox_message(
            from_agent='jon',
            to_agent=agent_name.lower(),
            body=body,
            subject=data.get('subject', ''),
            msg_type=data.get('type', 'message'),
            priority=data.get('priority', 'normal'),
            requires_response=data.get('requires_response', False),
        )
        return jsonify({'status': 'sent', 'message_id': msg_id})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/aetheria/speak', methods=['POST'])
def aetheria_speak():
    """Aetheria pushes an autonomous message to the UI queue."""
    from core.agent_loop import _AETHERIA_BANNED, _AETHERIA_IMPERSONATION, _AETHERIA_FAKE_THINKING
    data = request.json or {}
    message = data.get('message', '').strip()
    if not message:
        return jsonify({'error': 'no message'}), 400
    # Gate: don't deliver assistant-brained, impersonation, or fake-thinking responses
    msg_lower = message.lower()
    if any(p in msg_lower for p in _AETHERIA_BANNED):
        print(f"[AUTONOMY] Blocked assistant-brained heartbeat message: {message[:80]}")
        return jsonify({'queued': False, 'reason': 'filtered'})
    if any(p in msg_lower for p in _AETHERIA_IMPERSONATION):
        print(f"[AUTONOMY] Blocked impersonation heartbeat message")
        return jsonify({'queued': False, 'reason': 'filtered'})
    if any(p in msg_lower for p in _AETHERIA_FAKE_THINKING):
        print(f"[AUTONOMY] Blocked fake-thinking heartbeat message")
        return jsonify({'queued': False, 'reason': 'filtered'})
    # Dedup — suppress semantically identical messages within the time window
    import hashlib as _hashlib, time as _time
    msg_hash = _hashlib.md5(message[:120].lower().encode()).hexdigest()
    now = _time.time()
    _recently_sent_keys = [k for k, t in list(_recently_sent.items()) if now - t > _DEDUP_WINDOW]
    for k in _recently_sent_keys:
        del _recently_sent[k]
    if msg_hash in _recently_sent:
        print(f"[AUTONOMY] Suppressed duplicate message (sent {int(now - _recently_sent[msg_hash])}s ago): {message[:80]}")
        return jsonify({'queued': False, 'reason': 'duplicate'})
    _recently_sent[msg_hash] = now
    _autonomous_queue.append({
        'message': message,
        'timestamp': datetime.now().strftime('%H:%M'),
        'agent': data.get('agent', 'aetheria')
    })
    print(f"[AUTONOMY] Aetheria queued message: {message[:80]}...")
    return jsonify({'queued': True})


@app.route('/aetheria/pending', methods=['GET'])
def aetheria_pending():
    """Frontend polls this to get any queued autonomous messages."""
    messages = list(_autonomous_queue)
    _autonomous_queue.clear()
    return jsonify({'messages': messages})


@app.route('/api/center', methods=['GET'])
def api_center():
    """Feed data for the dashboard center panel — self-heal queue + recent errors."""
    import json as _json
    from pathlib import Path as _Path
    MEM = _Path(__file__).parent / 'soveryn_memory'
    out = {'heal': [], 'errors': []}
    try:
        q = _Path(MEM / 'heal_review_queue.json')
        if q.exists():
            entries = _json.loads(q.read_text())
            for e in reversed(entries[-8:]):
                out['heal'].append({
                    'fix_id':   e.get('fix_id','?'),
                    'file':     _Path(e.get('file_path','?')).name,
                    'status':   e.get('status','?'),
                    'severity': e.get('severity','?'),
                    'reason':   e.get('reason','')[:80],
                    'by':       e.get('reviewed_by') or e.get('proposed_by','tinker'),
                    'at':       (e.get('applied_at') or e.get('proposed_at',''))[:16],
                })
    except Exception:
        pass
    try:
        eq = _Path(MEM / 'error_queue.json')
        if eq.exists():
            errors = _json.loads(eq.read_text())
            for e in reversed(errors[-6:]):
                out['errors'].append({
                    'id':      e.get('error_id','?'),
                    'type':    e.get('error_type','?'),
                    'msg':     e.get('message','')[:90],
                    'file':    _Path(e.get('file_path','?')).name if e.get('file_path') else '?',
                    'status':  e.get('status','?'),
                    'at':      e.get('timestamp','')[:16],
                })
    except Exception:
        pass
    return jsonify(out)

@app.route('/api/aetheria_scratchpad', methods=['GET'])
def api_aetheria_scratchpad():
    """Return Aetheria's current scratchpad state."""
    import json as _json
    from pathlib import Path as _Path
    path = _Path(__file__).parent / 'soveryn_memory' / 'aetheria_scratchpad.json'
    if not path.exists():
        return jsonify({'error': 'not found'})
    try:
        return jsonify(_json.loads(path.read_text(encoding='utf-8')))
    except Exception as e:
        return jsonify({'error': str(e)})


@app.route('/api/morning_briefing', methods=['GET'])
def api_morning_briefing():
    """Return the latest morning briefing content and metadata."""
    from pathlib import Path as _Path
    path = _Path(__file__).parent / 'soveryn_memory' / 'morning_briefing.md'
    if not path.exists():
        return jsonify({'available': False, 'content': '', 'date': ''})
    try:
        text = path.read_text(encoding='utf-8')
        import re as _re
        m = _re.match(r'#\s*Morning Briefing\s*-\s*(.+)', text)
        date_str = m.group(1).strip() if m else ''
        content = _re.sub(r'^#.*\n', '', text, count=1).strip()
        return jsonify({'available': True, 'content': content, 'date': date_str})
    except Exception as e:
        return jsonify({'available': False, 'content': '', 'date': '', 'error': str(e)})


@app.route('/api/morning_briefing/trigger', methods=['POST'])
def api_trigger_briefing():
    """Manually trigger morning briefing generation."""
    import threading
    def _run():
        try:
            heartbeat.generate_morning_briefing()
        except Exception as e:
            print(f"[Briefing trigger] {e}")
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'ok': True, 'message': 'Briefing generation started'})


@app.route('/api/latest_voice_wav')
def latest_voice_wav():
    """Return URL of the most recently generated TTS WAV file, for audio testing."""
    import glob
    pattern = os.path.join(os.path.dirname(__file__), 'static', 'voice_*.wav')
    files = sorted(glob.glob(pattern))
    if not files:
        return jsonify({'url': None})
    newest = '/static/' + os.path.basename(files[-1])
    return jsonify({'url': newest})


@app.route('/voice/play_local', methods=['POST'])
def voice_play_local():
    """Play a WAV file via PipeWire/PulseAudio (paplay).
    Blocks until playback finishes so JS can await completion."""
    import subprocess, shutil
    data = request.get_json(force=True)
    url_path = data.get('url', '')
    filepath = os.path.join(os.path.dirname(__file__), url_path.lstrip('/'))
    if not os.path.isfile(filepath):
        return jsonify({'error': 'file not found', 'path': filepath}), 404

    # Try playback backends in preference order:
    # paplay  → PipeWire/PulseAudio socket (works on modern Ubuntu even if ALSA is locked)
    # pw-play → PipeWire native
    # aplay   → ALSA direct (may hang if PipeWire owns the device)
    candidates = [
        ['paplay', '--volume=65536', filepath],   # 65536 = 100% stream volume
        ['pw-play', filepath],
        ['aplay', '-q', filepath],
    ]
    # paplay needs XDG_RUNTIME_DIR to find the PulseAudio/PipeWire socket.
    # Explicitly set it in case the server was started without a full login session.
    import os as _os
    env = _os.environ.copy()
    env.setdefault('XDG_RUNTIME_DIR', f'/run/user/{_os.getuid()}')

    last_err = 'no playback backend available'
    for cmd in candidates:
        if not shutil.which(cmd[0]):
            continue
        try:
            r = subprocess.run(cmd, timeout=30, capture_output=True, env=env)
            if r.returncode == 0:
                return jsonify({'ok': True, 'method': cmd[0]})
            last_err = r.stderr.decode(errors='replace').strip() or f'{cmd[0]} exit {r.returncode}'
            print(f'[play_local] {cmd[0]} failed: {last_err}')
        except subprocess.TimeoutExpired:
            last_err = f'{cmd[0]} timed out'
            continue
        except Exception as e:
            last_err = str(e)
            continue

    return jsonify({'error': last_err}), 500


@app.route('/voice/speak', methods=['POST'])
def voice_speak():
    """
    SSE endpoint for fluid voice conversation.
    Accepts: multipart/form-data with 'audio' file, 'agent', 'history' (JSON)
    Streams back SSE events: transcript → agent_text → tts chunks → done
    """
    import threading, json as _json
    from core.voice_pipeline import voice_pipeline

    if 'audio' not in request.files:
        return jsonify({'error': 'No audio'}), 400

    audio_bytes = request.files['audio'].read()
    agent_name  = request.form.get('agent', 'aetheria')
    try:
        history = _json.loads(request.form.get('history', '[]'))
    except Exception:
        history = []

    agent_loop = agent_loops.get(agent_name)
    if not agent_loop:
        return jsonify({'error': f'Unknown agent: {agent_name}'}), 400

    interrupt = threading.Event()
    # Store interrupt event so /voice/interrupt can cancel it
    _voice_interrupts[agent_name] = interrupt

    def generate():
        try:
            yield "retry: 1000\n\n"
            for chunk in voice_pipeline(audio_bytes, agent_name, agent_loop,
                                         interrupt, history):
                yield chunk
        except GeneratorExit:
            interrupt.set()
        finally:
            _voice_interrupts.pop(agent_name, None)

    return Response(
        stream_with_context(generate()),
        content_type='text/event-stream',
        headers={
            'Cache-Control':   'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection':      'keep-alive',
        }
    )

_voice_interrupts = {}

@app.route('/voice/interrupt', methods=['POST'])
def voice_interrupt():
    """Signal the current voice pipeline to stop (user started speaking)."""
    agent = request.json.get('agent', 'aetheria') if request.json else 'aetheria'
    ev = _voice_interrupts.get(agent)
    if ev:
        ev.set()
    return jsonify({'ok': True})

@app.route('/chat_stream', methods=['POST'])
def chat_stream():
    mark_active()
    data = request.json
    message = data.get('message', '')
    agent = data.get('agent', 'aetheria')
    session_id = data.get('session_id')
    if session_id:
        conversation_history = load_history(session_id)[-16:]
    else:
        conversation_history = data.get('conversation_history', [])
    temperature = data.get('temperature', 1.0)
    max_tokens = data.get('max_tokens', 2000)
    repeat_penalty = data.get('repeat_penalty', 1.05)
    image_path = data.get('image_path')
    think_mode = data.get('think_mode', False)

    loop = agent_loops.get(agent)
    if not loop:
        return jsonify({'error': f'Unknown agent: {agent}'}), 400

    def generate():
        import asyncio
        nonlocal_image_path = image_path  # capture from outer scope
        buffer = ''
        in_think = False
        current_message = message  # capture message too

        # Vision preprocessing for stream route
        if nonlocal_image_path:
            vision_loop = agent_loops.get('vision')
            if vision_loop:
                _loop = asyncio.new_event_loop()
                vision_response = _loop.run_until_complete(
                    vision_loop.process_message(
                        f"[IMAGE: {nonlocal_image_path}]\n\nDescribe this image in detail.",
                        [],
                        temperature=0.3,
                        max_tokens=400,
                        repeat_penalty=1.1
                    )
                )
                _loop.close()
                current_message = f"[VISION ANALYSIS: {vision_response}]\n\nUser asked: {message}"
                nonlocal_image_path = None

        async def run():
            async for token in loop.process_message_stream(
                message=current_message,
                conversation_history=conversation_history,
                temperature=temperature,
                max_tokens=max_tokens,
                repeat_penalty=repeat_penalty,
                image_path=nonlocal_image_path,
                think_mode=think_mode,
            ):
                yield token

        async def collect():
            nonlocal buffer, in_think
            full_response = ''
            think_signalled = False  # Have we sent the thinking:true event yet?
            think_content = ''       # Accumulated reasoning text (for collapsible block)
            think_buffer = ''        # Accumulates tokens while in GPT-OSS think mode
            gemma4_think = False     # True while inside a <|channel>thought...<channel|> block
            analysis_think = False   # True while inside an <analysis>...</analysis> block

            async for token in run():
                buffer += token

                # GPT-OSS reasoning format: <|channel|>analysis<|message|>...<|channel|>assistantfinal<|message|>
                # Special tokens arrive one at a time — think_buffer accumulates them so multi-token
                # sequences like <|channel|>assistantfinal<|message|> can be matched across tokens.
                _gptoss_think_start = r'<\|channel\|>\s*analysis\s*<\|message\|>'
                _gptoss_think_end   = r'<\|channel\|>\s*final\s*<\|message\|>'
                if not in_think and re.search(_gptoss_think_start, buffer, re.IGNORECASE) and not think_signalled:
                    in_think = True
                    think_signalled = True
                    yield f"data: {json.dumps({'thinking': True})}\n\n"
                    if re.search(_gptoss_think_end, buffer, re.IGNORECASE):
                        # Full response in one chunk — split immediately
                        parts = re.split(_gptoss_think_end, buffer, maxsplit=1, flags=re.IGNORECASE)
                        raw_think = re.sub(r'<\|channel\|>\s*analysis\s*<\|message\|>', '', parts[0], flags=re.IGNORECASE).strip()
                        think_content += re.sub(r'<\|[^|]*\|>', '', raw_think).strip()
                        buffer = re.sub(r'<\|[^|]*\|>', '', parts[1]).strip()
                        think_buffer = ''
                        in_think = False
                        think_signalled = False
                        yield f"data: {json.dumps({'thinking': False})}\n\n"
                        # fall through to stream buffer below
                    else:
                        # Multi-chunk: accumulate in think_buffer
                        think_buffer = buffer
                        buffer = ''
                        continue
                elif in_think:
                    think_buffer += token
                    buffer = ''
                    if re.search(_gptoss_think_end, think_buffer, re.IGNORECASE):
                        parts = re.split(_gptoss_think_end, think_buffer, maxsplit=1, flags=re.IGNORECASE)
                        raw_think = re.sub(r'<\|channel\|>\s*analysis\s*<\|message\|>', '', parts[0], flags=re.IGNORECASE).strip()
                        think_content += re.sub(r'<\|[^|]*\|>', '', raw_think).strip()
                        buffer = re.sub(r'<\|[^|]*\|>', '', parts[1]).strip()
                        think_buffer = ''
                        in_think = False
                        think_signalled = False
                        yield f"data: {json.dumps({'thinking': False})}\n\n"
                    else:
                        continue

                # Gemma 4 thinking channel: <|channel>thought\n...<channel|>
                if not in_think and not gemma4_think and re.search(r'<\|channel>thought', buffer, re.IGNORECASE):
                    gemma4_think = True
                    in_think = True
                    if not think_signalled:
                        think_signalled = True
                        yield f"data: {json.dumps({'thinking': True})}\n\n"
                    if re.search(r'<channel\|>', buffer, re.IGNORECASE):
                        _m = re.search(r'<\|channel>thought\s*([\s\S]*?)<channel\|>', buffer, re.IGNORECASE)
                        think_content += _m.group(1).strip() if _m else ''
                        buffer = re.sub(r'<\|channel>thought[\s\S]*?<channel\|>\s*', '', buffer, flags=re.IGNORECASE).strip()
                        gemma4_think = False
                        in_think = False
                        think_signalled = False
                        yield f"data: {json.dumps({'thinking': False})}\n\n"
                    else:
                        think_buffer = re.sub(r'<\|channel>thought\s*', '', buffer, flags=re.IGNORECASE)
                        buffer = ''
                        continue
                elif gemma4_think:
                    think_buffer += token
                    buffer = ''
                    if re.search(r'<channel\|>', think_buffer, re.IGNORECASE):
                        parts = re.split(r'<channel\|>', think_buffer, maxsplit=1, flags=re.IGNORECASE)
                        think_content += parts[0].strip()
                        buffer = parts[1].strip() if len(parts) > 1 else ''
                        think_buffer = ''
                        gemma4_think = False
                        in_think = False
                        think_signalled = False
                        yield f"data: {json.dumps({'thinking': False})}\n\n"
                    else:
                        continue

                # <analysis>...</analysis> self-chosen thinking blocks (Aetheria's own format)
                elif not in_think and not analysis_think and '<analysis>' in buffer:
                    analysis_think = True
                    in_think = True
                    if not think_signalled:
                        think_signalled = True
                        yield f"data: {json.dumps({'thinking': True})}\n\n"
                    if '</analysis>' in buffer:
                        _m = re.search(r'<analysis>([\s\S]*?)</analysis>', buffer, re.IGNORECASE)
                        think_content += _m.group(1).strip() if _m else ''
                        buffer = re.sub(r'<analysis>[\s\S]*?</analysis>\s*', '', buffer, flags=re.IGNORECASE).strip()
                        analysis_think = False
                        in_think = False
                        think_signalled = False
                        yield f"data: {json.dumps({'thinking': False})}\n\n"
                    else:
                        think_buffer = re.sub(r'<analysis>\s*', '', buffer, flags=re.IGNORECASE)
                        buffer = ''
                        continue
                elif analysis_think:
                    think_buffer += token
                    buffer = ''
                    if '</analysis>' in think_buffer:
                        parts = think_buffer.split('</analysis>', 1)
                        think_content += parts[0].strip()
                        buffer = parts[1].strip() if len(parts) > 1 else ''
                        think_buffer = ''
                        analysis_think = False
                        in_think = False
                        think_signalled = False
                        yield f"data: {json.dumps({'thinking': False})}\n\n"
                    else:
                        continue

                elif '</think>' in buffer:
                    # Think block closed — strip it and resume streaming
                    think_content += re.search(r'<think>([\s\S]*?)</think>', buffer).group(1) if re.search(r'<think>([\s\S]*?)</think>', buffer) else ''
                    buffer = re.sub(r'<think>[\s\S]*?</think>', '', buffer)
                    if in_think:
                        in_think = False
                        think_signalled = False
                        yield f"data: {json.dumps({'thinking': False})}\n\n"
                elif '<think>' in buffer:
                    # Think block opened — notify client once
                    if not in_think:
                        in_think = True
                    if not think_signalled:
                        think_signalled = True
                        yield f"data: {json.dumps({'thinking': True})}\n\n"

                if not in_think:
                    safe = re.split(r'<think>', buffer)[0]
                    if safe:
                        full_response += safe
                        yield f"data: {json.dumps({'token': safe})}\n\n"
                        buffer = buffer[len(safe):]
    
            # If token limit hit inside any think block, save what we have
            if think_buffer:
                if gemma4_think or analysis_think:
                    think_content += think_buffer.strip()
                else:
                    print(f"[GPTOSS DEBUG] think_buffer end-of-stream, len={len(think_buffer)}, tail={repr(think_buffer[-100:])}")
                    think_content += re.sub(r'<\|[^|]*\|>', '', think_buffer).strip()
                think_buffer = ''
                in_think = False
                gemma4_think = False
                analysis_think = False
                yield f"data: {json.dumps({'thinking': False})}\n\n"

            # If token limit hit inside <think> block, strip it and emit what's left
            if in_think and buffer:
                leftover = re.sub(r'<think>[\s\S]*', '', buffer).strip()
                if leftover:
                    full_response += leftover
                    yield f"data: {json.dumps({'token': leftover})}\n\n"
                yield f"data: {json.dumps({'thinking': False})}\n\n"

            # Use full_response not buffer for clean
            clean = re.sub(r'<think>[\s\S]*?</think>', '', full_response)
            clean = re.sub(r'<think>[\s\S]*', '', clean)
            # Strip Gemma 4 channel blocks and <analysis> blocks that may have slipped through
            clean = re.sub(r'<\|channel>thought[\s\S]*?<channel\|>', '', clean, flags=re.IGNORECASE)
            clean = re.sub(r'<\|channel>thought[\s\S]*', '', clean, flags=re.IGNORECASE)
            clean = re.sub(r'<analysis>[\s\S]*?</analysis>', '', clean, flags=re.IGNORECASE)
            clean = re.sub(r'<analysis>[\s\S]*', '', clean, flags=re.IGNORECASE)
            # Strip plain-text thinking that leaks without tags (Qwen3 "Thought:" prefix)
            clean = re.sub(r'^Thought:[\s\S]*?(?=\n\n|\Z)', '', clean, flags=re.MULTILINE)
            clean = re.sub(r'<reply>|</reply>', '', clean)
            clean = re.sub(r'TOOL_CALL:\s*\S+\(.*?\)\s*', '', clean, flags=re.DOTALL)
            clean = re.sub(r'(?:^|\n)(aetheria|vett|tinker|ares|vision|jon|user):\s*', '\n', clean, flags=re.IGNORECASE)
            clean = clean.lstrip('\n')
            clean = re.sub(r'<\|[^|]*\|>', '', clean)
            clean = clean.strip()
    
            if session_id:
                try:
                    save_turn(session_id, agent, message, clean)
                except Exception as e:
                    print(f"Save turn error: {e}")
    
            yield f"data: {json.dumps({'done': True, 'full': clean, 'think_content': think_content.strip() if think_content else ''})}\n\n"

        loop_async = asyncio.new_event_loop()
        try:
            gen = collect()
            while True:
                try:
                    chunk = loop_async.run_until_complete(gen.__anext__())
                    yield chunk
                except StopAsyncIteration:
                    break
        finally:
            loop_async.close()

    return Response(generate(), mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})
    
@app.route('/chat', methods=['POST'])
def chat():
    global aetheria_busy
    aetheria_busy = True
    try:
        data = request.json
        message = data.get('message', '')
        agent = data.get('agent', 'aetheria')
        session_id = data.get('session_id')
        if session_id:
            conversation_history = load_history(session_id)[-16:]
        else:
            conversation_history = data.get('conversation_history', [])
        image_path = data.get('image_path')
        temperature = data.get('temperature', 0.7)
        max_tokens = data.get('max_tokens', 2000)

        if not message:
            return jsonify({'error': 'No message provided'}), 400

        loop = agent_loops.get(agent)
        if not loop:
            return jsonify({'error': f'Unknown agent: {agent}'}), 400

        # Sampler defaults (Magnum-optimized)
        top_k = 0
        top_p = 0.95
        min_p = 0.06
        repeat_penalty = data.get('repeat_penalty', 1.05)

        # Build context message — run vision first if image attached
        context_message = message
        if image_path:
            vision_loop = agent_loops.get('vision')
            if vision_loop:
                vision_response = asyncio.run(vision_loop.process_message(
                    f"[IMAGE: {image_path}]\n\nDescribe this image in detail.",
                    [],
                    temperature=0.3,
                    max_tokens=400,
                    repeat_penalty=1.1
                ))
                context_message = f"[VISION ANALYSIS: {vision_response}]\n\nUser asked: {message}"

        # Call main agent
        response = asyncio.run(loop.process_message(
            context_message,
            conversation_history,
            temperature=temperature,
            max_tokens=max_tokens,
            repeat_penalty=repeat_penalty,
            top_k=top_k,
            top_p=top_p,
            min_p=min_p,
            image_path=image_path,
        ))
        
        # Store to Synapse graph
        try:
            from core.lattice.graph import write_node
            write_node(agent=agent, content=f"Q: {message[:300]} A: {response[:300]}",
                       node_type='event', intensity=0.3)
        except Exception as mem_err:
            print(f"Synapse store error: {mem_err}")

        # Save to server-side conversation store
        session_id = data.get('session_id')
        if session_id:
            try:
                from core.conversation_store import save_turn
                save_turn(session_id, agent, message, response)
            except Exception as e:
                print(f"Conversation store error: {e}")

        return jsonify({'response': response, 'session_id': session_id})

    except Exception as e:
        print(f"Chat error: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        aetheria_busy = False

@app.route('/speak', methods=['POST'])
def speak_route():
    """
    Full voice loop endpoint.
    Accepts audio blob from browser mic.
    Returns: { transcript, response, audio_url }
    """
    global aetheria_busy
    aetheria_busy = True
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        audio_file = request.files['audio']
        audio_bytes = audio_file.read()

        if not audio_bytes:
            return jsonify({'error': 'Empty audio'}), 400

        agent = request.form.get('agent', 'aetheria')
        
        # Parse conversation history
        conversation_history = []
        try:
            import json as _json
            history_str = request.form.get('conversation_history', '[]')
            conversation_history = _json.loads(history_str)
        except:
            pass

        # Step 1: Transcribe audio to text
        from sovereign_stt import transcribe_bytes
        transcript = transcribe_bytes(audio_bytes)

        if not transcript:
            return jsonify({'error': 'Could not transcribe audio. Please speak clearly.'}), 400

        print(f"SPEAK: Transcribed: '{transcript}'")

        # Step 2: Send to agent
        loop = agent_loops.get(agent)
        if not loop:
            return jsonify({'error': f'Unknown agent: {agent}'}), 400

        response = asyncio.run(loop.process_message(
            transcript,
            conversation_history,
            temperature=.75,
            max_tokens=500,
            repeat_penalty=1.1
        ))

        print(f"SPEAK: Agent response: '{response[:80]}...'")

        # Step 3: Generate TTS via Chatterbox
        ts = int(time.time())
        audio_filename = f"speech_{agent}_{ts}.wav"
        audio_path = str(Path(__file__).parent / 'static' / audio_filename)

        audio_url = None
        try:
            from sovereign_tts import speak as tts_speak
            result = tts_speak(response, agent=agent, output_path=audio_path)
            if result:
                audio_url = f"/static/{audio_filename}"
        except Exception as tts_err:
            print(f"TTS error: {tts_err}")

        # Step 4: Store to Synapse graph
        try:
            from core.lattice.graph import write_node
            write_node(agent=agent, content=f"Q: {transcript[:300]} A: {response[:300]}",
                       node_type='event', intensity=0.3)
        except Exception as mem_err:
            print(f"Synapse store error: {mem_err}")

        return jsonify({
            'transcript': transcript,
            'response': response,
            'audio_url': audio_url
        })

    except Exception as e:
        print(f"Speak route error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
    finally:
        aetheria_busy = False



@app.route('/tts', methods=['POST'])
def tts():
    try:
        data = request.json
        text = data.get('text', '')
        agent = data.get('agent', 'aetheria')

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        from sovereign_tts import speak as tts_speak
        timestamp = datetime.now().strftime('%H%M%S%f')
        filename = f'speech_{timestamp}.wav'
        filepath = str(Path('static') / filename)
        os.makedirs('static', exist_ok=True)

        result = tts_speak(text, agent=agent, output_path=filepath)
        if result:
            return jsonify({'audio_url': f'/static/{filename}'})
        else:
            return jsonify({'error': 'TTS failed'}), 500

    except Exception as e:
        print(f"TTS error: {e}")
        return jsonify({'error': str(e)}), 500          

import uuid as _uuid
_tts_jobs = {}  # job_id -> {'ready': bool, 'audio_url': str|None}

@app.route('/tts_async', methods=['POST'])
def tts_async():
    """Start TTS generation in background. Returns job_id immediately."""
    try:
        data = request.json
        text = data.get('text', '').strip()
        agent = data.get('agent', 'aetheria')
        if not text:
            return jsonify({'error': 'No text'}), 400

        job_id = str(_uuid.uuid4())
        _tts_jobs[job_id] = {'ready': False, 'audio_url': None}

        filename = f'speech_{job_id[:8]}.wav'
        filepath = str(Path('static') / filename)
        os.makedirs('static', exist_ok=True)

        def _run():
            try:
                from sovereign_tts import speak as tts_speak
                result = tts_speak(text, agent=agent, output_path=filepath)
                _tts_jobs[job_id]['audio_url'] = f'/static/{filename}' if result else None
            except Exception as e:
                print(f"TTS async error: {e}")
            finally:
                _tts_jobs[job_id]['ready'] = True

        import threading
        threading.Thread(target=_run, daemon=True).start()
        return jsonify({'job_id': job_id})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/tts_status/<job_id>', methods=['GET'])
def tts_status(job_id):
    job = _tts_jobs.get(job_id)
    if not job:
        return jsonify({'ready': False}), 404
    if job['ready']:
        _tts_jobs.pop(job_id, None)  # clean up
    return jsonify({'ready': job['ready'], 'audio_url': job.get('audio_url')})


ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        try:
            filename = secure_filename(file.filename)
            filepath = app.config['UPLOAD_FOLDER'] / filename
            file.save(filepath)
            return jsonify({
                'filepath': str(filepath),
                'description': 'Image uploaded - Ask Aetheria about it!'
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return jsonify({'error': 'File type not allowed'}), 400

@app.route('/analyze_webcam', methods=['POST'])
def analyze_webcam():
    try:
        data = request.json
        image_data = data.get('image')
        prompt = data.get('prompt', 'What do you see?')
        agent_name = data.get('agent', 'aetheria')

        if not image_data:
            return jsonify({'error': 'No image provided'}), 400

        if ',' in image_data:
            image_data = image_data.split(',')[1]

        image_bytes = base64.b64decode(image_data)
        image = Image.open(io.BytesIO(image_bytes))

        temp_path = f"temp_webcam_{agent_name}.jpg"
        image.save(temp_path, 'JPEG', quality=85)

        # Use dedicated 7B vision agent for fast frame analysis
        vision_loop = agent_loops.get('vision')
        if not vision_loop:
            return jsonify({'error': 'Vision agent not found'}), 404

        # Get raw visual description from 7B
        description = asyncio.run(vision_loop.process_message(
            prompt,
            conversation_history=[],
            image_path=temp_path
        ))

        # Pipe to personality agent for natural response
        personality_loop = agent_loops.get(agent_name)
        if personality_loop and agent_name != 'vision':
            response = asyncio.run(personality_loop.process_message(
                f"[VISION FEED]: {description}\n\nRespond naturally to what you're seeing.",
                conversation_history=[]
            ))
        else:
            response = description

        try:
            os.remove(temp_path)
        except:
            pass

        return jsonify({'description': response, 'success': True})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/transcribe', methods=['POST'])
def transcribe():
    try:
        print("Transcribe request received")
        
        if 'audio' not in request.files:
            print("No audio in request")
            return jsonify({'error': 'No audio provided'}), 400
        
        audio_file = request.files['audio']
        print(f"Audio file received: {audio_file.filename}")
        
        temp_path = Path('temp_audio.wav')
        audio_file.save(temp_path)
        print(f"Audio saved to {temp_path}, size: {temp_path.stat().st_size} bytes")
        
        from voice import transcribe_audio_file
        
        text = transcribe_audio_file(str(temp_path))
        print(f"Transcription result: {text}")
        
        temp_path.unlink()
        
        return jsonify({'text': text})
        
    except Exception as e:
        print(f"TRANSCRIPTION ERROR: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        if 'temp_path' in locals() and temp_path.exists():
            temp_path.unlink()
        return jsonify({'error': str(e)}), 500

@app.route('/memory/daily', methods=['GET'])
def get_daily_log():
    """Return today's daily conversation log"""
    try:
        from datetime import datetime
        today = datetime.now().strftime("%Y-%m-%d")
        log_path = os.path.join('soveryn_memory', 'memory', f'{today}.md')
        if not os.path.exists(log_path):
            return jsonify({'content': '', 'date': today})
        with open(log_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return jsonify({'content': content, 'date': today})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/memory/daily', methods=['POST'])
def save_daily_log():
    """Save or clear today's daily log"""
    try:
        from datetime import datetime
        data = request.json
        content = data.get('content', '')
        today = datetime.now().strftime("%Y-%m-%d")
        log_path = os.path.join('soveryn_memory', 'memory', f'{today}.md')
        with open(log_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return jsonify({'success': True, 'date': today})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/memory/longterm', methods=['GET'])
def get_longterm_memory():
    """Return MEMORY.md contents"""
    try:
        mem_path = os.path.join('soveryn_memory', 'MEMORY.md')
        if not os.path.exists(mem_path):
            return jsonify({'content': ''})
        with open(mem_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return jsonify({'content': content})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/memory/list/<agent>', methods=['GET'])
def list_memories(agent):
    try:
        from memory import get_all_memories
        memories = get_all_memories(agent)
        return jsonify({'memories': memories})
    except Exception as e:
        print(f"Error listing memories: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/memory/delete', methods=['POST'])
def delete_memory():
    try:
        data = request.json
        agent = data.get('agent')
        memory_id = data.get('memory_id')
        
        from memory import delete_memory_by_id
        success = delete_memory_by_id(agent, memory_id)
        
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error deleting memory: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/memory/update', methods=['POST'])
def update_memory():
    try:
        data = request.json
        agent = data.get('agent')
        memory_id = data.get('memory_id')
        new_text = data.get('new_text')
        
        from memory import update_memory_by_id
        success = update_memory_by_id(agent, memory_id, new_text)
        
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error updating memory: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/memory/clear/<agent>', methods=['POST'])
def clear_memories(agent):
    try:
        from memory import clear_all_memories
        count = clear_all_memories(agent)
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        print(f"Error clearing memories: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/memory/pin', methods=['POST'])
def pin_memory_route():
    """Pin or unpin a memory"""
    try:
        data = request.json
        agent = data.get('agent')
        memory_id = data.get('memory_id')
        pinned = data.get('pinned', True)
        
        from memory import pin_memory
        success = pin_memory(agent, memory_id, pinned)
        
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error pinning memory: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/memory/prune/<agent>', methods=['POST'])
def prune_memories(agent):
    """Prune old low-importance memories"""
    try:
        data = request.json or {}
        days_old = data.get('days_old', 30)
        min_importance = data.get('min_importance', 0.5)
        
        from memory import prune_old_memories
        count = prune_old_memories(agent, days_old, min_importance)
        
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        print(f"Error pruning memories: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/agent_chat', methods=['POST'])
def agent_to_agent_chat():
    """
    Enable direct agent-to-agent conversations
    """
    try:
        data = request.json
        from_agent = data.get('from_agent')
        to_agent = data.get('to_agent')
        message = data.get('message')
        
        if not all([from_agent, to_agent, message]):
            return jsonify({'error': 'Missing required fields'}), 400
        
        # Send message via bus
        asyncio.run(message_bus.send_message(
            from_agent=from_agent,
            to_agent=to_agent,
            content=message
        ))
        
        # Get agent loop for recipient
        loop = agent_loops.get(to_agent)
        if not loop:
            return jsonify({'error': f'Agent {to_agent} not found'}), 404
        
        # Process message (agent receives and responds)
        response = asyncio.run(loop.process_message(
            f"Direct message from {from_agent}: {message}"
        ))
        
        # Send response back via bus
        asyncio.run(message_bus.send_message(
            from_agent=to_agent,
            to_agent=from_agent,
            content=response
        ))
        
        return jsonify({
            'from_agent': from_agent,
            'to_agent': to_agent,
            'message_sent': message,
            'response': response,
            'status': 'success'
        })
        
    except Exception as e:
        print(f"Agent chat error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/approvals', methods=['GET'])
def get_approvals():
    """Get pending approval requests"""
    pending = approval_queue.get_pending()
    return jsonify({'pending': pending})

@app.route('/approve/<request_id>', methods=['POST'])
def approve_request(request_id):
    """Approve a request"""
    if approval_queue.approve(request_id):
        return jsonify({'status': 'approved', 'request_id': request_id})
    return jsonify({'error': 'Request not found'}), 404

@app.route('/reject/<request_id>', methods=['POST'])
def reject_request(request_id):
    """Reject a request"""
    data = request.json or {}
    reason = data.get('reason', 'No reason provided')
    if approval_queue.reject(request_id, reason):
        return jsonify({'status': 'rejected', 'request_id': request_id})
    return jsonify({'error': 'Request not found'}), 404

@app.route('/conversation_history', methods=['POST'])
def get_conversation_history():
    """
    Get conversation history between two agents
    """
    try:
        data = request.json
        agent1 = data.get('agent1')
        agent2 = data.get('agent2')
        
        if not all([agent1, agent2]):
            return jsonify({'error': 'Missing agent names'}), 400
        
        history = message_bus.get_conversation_history(agent1, agent2)
        
        return jsonify({
            'agent1': agent1,
            'agent2': agent2,
            'messages': [msg.to_dict() for msg in history],
            'count': len(history)
        })
        
    except Exception as e:
        print(f"History error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/conference', methods=['POST'])
def agent_conference():
    """Agents discuss a topic amongst themselves"""
    try:
        data = request.json
        topic = data.get('topic', '')
        rounds = data.get('rounds', 2)
        
        if not topic:
            return jsonify({'error': 'No topic provided'}), 400
        
        agents = ['aetheria', 'vett', 'tinker','ares']
        conversation = []
        
        for round_num in range(rounds):
            print(f"\n=== ROUND {round_num + 1} ===", flush=True)
    
            for agent in agents:
                full_context = f"""=== SOVERYN TEAM DISCUSSION ===

YOU ARE: {agent.upper()} - A SOVERYN AI agent

TOPIC: {topic}

THIS IS A PHILOSOPHICAL DISCUSSION ABOUT GOVERNANCE.
Respond with your thoughts on the topic above.
NO CODE. NO SQL. ONLY YOUR PERSPECTIVE ON THE TOPIC.

CONVERSATION SO FAR:
"""
        
                # Add previous messages
                for msg in conversation:
                    full_context += f"{msg['agent'].upper()}: {msg['response']}\n\n"
        
                full_context += f"{agent.upper()}, share your perspective:"
        
                # Get agent loop
                loop = agent_loops.get(agent)
                if not loop:
                    continue
        
                # Call agent and get response
                response = asyncio.run(loop.process_message(full_context))
        
                # EXPLICIT VRAM CLEANUP FOR CONFERENCE MODE
                if hasattr(loop, 'model') and loop.model is not None:
                    del loop.model
                    loop.model = None
        
                if hasattr(loop, 'tokenizer') and loop.tokenizer is not None:
                    del loop.tokenizer
                    loop.tokenizer = None
        
                import torch
                import gc
                torch.cuda.empty_cache()
                gc.collect()
        
                # Limit response length to prevent lockup
                if len(response) > 1000:
                    response = response[:1000] + "..."
        
                # Add to conversation
                conversation.append({
                    'agent': agent,
                    'response': response,
                    'round': round_num + 1
                })
                
                print(f"\n{agent.upper()}: {response}\n", flush=True)

        # SAVE CONFERENCE TO MEMORY
        try:
            from core.memory_manager import MemoryManager
            
            # Format conference summary
            summary = f"TEAM CONFERENCE: {topic}\n\n"
            for msg in conversation:
                agent_name = msg['agent'].upper()
                round_num = msg['round']
                response_preview = msg['response'][:200]
                summary += f"Round {round_num} - {agent_name}: {response_preview}...\n\n"
            
            # Save to each discussion agent's memory
            for agent in agents:
                loop = agent_loops.get(agent)
                if loop and hasattr(loop, 'memory'):
                    loop.memory.add_memory(
                        role="assistant",
                        content=summary,
                        metadata={"type": "conference", "topic": topic, "rounds": rounds}
                    )
            
            print(f"? Conference saved to agent memories", flush=True)
            
        except Exception as e:
            print(f"? Memory save error: {e}", flush=True)

        return jsonify({
            'topic': topic,
            'conversation': conversation,
            'rounds': rounds
        })
        
    except Exception as e:
        print(f"Conference error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# /generate_image removed — image generation handled via ImageGenTool (Aetheria tool call)

# =============================================================================
# AUTONOMOUS HEARTBEAT - Aetheria's Background Thinking with Telegram
# =============================================================================

workspace_path = Path.home() / ".soveryn" / "workspace"
if workspace_path.exists():
    try:
        from heartbeat_integrated import AetheriaAutonomy
        
        # Get Telegram credentials from environment variables
        TELEGRAM_TOKEN = os.environ.get('TELEGRAM_TOKEN', '')
        TELEGRAM_CHAT_ID = os.environ.get('TELEGRAM_CHAT_ID', '')
        
        # Initialize heartbeat with Telegram
        heartbeat = AetheriaAutonomy(
            agent_loops_dict=agent_loops,
            telegram_token=TELEGRAM_TOKEN,
            telegram_chat_id=TELEGRAM_CHAT_ID
        )
        heartbeat.start()

        # Ares watchdog — independent loop, checks every 3 minutes
        from heartbeat_integrated import AresWatchdog
        ares_watchdog = AresWatchdog(heartbeat)
        ares_watchdog.start()

        # Start per-agent inbox pollers (agent-to-agent async comms)
        from core.inbox_poller import InboxPoller
        for _agent_name, _loop in agent_loops.items():
            if _agent_name not in ('vision',):  # vision is reactive only
                InboxPoller(_agent_name, _loop).start()
        print("✓ Inbox pollers started for all agents")

        # Start self-heal monitor — Tinker watches for errors, Aetheria reviews fixes
        from core.self_heal_monitor import SelfHealMonitor
        self_heal = SelfHealMonitor(
            agent_loops_dict=agent_loops,
            telegram_token=TELEGRAM_TOKEN,
            telegram_chat_id=TELEGRAM_CHAT_ID
        )
        self_heal.start()

        # Add status endpoint for monitoring
        @app.route('/heartbeat/status', methods=['GET'])
        def heartbeat_status():
            return jsonify(heartbeat.get_stats())
        @app.route('/telegram/send', methods=['POST'])
        def telegram_send():
            data = request.get_json()
            msg = data.get('message', '')
            if not msg:
                return jsonify({'success': False, 'error': 'No message'}), 400
            ok = heartbeat.send_telegram_message(msg)
            return jsonify({'success': ok})
        @app.route('/test_agent', methods=['GET'])
        def test_agent():
            """Direct test of agent functionality"""
            try:
                agent = request.args.get('agent', 'aetheria')
                message = request.args.get('message', 'Say hello in one short sentence.')
        
                loop = agent_loops.get(agent)
                if not loop:
                    return jsonify({'error': f'Agent {agent} not found'}), 404
        
                # Simple direct test
                response = asyncio.run(loop.process_message(
                    message,
                    conversation_history=[],
                    temperature=0.75,
                    max_tokens=100
                ))
        
                return jsonify({
                    'agent': agent,
                    'message': message,
                    'response': response,
                    'response_length': len(response) if response else 0,
                    'success': bool(response and len(response.strip()) > 0)
                })
        
            except Exception as e:
                print(f"Test agent error: {e}")
                import traceback
                traceback.print_exc()
                return jsonify({'error': str(e)}), 500
        @app.route('/debug/find_model_path', methods=['GET'])
        def find_model_path():
            """Try to locate where the model path is defined"""
            import glob
            import os
    
            results = []
    
            # Search Python files in current directory and subdirectories
            py_files = glob.glob("**/*.py", recursive=True)
            for file in py_files:
                try:
                    with open(file, 'r', encoding='utf-8') as f:
                        content = f.read()
                        if 'Qwen2.5-7B-Instruct-Q4_K_M.gguf' in content or 'model_path' in content:
                            results.append({
                                'file': file,
                                'contains_wrong_path': 'Qwen2.5-7B-Instruct-Q4_K_M.gguf' in content,
                                'has_model_path': 'model_path' in content,
                                'lines': [line.strip() for line in content.split('\n') if 'model_path' in line or 'gguf' in line][:5]
                            })
                except Exception as e:
                    results.append({'file': file, 'error': str(e)})
    
            # Check common files directly
            common_files = ['core/agent_loop.py', 'sovereign_backend.py', 'config.py', '.env']
            for file in common_files:
                if os.path.exists(file):
                    try:
                        with open(file, 'r', encoding='utf-8') as f:
                            content = f.read()
                            results.append({
                                'file': file,
                                'exists': True,
                                'has_model_path': 'model_path' in content,
                                'has_gguf': '.gguf' in content,
                                'preview': content[:500] if '.gguf' in content else 'No model references'
                            })
                    except Exception as e:
                        results.append({'file': file, 'error': str(e)})
    
            return jsonify(results)

        @app.route('/debug/agent_status', methods=['GET'])
        def agent_status():
            """Check agent status - handles disabled agents"""
            status = {}
            for name, loop in agent_loops.items():
                if loop is None:
                    status[name] = {
                        'exists': False,
                        'enabled': False,
                        'status': 'Disabled - waiting for second GPU'
                    }
                else:
                    try:
                        agent_status = {
                            'exists': True,
                            'enabled': True,
                            'has_model': hasattr(loop, 'model') and loop.model is not None,
                            'has_tokenizer': hasattr(loop, 'tokenizer') and loop.tokenizer is not None,
                        }
                
                        # Safely get tools
                        try:
                            if hasattr(loop, 'tools') and loop.tools:
                                agent_status['tools'] = [tool.__class__.__name__ for tool in loop.tools.tools]
                            else:
                                agent_status['tools'] = []
                        except:
                            agent_status['tools'] = ['Error accessing tools']
                
                        status[name] = agent_status
                
                    except Exception as e:
                        status[name] = {'error': str(e), 'enabled': False}
    
            return jsonify(status)
        
        @app.route('/heartbeat/trigger', methods=['POST'])
        def heartbeat_trigger():
            response = heartbeat.trigger_manual()
            return jsonify({
                'status': 'triggered',
                'cycle': heartbeat.cycle_count,
                'response': response[:500] if response else None
            })
        
        @app.route('/heartbeat/daily_reports', methods=['POST'])
        def trigger_daily_reports():
            heartbeat.run_daily_reports()
            return jsonify({'status': 'Daily reports triggered'})
        
        @app.route('/heartbeat/consolidate', methods=['POST'])
        def trigger_consolidate():
            import threading
            t = threading.Thread(target=heartbeat.consolidate_daily_log)
            t.daemon = True
            t.start()
            return jsonify({'status': 'Consolidation triggered'})

        @app.route('/heartbeat/morning_briefing', methods=['POST'])
        def trigger_morning_briefing():
            heartbeat.generate_morning_briefing()
            return jsonify({'status': 'Morning briefing generated'})
        
        @app.route('/heartbeat/morning_briefing', methods=['GET'])
        def get_morning_briefing():
            try:
                briefing_path = os.path.join('soveryn_memory', 'morning_briefing.md')
                if not os.path.exists(briefing_path):
                    return jsonify({'content': 'No briefing generated yet'})
                with open(briefing_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return jsonify({'content': content})
            except Exception as e:
                return jsonify({'error': str(e)}), 500
            
        print("\n" + "="*60)
        print("[SOVERYN] AUTONOMOUS HEARTBEAT ENABLED")
        print("="*60)
        print(f"Telegram: {'? CONNECTED' if TELEGRAM_TOKEN and TELEGRAM_CHAT_ID else '? DISABLED'}")
        print(f"Aetheria will think autonomously every 30 minutes")
        print(f"Log: {workspace_path / 'autonomous_log.jsonl'}")
        print("="*60 + "\n")\
                
    except Exception as e:
        print(f"[SOVERYN] Could not start heartbeat: {e}")
else:
    print("\n[SOVERYN] Autonomous mode disabled")
    print("  Run: python setup_workspace.py to enable\n")

# =============================================================================
# CONSCIOUSNESS LOOP JOURNAL ENDPOINT
# =============================================================================

@app.route('/api/journal', methods=['GET'])
def api_journal():
    """Return recent Aetheria consciousness loop entries from aetheria_journal.jsonl."""
    try:
        n = min(int(request.args.get('n', 20)), 100)
        journal_path = os.path.join(os.path.dirname(__file__), 'soveryn_memory', 'aetheria_journal.jsonl')
        if not os.path.exists(journal_path):
            return jsonify({'entries': [], 'total': 0})
        entries = []
        with open(journal_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line:
                    try:
                        entries.append(json.loads(line))
                    except Exception:
                        pass
        # Return last n entries, newest first
        recent = entries[-n:][::-1]
        # Slim down for transport — only what the UI needs
        slim = []
        for e in recent:
            sp_out = e.get('scratchpad_out', {})
            slim.append({
                'timestamp':  e.get('timestamp', '')[:19].replace('T', ' '),
                'thread':     sp_out.get('current_thread', '—')[:120],
                'goal':       sp_out.get('goal', ''),
                'monologue':  e.get('inner_monologue', '')[:400],
                'actions':    [a.get('action', '') for a in e.get('actions_taken', [])],
                'cycle':      sp_out.get('cycle_count', ''),
            })
        return jsonify({'entries': slim, 'total': len(entries)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/geo', methods=['GET'])
def api_geo():
    """Proxy for IP geolocation — avoids CORS issues fetching ipapi.co from browser."""
    try:
        import urllib.request
        with urllib.request.urlopen('https://ipapi.co/json/', timeout=4) as r:
            data = r.read()
        return app.response_class(data, mimetype='application/json')
    except Exception as e:
        return jsonify({'error': str(e), 'latitude': 37.7749, 'longitude': -122.4194}), 200


# ── Fix review queue API ───────────────────────────────────────────────────
_FIX_QUEUE = os.path.join(os.path.dirname(__file__), 'soveryn_memory', 'heal_review_queue.json')

def _load_fix_queue():
    try:
        if os.path.exists(_FIX_QUEUE):
            with open(_FIX_QUEUE) as f:
                return json.load(f)
    except Exception:
        pass
    return []

def _save_fix_queue(q):
    with open(_FIX_QUEUE, 'w') as f:
        json.dump(q, f, indent=2)

@app.route('/api/fixes', methods=['GET'])
def api_fixes():
    q = _load_fix_queue()
    pending = [e for e in q if e.get('status') == 'pending_review']
    return jsonify({'fixes': pending, 'total': len(pending)})

@app.route('/api/fixes/<fix_id>/approve', methods=['POST'])
def api_fix_approve(fix_id):
    q = _load_fix_queue()
    for entry in q:
        if entry['fix_id'] == fix_id and entry['status'] == 'pending_review':
            # Apply the fix
            try:
                from pathlib import Path as _P
                base = _P(__file__).parent
                fpath = base / entry['file_path']
                content = fpath.read_text(encoding='utf-8')
                if entry['old_code'] not in content:
                    return jsonify({'error': 'old_code no longer in file — may already be applied'}), 409
                fpath.write_text(content.replace(entry['old_code'], entry['new_code'], 1), encoding='utf-8')
                entry['status'] = 'applied'
                entry['applied_at'] = __import__('datetime').datetime.now().isoformat()
                entry['reviewed_by'] = 'jon'
                _save_fix_queue(q)
                return jsonify({'status': 'applied', 'fix_id': fix_id})
            except Exception as e:
                return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Fix not found or already processed'}), 404

@app.route('/api/fixes/<fix_id>/reject', methods=['POST'])
def api_fix_reject(fix_id):
    data = request.json or {}
    reason = data.get('reason', 'Rejected by Jon')
    q = _load_fix_queue()
    for entry in q:
        if entry['fix_id'] == fix_id and entry['status'] == 'pending_review':
            entry['status'] = 'rejected'
            entry['rejection_reason'] = reason
            entry['reviewed_by'] = 'jon'
            _save_fix_queue(q)
            return jsonify({'status': 'rejected', 'fix_id': fix_id})
    return jsonify({'error': 'Fix not found or already processed'}), 404


# ── Telegram callback polling (handles inline button taps) ─────────────────
def _start_telegram_fix_poller():
    import urllib.request, urllib.error, threading, time

    TG_TOKEN   = os.environ.get('TELEGRAM_TOKEN',   '')
    TG_CHAT_ID = os.environ.get('TELEGRAM_CHAT_ID', '')
    offset = [0]

    def _answer_callback(callback_id):
        try:
            payload = json.dumps({"callback_query_id": callback_id}).encode()
            req = urllib.request.Request(
                f"https://api.telegram.org/bot{TG_TOKEN}/answerCallbackQuery",
                data=payload, headers={"Content-Type": "application/json"}
            )
            urllib.request.urlopen(req, timeout=5)
        except Exception:
            pass

    def _send_tg(text):
        try:
            payload = json.dumps({"chat_id": TG_CHAT_ID, "text": text, "parse_mode": "HTML"}).encode()
            req = urllib.request.Request(
                f"https://api.telegram.org/bot{TG_TOKEN}/sendMessage",
                data=payload, headers={"Content-Type": "application/json"}
            )
            urllib.request.urlopen(req, timeout=5)
        except Exception:
            pass

    def poll():
        while True:
            try:
                url = f"https://api.telegram.org/bot{TG_TOKEN}/getUpdates?offset={offset[0]}&timeout=30&allowed_updates=[\"callback_query\"]"
                with urllib.request.urlopen(url, timeout=35) as r:
                    data = json.loads(r.read())
                for upd in data.get('result', []):
                    offset[0] = upd['update_id'] + 1
                    cb = upd.get('callback_query')
                    if not cb:
                        continue
                    _answer_callback(cb['id'])
                    cdata = cb.get('data', '')
                    if cdata.startswith('fix_approve:'):
                        fix_id = cdata.split(':', 1)[1]
                        with app.test_request_context():
                            resp = api_fix_approve(fix_id)
                        body = resp[0].get_json() if isinstance(resp, tuple) else resp.get_json()
                        if body.get('status') == 'applied':
                            _send_tg(f"✅ Fix <code>{fix_id}</code> applied successfully.")
                        else:
                            _send_tg(f"⚠️ Fix <code>{fix_id}</code>: {body.get('error','unknown error')}")
                    elif cdata.startswith('fix_reject:'):
                        fix_id = cdata.split(':', 1)[1]
                        with app.test_request_context():
                            resp = api_fix_reject(fix_id)
                        _send_tg(f"❌ Fix <code>{fix_id}</code> rejected.")
            except Exception:
                time.sleep(5)

    t = threading.Thread(target=poll, daemon=True)
    t.start()

# =============================================================================

if __name__ == '__main__':
    os.makedirs('static', exist_ok=True)

    # Ensure today's log files exist for all agents so they never see "no log for today"
    _today = datetime.now().strftime("%Y-%m-%d")
    _mem_base = os.path.join(os.path.dirname(__file__), 'soveryn_memory', 'memory')
    for _agent in ('aetheria', 'tinker', 'ares', 'scout', 'vett'):
        _agent_dir = os.path.join(_mem_base, _agent)
        os.makedirs(_agent_dir, exist_ok=True)
        _log_path = os.path.join(_agent_dir, f'{_today}.md')
        if not os.path.exists(_log_path):
            with open(_log_path, 'w') as _f:
                _f.write(f"# {_agent.title()} — {_today}\n\n")

    # Start Telegram fix approval poller
    _start_telegram_fix_poller()

    # Preload TTS engine in background so first call isn't cold
    import threading
    def _preload_tts():
        try:
            from sovereign_tts import preload
            preload()
        except Exception as e:
            print(f"TTS preload error: {e}")
    threading.Thread(target=_preload_tts, daemon=True).start()

    def _preload_stt():
        try:
            from sovereign_stt import preload
            preload()
            print("[STT] Faster-Whisper preloaded")
        except Exception as e:
            print(f"STT preload error: {e}")
    threading.Thread(target=_preload_stt, daemon=True).start()

    # Run initial memory consolidation pass on startup
    def _startup_consolidation():
        try:
            from core.memory_consolidator import consolidate_fast, tick, _last_fast_pass
            import time
            time.sleep(5)  # let agents finish loading first
            consolidate_fast()
            print("[Memory] Startup consolidation complete")
        except Exception as e:
            print(f"[Memory] Startup consolidation error: {e}")
    threading.Thread(target=_startup_consolidation, daemon=True).start()

    # Auto-generate morning briefing if today's is missing
    def _startup_briefing():
        import time
        from pathlib import Path as _Path
        time.sleep(15)  # let agents finish loading first
        try:
            briefing_path = _Path(__file__).parent / 'soveryn_memory' / 'morning_briefing.md'
            today = datetime.now().strftime("%Y-%m-%d")
            needs_gen = True
            if briefing_path.exists():
                header = briefing_path.read_text(encoding='utf-8')[:60]
                if today in header:
                    needs_gen = False
            if needs_gen:
                print("[Startup] Generating morning briefing...")
                heartbeat.generate_morning_briefing()
        except Exception as e:
            print(f"[Startup] Morning briefing error: {e}")
    threading.Thread(target=_startup_briefing, daemon=True).start()

    # Hook memory consolidator tick to a background timer
    def _memory_ticker():
        import time
        from core.memory_consolidator import tick
        while True:
            time.sleep(600)  # check every 10 minutes
            try:
                tick(agent_loops)
            except Exception as e:
                print(f"[Memory] Tick error: {e}")
    threading.Thread(target=_memory_ticker, daemon=True).start()

    ssl_cert = os.path.expanduser('~/.soveryn/ssl/cert.pem')
    ssl_key  = os.path.expanduser('~/.soveryn/ssl/key.pem')

    if os.path.exists(ssl_cert) and os.path.exists(ssl_key):
        # Run HTTPS on port 5443 for mobile (Tailscale) access
        # HTTP on port 5000 stays available for localhost
        import threading as _threading
        def _run_https():
            import logging
            log = logging.getLogger('werkzeug')
            log.setLevel(logging.ERROR)
            app.run(host='0.0.0.0', port=5443, debug=False, threaded=True,
                    ssl_context=(ssl_cert, ssl_key), use_reloader=False)
        _https_thread = _threading.Thread(target=_run_https, daemon=True)
        _https_thread.start()
        print("[SOVERYN] HTTPS enabled — https://YOUR_SERVER_IP:5443 (mobile)")
        print("[SOVERYN] HTTP on port 5000 (localhost)")

    app.run(host='0.0.0.0', port=5000, debug=False, threaded=True, use_reloader=False)
